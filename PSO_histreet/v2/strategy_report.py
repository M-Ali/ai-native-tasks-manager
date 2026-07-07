"""
PSO Lubricants — Pricing Strategy Word Report
Reads PSO_Pricing_Strategy.csv and generates a fully formatted Word document
covering all 8 frameworks with SKU tables, findings, and strategic narrative.
"""
from pathlib import Path
from datetime import datetime
import pandas as pd
import numpy as np

from docx import Document
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CSV_IN  = Path("output/reports/PSO_Pricing_Strategy.csv")
OUT     = Path("output/reports/PSO_Pricing_Strategy_Report.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

# ── Brand colours ─────────────────────────────────────────────────
PSO_GREEN   = RGBColor(0x00, 0x70, 0x3C)
PSO_LIGHT   = RGBColor(0xE8, 0xF5, 0xEE)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
DARK        = RGBColor(0x1A, 0x1A, 0x2E)
AMBER       = RGBColor(0xFF, 0xF3, 0xCD)
AMBER_DARK  = RGBColor(0xFF, 0xC1, 0x07)
RED_LIGHT   = RGBColor(0xF8, 0xD7, 0xDA)
GREEN_LIGHT = RGBColor(0xD4, 0xED, 0xDA)
BLUE_LIGHT  = RGBColor(0xD1, 0xEC, 0xF1)
GREY_LIGHT  = RGBColor(0xF2, 0xF2, 0xF2)
GREY_MID    = RGBColor(0xC0, 0xC0, 0xC0)
TIER_COLORS = {
    "Super Premium": RGBColor(0x1A, 0x53, 0x76),
    "Premium":       PSO_GREEN,
    "Mainstream":    RGBColor(0xE6, 0x7E, 0x22),
    "Economy":       RGBColor(0x7F, 0x8C, 0x8D),
}


# ── Document helpers ──────────────────────────────────────────────

def _rgb_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _rgb_hex(rgb))
    tcPr.append(shd)


def _set_cell_text(cell, text: str, bold=False, size=9.5,
                   color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text) if text is not None else "—")
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    c = color or PSO_GREEN
    for run in h.runs:
        run.font.color.rgb = c
    h.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    h.paragraph_format.space_after  = Pt(4)
    return h


def _para(doc, text, bold=False, size=11, color=None,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _bullet(doc, text, level=0, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def _callout(doc, text, bg=BLUE_LIGHT, bold=False, size=10.5):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, bg)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _divider(doc, color="00703C"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


def _fmt(val, prefix="Rs ", suffix="/L", decimals=0):
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return "—"
    return f"{prefix}{val:,.{decimals}f}{suffix}"


def _signal_color(signal: str) -> RGBColor:
    if signal == "PREMIUM":
        return RED_LIGHT
    if signal == "VALUE POSITION":
        return GREEN_LIGHT
    return BLUE_LIGHT


# ── Cover page ────────────────────────────────────────────────────

def _cover(doc, df):
    _para(doc, "PSO LUBRICANTS", bold=True, size=28,
          color=PSO_GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(doc, "Competitive Pricing Strategy Report", bold=True, size=17,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    _para(doc, "8-Framework SKU-Level Analysis", size=13,
          color=RGBColor(0x60, 0x60, 0x60), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _divider(doc)
    _para(doc, f"Date: {datetime.now().strftime('%d %B %Y')}",
          size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
          color=RGBColor(0x80, 0x80, 0x80), space_after=2)
    _para(doc, f"SKUs Analysed: {len(df)}  |  Brands Covered: {df['Brand'].nunique()}  |  "
               f"Grades: {df['Grade'].nunique()}  |  Market: Daraz.pk",
          size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
          color=RGBColor(0x80, 0x80, 0x80), space_after=4)
    doc.add_page_break()


# ── Executive summary ─────────────────────────────────────────────

def _exec_summary(doc, df):
    _heading(doc, "1. Executive Summary")
    _divider(doc)

    premium  = df[df["Market_Signal"] == "PREMIUM"]
    at_mkt   = df[df["Market_Signal"] == "AT MARKET"]
    value    = df[df["Market_Signal"] == "VALUE POSITION"]
    no_data  = df[df["Market_Signal"] == "INSUFFICIENT DATA"]

    _callout(doc, (
        f"This report applies 8 international pricing frameworks to {len(df)} PSO Carient / DEO / "
        f"Blaze SKUs across {df['Grade'].nunique()} grades and {df['Pack_L'].nunique()} pack sizes. "
        f"Competitor benchmarks are sourced from live Daraz.pk listings "
        f"({df['Mkt_Listing_Count'].sum():.0f} market observations across "
        f"Shell, Caltex, Total, Aramco, ZIC, and Kixx)."
    ), bg=BLUE_LIGHT)

    _para(doc, "Portfolio Pricing Position (vs Market Median):", bold=True, size=11)
    findings = [
        (f"{len(premium)} SKUs ABOVE market median (>+10%): "
         f"{', '.join(premium['Brand'].str.replace('PSO ','') + ' ' + premium['Grade'] + ' ' + premium['Pack_L'].astype(str) + 'L')}",
         RED_LIGHT),
        (f"{len(at_mkt)} SKUs AT MARKET (within ±10%): broadly well-positioned.", BLUE_LIGHT),
        (f"{len(value)} SKUs BELOW market median (>-10%): margin being left on the table — "
         f"scope to increase without losing volume.",  GREEN_LIGHT),
    ]
    for text, color in findings:
        _callout(doc, text, bg=color, size=10.5)

    _para(doc, "Strategic Priorities:", bold=True, size=11, space_after=3)
    priorities = [
        "Carient 5W-30 (Ultra + FS): priced 100-128% above market median — "
        "either reduce to capture volume or invest in spec communication to justify premium.",
        "Carient Plus, SPRO, Blaze 4T (10W-40, 20W-50): priced 18-28% below market — "
        "immediate opportunity to recover Rs 100-300/L without competitive risk.",
        "DEO bulk (15W-40 10L): value-positioned at -14% vs market — "
        "fleet channel pricing may be appropriate but confirm it clears the F4 waterfall floor.",
        "0W-20 (Ultra): at market. Fastest-growing grade globally. Hold price, invest in visibility.",
        "Carient FS 10W-40 4L and DEO 5000 10W-40 4L: showing as PREMIUM signal — "
        "verify pack size mapping in competitor data before raising.",
    ]
    for p in priorities:
        _bullet(doc, p)
    doc.add_page_break()


# ── Framework overview ─────────────────────────────────────────────

FRAMEWORK_DESCRIPTIONS = [
    ("F1", "Value-Based Tiering (VBT)",
     "Shell / Castrol model",
     "Maps each Carient variant to its Shell tier equivalent (Ultra→Helix Ultra, "
     "FS→HX8/HX7, Plus→HX5, SPRO→HX3). Applies a PSO brand discount (2–8%) reflecting "
     "current brand perception gap vs Shell. This gap should narrow as Carient equity grows.",
     "30%"),
    ("F2", "Competitive Reference Pricing (CRP)",
     "Market median anchor",
     "Prices 3% below the Daraz.pk median for each grade × pack combination. "
     "Positions PSO as best-value-in-class without being cheapest. "
     "10% band (±) defines the AT MARKET zone.",
     "30%"),
    ("F3", "Price-Pack Architecture (PPA)",
     "Unilever / P&G model",
     "Derives per-litre price for each pack size from the 4L market anchor. "
     "1L pack commands +40% per-litre premium; 10L carries -12% discount. "
     "Ensures the size-price ladder is legible and consistent across the range.",
     "20%"),
    ("F4", "McKinsey Price Waterfall",
     "Hard floor / margin protection",
     "Builds list price from cost upward: base oil import + additive package + packaging + "
     "target pocket margin. Fleet channel (-20%) is the deepest discount. "
     "Any recommended price below this floor is overridden upward. Acts as a hard constraint.",
     "10% (hard floor)"),
    ("F5", "OEM / API Spec Premium",
     "International OEM approval model",
     "Prices the specification uplift over an economy baseline. API SP (Carient Ultra) "
     "justifies +28% vs economy grade; API SN+ (FS) +18%; API SN (Plus) +8%. "
     "Premium is only realised if spec is communicated clearly at point of sale.",
     "10%"),
    ("F6", "Geographic Segmentation",
     "Zone-based price adjustment",
     "Applies regional growth signals from PSO's own FY24-25 data. "
     "Central DEO volume +18.6% → 5% regional premium. North MCO volume -27.3% → "
     "8% promotional discount. Applied as quarterly overlays, not permanent list changes.",
     "Overlay (informational)"),
    ("F7", "Base Oil Index Floor (BOI)",
     "ExxonMobil / TotalEnergies cost-index model",
     "Minimum price derived from base oil import cost × additive × 2.8x markup. "
     "0W-20 / 5W-30 / 5W-40 (Group III): BOI floor ~Rs 2,800–3,000/L. "
     "15W-40 (Group II): ~Rs 1,600/L. 20W-50 (Group I): ~Rs 900/L.",
     "Hard floor (informational)"),
    ("F8", "Channel Pricing",
     "Retail / workshop / fleet / distributor overlay",
     "Derives channel-specific prices from the final recommended list price. "
     "Retail pump: full list. Workshop: -12%. Distributor: -15%. Fleet contract: -20%. "
     "Ensures each channel generates an acceptable pocket margin.",
     "Overlay (applied last)"),
]

def _framework_overview(doc):
    _heading(doc, "2. Pricing Frameworks Applied")
    _divider(doc)
    _para(doc,
          "Eight internationally proven frameworks are applied to every PSO SKU. "
          "F1 and F2 carry the highest weights in the final synthesis (30% each) as they "
          "reflect actual market evidence. F4 and F7 act as hard floors — if the weighted "
          "average falls below either, the recommendation is raised to the floor.",
          size=10.5, space_after=8)

    tbl = doc.add_table(rows=1 + len(FRAMEWORK_DESCRIPTIONS), cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Code", "Framework", "Origin Model", "Logic", "Weight in Synthesis"]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        _set_cell_bg(cell, PSO_GREEN)
        _set_cell_text(cell, h, bold=True, color=WHITE, size=9)

    widths = [Inches(0.4), Inches(1.4), Inches(1.2), Inches(2.8), Inches(0.8)]
    for i, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[i].width = w

    alt = False
    for r_idx, (code, name, origin, logic, weight) in enumerate(FRAMEWORK_DESCRIPTIONS):
        bg = GREY_LIGHT if alt else WHITE
        alt = not alt
        row = tbl.rows[r_idx + 1]
        data = [code, name, origin, logic, weight]
        for c_idx, val in enumerate(data):
            _set_cell_bg(row.cells[c_idx], bg)
            _set_cell_text(row.cells[c_idx], val, size=9,
                           bold=(c_idx == 0),
                           color=PSO_GREEN if c_idx == 0 else DARK)

    doc.add_paragraph()
    doc.add_page_break()


# ── SKU master table ───────────────────────────────────────────────

def _sku_master_table(doc, df):
    _heading(doc, "3. SKU Pricing Recommendation Matrix")
    _divider(doc)
    _para(doc,
          "Final recommended price per litre and per pack (PKR) for each PSO SKU, "
          "with all 8 framework prices shown. Highlighted cells indicate the binding constraint "
          "or dominant framework for each SKU.",
          size=10.5, space_after=6)

    # --- Summary table (compact) ---
    _heading(doc, "3.1  Recommended Prices — All SKUs", level=2)

    cols = ["Brand", "Grade", "Pack", "F1 VBT", "F2 CRP", "F3 PPA",
            "F4 Floor", "F5 Spec", "Rec /L", "Rec PKR", "vs Mkt", "Signal"]
    tbl = doc.add_table(rows=1 + len(df), cols=len(cols))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(cols):
        cell = tbl.cell(0, i)
        _set_cell_bg(cell, PSO_GREEN)
        _set_cell_text(cell, h, bold=True, color=WHITE, size=8.5,
                       align=WD_ALIGN_PARAGRAPH.CENTER)

    for r_idx, (_, row) in enumerate(df.iterrows()):
        signal = str(row.get("Market_Signal", ""))
        bg = _signal_color(signal)
        tr = tbl.rows[r_idx + 1]

        def cell_val(col_idx, val, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
            _set_cell_bg(tr.cells[col_idx], bg)
            _set_cell_text(tr.cells[col_idx], val, bold=bold, size=8.5, align=align)

        cell_val(0,  str(row["Brand"]).replace("PSO ", ""), bold=True,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        cell_val(1,  str(row["Grade"]))
        cell_val(2,  f"{row['Pack_L']:.0f}L")
        cell_val(3,  _fmt(row.get("F1_VBT_Price_L"), "Rs ", "/L"))
        cell_val(4,  _fmt(row.get("F2_CRP_Price_L"), "Rs ", "/L"))
        cell_val(5,  _fmt(row.get("F3_PPA_Price_L"), "Rs ", "/L"))
        cell_val(6,  _fmt(row.get("F4_Waterfall_Floor_L"), "Rs ", "/L"))
        cell_val(7,  _fmt(row.get("F5_Spec_Price_L"), "Rs ", "/L"))
        cell_val(8,  _fmt(row.get("Recommended_Price_L"), "Rs ", "/L"), bold=True)
        cell_val(9,  _fmt(row.get("Recommended_Price_PKR"), "Rs ", ""), bold=True)
        vm = row.get("Rec_vs_Market_Med_Pct")
        cell_val(10, f"{vm:+.1f}%" if pd.notna(vm) else "—")
        cell_val(11, signal)

    doc.add_paragraph()

    # Legend
    legend_items = [
        (RED_LIGHT,   "PREMIUM — priced >10% above market; justify via spec or reduce"),
        (BLUE_LIGHT,  "AT MARKET — within ±10% of market median; competitive"),
        (GREEN_LIGHT, "VALUE POSITION — priced >10% below market; increase to capture margin"),
    ]
    for bg, text in legend_items:
        tbl_l = doc.add_table(rows=1, cols=1)
        tbl_l.style = "Table Grid"
        _set_cell_bg(tbl_l.cell(0, 0), bg)
        _set_cell_text(tbl_l.cell(0, 0), text, size=9)
    doc.add_paragraph()
    doc.add_page_break()


# ── Per-framework deep-dive sections ─────────────────────────────

def _framework_section(doc, df, code, title, description, col_price, col_rationale,
                        section_num, insight_fn=None):
    _heading(doc, f"{section_num}. {code}: {title}")
    _divider(doc)
    _callout(doc, description, bg=BLUE_LIGHT, size=10.5)

    if insight_fn:
        insight_fn(doc, df)

    # Table
    tbl = doc.add_table(rows=1 + len(df), cols=4)
    tbl.style = "Table Grid"

    headers = ["Brand", "Grade / Pack", f"{code} Recommended Price/L", "Rationale"]
    widths  = [Inches(1.4), Inches(0.9), Inches(1.2), Inches(3.1)]
    for i, (h, w) in enumerate(zip(headers, widths)):
        cell = tbl.cell(0, i)
        _set_cell_bg(cell, PSO_GREEN)
        _set_cell_text(cell, h, bold=True, color=WHITE, size=9)
        for row in tbl.rows:
            row.cells[i].width = w

    alt = False
    for r_idx, (_, row) in enumerate(df.iterrows()):
        bg = GREY_LIGHT if alt else WHITE
        alt = not alt
        tr = tbl.rows[r_idx + 1]
        price = row.get(col_price)
        rat   = str(row.get(col_rationale, "")).strip()

        _set_cell_bg(tr.cells[0], bg)
        _set_cell_text(tr.cells[0], str(row["Brand"]).replace("PSO ", ""),
                       bold=True, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_bg(tr.cells[1], bg)
        _set_cell_text(tr.cells[1], f"{row['Grade']}  {row['Pack_L']:.0f}L",
                       size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_bg(tr.cells[2], bg)
        _set_cell_text(tr.cells[2], _fmt(price, "Rs ", "/L"),
                       bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_bg(tr.cells[3], bg)
        _set_cell_text(tr.cells[3], rat, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()
    doc.add_page_break()


# ── Framework-specific insight functions ──────────────────────────

def _f1_insights(doc, df):
    _para(doc, "Key Observations:", bold=True, size=10.5, space_after=3)
    _bullet(doc, "Carient Ultra (super premium): discounted 8% vs Shell Helix Ultra — "
                 "reduce discount to 4% once Carient Ultra achieves API SP certification visibility.")
    _bullet(doc, "Carient FS (premium): priced at Shell HX7/HX8 -6%. Correct positioning for "
                 "a domestic brand competing with an established international premium tier.")
    _bullet(doc, "Carient Plus / SPRO (mainstream/economy): discounts of 2-4% vs Shell equivalents "
                 "are narrow — appropriate given PSO's distribution strength and pump-channel access.")
    doc.add_paragraph()


def _f2_insights(doc, df):
    _para(doc, "Key Observations:", bold=True, size=10.5, space_after=3)
    outliers = df[abs(df["Rec_vs_Market_Med_Pct"].fillna(0)) > 20]
    for _, row in outliers.iterrows():
        vm = row.get("Rec_vs_Market_Med_Pct")
        if pd.isna(vm):
            continue
        direction = "above" if vm > 0 else "below"
        _bullet(doc, f"{row['Brand'].replace('PSO ','')} {row['Grade']} {row['Pack_L']:.0f}L: "
                     f"recommended price is {abs(vm):.0f}% {direction} market median. "
                     + ("Review spec communication strategy." if vm > 0 else
                        "Increase price — significant margin opportunity."))
    doc.add_paragraph()


def _f3_insights(doc, df):
    _para(doc, "Key Observations:", bold=True, size=10.5, space_after=3)
    _bullet(doc, "1L packs should carry a 40% per-litre premium over 4L. "
                 "If PSO's current 1L:4L ratio is below 1.3:1, the pack ladder is compressed "
                 "and workshop buyers have low incentive to move consumers to 4L purchases.")
    _bullet(doc, "10L DEO packs: PPA recommends -12% vs 4L anchor. Fleet buyers expect "
                 "bulk savings; maintaining a narrower gap will lose fleet contracts to ZIC/Kixx "
                 "who are more aggressively priced on bulk.")
    _bullet(doc, "Currently no 800ml or 1.5L pack data captured from Daraz — "
                 "if PSO offers these sizes, the PPA multipliers (1.55× and 1.30×) "
                 "should be applied to set their prices.")
    doc.add_paragraph()


def _f4_insights(doc, df):
    _para(doc, "Key Observations:", bold=True, size=10.5, space_after=3)
    _bullet(doc, "F4 is a hard floor, not a target. No SKU should be priced below "
                 "the waterfall floor — doing so means PSO is subsidising the fleet channel "
                 "from retail margins, which is unsustainable.")
    _bullet(doc, "Highest floor is on 0W-20 and 5W-30 grades (Group III base oil + high additive cost). "
                 "These are the grades most exposed to base oil import price volatility.")
    _bullet(doc, "Recommendation: PSO should implement a quarterly price review mechanism "
                 "tied to ICIS base oil index and PKR/USD exchange rate. A 10% base oil cost "
                 "increase should trigger a review within 30 days.")
    doc.add_paragraph()


def _f5_insights(doc, df):
    _para(doc, "Key Observations:", bold=True, size=10.5, space_after=3)
    _bullet(doc, "The API SP premium (+28%) is only justifiable if Carient Ultra's "
                 "API SP certification is prominently displayed on pack, at pump, and in workshop POS. "
                 "Without this communication, consumers will not pay the premium.")
    _bullet(doc, "Carient FS at +18% over economy baseline is consistent with international "
                 "semi-synthetic positioning. Ensure 'Semi-Synthetic' claim is on label.")
    _bullet(doc, "Spec premium realisation gap: the F5 price and F2 CRP price diverge significantly "
                 "for 5W-30 grades — indicating the market is not fully paying for spec. "
                 "This is a brand/communication problem, not a pricing problem.")
    doc.add_paragraph()


def _f6_insights(doc, df):
    _para(doc, "Key Observations:", bold=True, size=10.5, space_after=3)
    _bullet(doc, "Central region DEO: +18.6% volume growth YoY is a strong signal of fleet expansion "
                 "(construction, agriculture). PSO can apply a 5% premium in Central without "
                 "risking volume loss. Implement via distributor invoice pricing, not pack price change.")
    _bullet(doc, "North MCO: -27.3% decline suggests competitive pressure from ZIC/Kixx imports "
                 "and possibly informal market. A -8% promotional pricing campaign in North "
                 "is recommended for Q3 to defend volume before the competitive position worsens.")
    _bullet(doc, "South region: baseline pricing — no adjustment recommended. "
                 "Monitor for any volume softness in 20W-50 which may reflect vehicle fleet upgrade "
                 "(shift to 10W-40 multigrade).")
    doc.add_paragraph()


def _f7_insights(doc, df):
    _para(doc, "Key Observations:", bold=True, size=10.5, space_after=3)
    _bullet(doc, "Group III base oil (0W-20, 5W-30, 5W-40) has the highest cost floor "
                 "(Rs 2,800–3,000+/L). PSO must secure long-term Group III supply contracts "
                 "to protect margin as these grades grow.")
    _bullet(doc, "Group I base oil (20W-50) has a floor of ~Rs 900/L — consistent with "
                 "Carient SPRO's current positioning as an economy grade.")
    _bullet(doc, "If base oil import costs rise 15% (ICIS Group III spot has been volatile), "
                 "the Group III floor rises to ~Rs 3,200–3,500/L — which would compress "
                 "margins on 5W-30 and 5W-40 at current recommended prices. "
                 "Build a 10% base oil buffer into pricing reviews.")
    doc.add_paragraph()


def _f8_insights(doc, df):
    _para(doc, "Key Observations:", bold=True, size=10.5, space_after=3)
    _bullet(doc, "Workshop channel (-12%): PSO's authorised workshop programme is the primary "
                 "volume channel for PCMO. At 12% discount, workshop pocket price must still "
                 "clear the F4 waterfall floor — confirm this for all grades before finalising "
                 "workshop price lists.")
    _bullet(doc, "Fleet contracts (-20%): DEO 15W-40 and 10W-40 bulk. At -20%, the pocket "
                 "price should still exceed the BOI floor (F7). For DEO 3000 20W-50, this "
                 "is the tightest margin grade — do not go beyond 20% fleet discount.")
    _bullet(doc, "Distributor margin (-15%): PSO must ensure distributor NTP (net trade price) "
                 "is uniform across regions to prevent grey market arbitrage between "
                 "Central (premium) and North (discounted) zones.")
    doc.add_paragraph()


# ── Geographic table ──────────────────────────────────────────────

def _geo_section(doc, df, section_num):
    _heading(doc, f"{section_num}. F6: Geographic Segmentation")
    _divider(doc)
    _callout(doc, (
        "Regional price adjustments derived from PSO's own FY24-25 volume growth data "
        "(Lubes Data Final.xlsx). These are not permanent list price changes — apply as "
        "quarterly distributor invoice adjustments or promotional pricing."
    ), bg=BLUE_LIGHT)
    _f6_insights(doc, df)

    tbl = doc.add_table(rows=1 + len(df), cols=6)
    tbl.style = "Table Grid"
    headers = ["Brand", "Grade / Pack", "National /L", "South /L", "Central /L", "North /L"]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        _set_cell_bg(cell, PSO_GREEN)
        _set_cell_text(cell, h, bold=True, color=WHITE, size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)

    alt = False
    for r_idx, (_, row) in enumerate(df.iterrows()):
        bg = GREY_LIGHT if alt else WHITE
        alt = not alt
        tr = tbl.rows[r_idx + 1]
        vals = [
            str(row["Brand"]).replace("PSO ", ""),
            f"{row['Grade']}  {row['Pack_L']:.0f}L",
            _fmt(row.get("Recommended_Price_L"), "Rs ", "/L"),
            _fmt(row.get("F6_South_Price_L"),   "Rs ", "/L"),
            _fmt(row.get("F6_Central_Price_L"), "Rs ", "/L"),
            _fmt(row.get("F6_North_Price_L"),   "Rs ", "/L"),
        ]
        aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
        # Colour Central green (premium) and North amber (discount)
        cell_bgs = [bg, bg, bg, bg,
                    GREEN_LIGHT if row.get("F6_Central_Price_L") else bg,
                    AMBER        if row.get("F6_North_Price_L")   else bg]
        for c_idx, (val, align, cbg) in enumerate(zip(vals, aligns, cell_bgs)):
            _set_cell_bg(tr.cells[c_idx], cbg)
            _set_cell_text(tr.cells[c_idx], val, size=8.5,
                           bold=(c_idx == 0), align=align)

    doc.add_paragraph()
    doc.add_page_break()


# ── Channel pricing section ───────────────────────────────────────

def _channel_section(doc, df, section_num):
    _heading(doc, f"{section_num}. F8: Channel Pricing")
    _divider(doc)
    _callout(doc, (
        "Channel prices are derived by applying fixed discount rates to the final "
        "recommended list price. These represent the NET price PSO receives per litre "
        "in each channel — not what the end consumer pays."
    ), bg=BLUE_LIGHT)
    _f8_insights(doc, df)

    tbl = doc.add_table(rows=1 + len(df), cols=7)
    tbl.style = "Table Grid"
    headers = ["Brand", "Grade / Pack", "List /L",
               "Retail Pump", "Workshop\n(-12%)", "Distributor\n(-15%)", "Fleet\n(-20%)"]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        _set_cell_bg(cell, PSO_GREEN)
        _set_cell_text(cell, h, bold=True, color=WHITE, size=8.5,
                       align=WD_ALIGN_PARAGRAPH.CENTER)

    alt = False
    for r_idx, (_, row) in enumerate(df.iterrows()):
        bg = GREY_LIGHT if alt else WHITE
        alt = not alt
        tr = tbl.rows[r_idx + 1]
        vals = [
            str(row["Brand"]).replace("PSO ", ""),
            f"{row['Grade']}  {row['Pack_L']:.0f}L",
            _fmt(row.get("Recommended_Price_L"),   "Rs ", "/L"),
            _fmt(row.get("F8_Retail_Pump_L"),       "Rs ", "/L"),
            _fmt(row.get("F8_Workshop_L"),          "Rs ", "/L"),
            _fmt(row.get("F8_Distributor_L"),       "Rs ", "/L"),
            _fmt(row.get("F8_Fleet_Contract_L"),    "Rs ", "/L"),
        ]
        for c_idx, val in enumerate(vals):
            _set_cell_bg(tr.cells[c_idx], bg)
            _set_cell_text(tr.cells[c_idx], val, size=8.5,
                           bold=(c_idx in (0, 2)),
                           align=WD_ALIGN_PARAGRAPH.LEFT if c_idx < 2 else WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    doc.add_page_break()


# ── Strategic recommendations ──────────────────────────────────────

def _strategy_section(doc, df):
    _heading(doc, "12. Strategic Pricing Recommendations")
    _divider(doc)

    recs = [
        ("Immediate: Raise Carient Plus & SPRO prices",
         GREEN_LIGHT,
         "VALUE POSITION SKUs (Plus 10W-40, SPRO 20W-50, Blaze 4T) are priced 18-28% "
         "below market median. Increase by Rs 100-200/L in Q3. These grades have the "
         "least competitive risk due to PSO's pump-channel advantage and consumer loyalty. "
         "Estimated margin uplift: Rs 80-150M annually (assuming volumes hold)."),

        ("Short-term: Resolve Carient 5W-30 PREMIUM signal",
         RED_LIGHT,
         "Ultra and FS 5W-30 are priced 100-128% above market median. Two options: "
         "(a) Invest in API SP / ACEA spec communication to justify the premium — "
         "develop workshop technician education and pack callout; or "
         "(b) Reduce price to CRP-recommended level (Rs 2,000-2,300/L) and compete on "
         "volume. Option (a) preserves margin; Option (b) grows volume. Decide by brand strategy."),

        ("Short-term: Protect DEO bulk pricing floor",
         AMBER,
         "DEO 5000 15W-40 10L is at VALUE POSITION (-14% vs market). "
         "Fleet channel at -20% will bring pocket price close to the F4 floor. "
         "Apply the F4 waterfall check before signing any fleet contract below Rs 1,200/L. "
         "Negotiate fleet contracts in PKR with a base oil escalation clause."),

        ("Medium-term: Establish quarterly price review cycle",
         BLUE_LIGHT,
         "Pakistan's base oil import cost and PKR/USD rate are volatile. "
         "Implement a structured quarterly price review: "
         "(1) Pull ICIS Group II/III spot prices. "
         "(2) Run this pricing strategy script with updated competitor Daraz scrape. "
         "(3) Apply F4 BOI floor update. "
         "(4) Recommend list price adjustments to commercial team. "
         "This app supports that cycle — just run: uv run python main.py all"),

        ("Medium-term: Geographic pricing pilot in Central",
         GREEN_LIGHT,
         "Central DEO volume +18.6% YoY signals strong fleet demand. "
         "Pilot a 5% premium via Central distributor invoice pricing (not pack label change). "
         "If volume holds, make permanent. If volume drops, withdraw within one quarter. "
         "Target: additional Rs 40-60/L margin on ~15% of DEO Central volume."),

        ("Medium-term: North MCO promotional defence",
         AMBER,
         "MCO North declining -27.3% — early warning of competitive displacement by "
         "ZIC X7 and Kixx G1 (both priced aggressively on Daraz). "
         "Launch a 3-month promotional price (-8% = ~Rs 130/L discount on Blaze 4T) "
         "in North through PSO pump channel. Measure volume recovery. "
         "If Blaze Xtreme is in North, use it as a quality anchor while SPRO/4T defends volume."),

        ("Long-term: List Carient on Daraz",
         BLUE_LIGHT,
         "PSO Carient has no visible presence on Daraz as a first-party seller. "
         "All 'PSO' listings captured are third-party, often at distressed prices. "
         "Establishing an official Carient storefront on Daraz achieves three things: "
         "(1) Price anchor — the market sees PSO's intended list price. "
         "(2) Direct consumer channel — captures e-commerce-native buyers. "
         "(3) Competitive intelligence — Daraz provides real-time sell-through data."),
    ]

    for title, bg, body in recs:
        _callout(doc, f"{title}\n\n{body}", bg=bg, size=10.5)

    doc.add_page_break()


# ── Methodology & limitations ─────────────────────────────────────

def _methodology(doc):
    _heading(doc, "13. Methodology & Limitations")
    _divider(doc)
    _callout(doc, (
        "All competitor prices are scraped from Daraz.pk, Pakistan's largest e-commerce "
        "marketplace. Prices reflect active third-party seller listings and may include "
        "promotional, bulk-buy, or distressed prices. Median price per litre is used throughout "
        "to reduce outlier distortion. Grade and pack size are extracted via regex pattern "
        "matching from product titles — accuracy depends on seller listing quality."
    ), bg=AMBER)

    limits = [
        "No PSO internal list price data used — recommendations are market-derived only. "
        "Overlay with actual PSO price lists before finalising.",
        "Daraz prices may lag physical retail prices by days to weeks.",
        "Equimoli brand returned no Daraz listings — insufficient data for benchmarking.",
        "Aramco and Caltex Havoline Ultra / Pro-DS had limited listings — "
        "CRP benchmarks for those tiers rely on Shell/ZIC equivalents.",
        "Channel prices (F8) assume PSO's published discount schedule. "
        "If actual workshop or fleet discounts differ, recalculate from the F4 floor.",
        "Geographic adjustments (F6) are based on FY24-25 volume growth data. "
        "CY26 YTD signals should be incorporated once H1 data is available.",
        "Base oil cost estimates (F7) are indicative. Actual PSO procurement costs may differ.",
    ]
    for l in limits:
        _bullet(doc, l)


# ── Master generator ──────────────────────────────────────────────

def generate():
    if not CSV_IN.exists():
        print("ERROR: PSO_Pricing_Strategy.csv not found. Run pricing_strategy.py first.")
        return

    df = pd.read_csv(CSV_IN)
    print(f"Loaded {len(df)} SKUs from {CSV_IN.name}")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # Set default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    _cover(doc, df)
    _exec_summary(doc, df)
    _framework_overview(doc)
    _sku_master_table(doc, df)

    # Per-framework deep-dive sections
    _framework_section(
        doc, df, "F1", "Value-Based Tiering",
        "Maps each PSO Carient variant to its Shell tier equivalent and applies a PSO brand "
        "discount (2–8%). As Carient brand equity grows, this discount should narrow — "
        "the VBT price is the aspirational ceiling for each tier.",
        "F1_VBT_Price_L", "F1_VBT_Rationale", 5, _f1_insights)

    _framework_section(
        doc, df, "F2", "Competitive Reference Pricing (CRP)",
        "Prices 3% below the Daraz.pk market median per grade × pack. "
        "The ±10% band defines the AT MARKET zone. SKUs outside this band warrant action.",
        "F2_CRP_Price_L", "F2_CRP_Rationale", 6, _f2_insights)

    _framework_section(
        doc, df, "F3", "Price-Pack Architecture (PPA)",
        "Derives per-litre price for each pack size from the 4L market anchor. "
        "Ensures the pack-price ladder is consistent and legible to buyers.",
        "F3_PPA_Price_L", "F3_PPA_Rationale", 7, _f3_insights)

    _framework_section(
        doc, df, "F4", "Price Waterfall (Hard Floor)",
        "Minimum viable list price built from cost upward. Any final recommendation "
        "below this floor is automatically raised. Acts as a margin protection hard constraint.",
        "F4_Waterfall_Floor_L", "F4_Waterfall_Rationale", 8, _f4_insights)

    _framework_section(
        doc, df, "F5", "OEM / API Spec Premium",
        "Prices the specification uplift over an economy-tier baseline. "
        "The premium is only realised if the spec is communicated at point of sale.",
        "F5_Spec_Price_L", "F5_Spec_Rationale", 9, _f5_insights)

    _geo_section(doc, df, 10)

    _framework_section(
        doc, df, "F7", "Base Oil Index Floor (BOI)",
        "Minimum price derived from base oil import cost × additive × 2.8x markup. "
        "Serves as a second hard floor alongside F4.",
        "F7_BOI_Floor_L", "F7_BOI_Rationale", 11, _f7_insights)

    _channel_section(doc, df, 12)
    _strategy_section(doc, df)
    _methodology(doc)

    doc.save(OUT)
    print(f"\nReport saved: {OUT}")
    print(f"Pages: ~{len(df) // 3 + 20} (estimated)")


if __name__ == "__main__":
    generate()
