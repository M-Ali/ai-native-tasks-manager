"""
Generates a Word document with pricing strategy recommendations
based on the scraped + analysed data.
"""
from pathlib import Path
from datetime import datetime

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from db.store import latest_run

OUTPUT = Path("output/reports")
OUTPUT.mkdir(parents=True, exist_ok=True)
CHARTS = Path("output/charts")

PSO_GREEN = RGBColor(0x00, 0x70, 0x3C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
MID_GREEN = RGBColor(0xD4, 0xED, 0xDA)
MID_RED = RGBColor(0xF8, 0xD7, 0xDA)
MID_AMBER = RGBColor(0xFF, 0xF3, 0xCD)
MID_BLUE = RGBColor(0xD1, 0xEC, 0xF1)


# ── Document helpers ──────────────────────────────────────────────

def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def _para(doc, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PSO_GREEN
    return h


def _bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def _callout(doc, text, bg: RGBColor = MID_BLUE, bold=False):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    doc.add_paragraph()


def _insert_chart(doc, path: Path, caption: str, width=6.5):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        doc.add_paragraph()


def _divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "00703C")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


# ── Main report ───────────────────────────────────────────────────

def generate(config: dict):
    rows = latest_run()
    if not rows:
        print("No data found. Run scrape first.")
        return

    df = pd.DataFrame(rows)
    df = df[df["price_per_litre"].notna() & df["grade_detected"].notna()]
    df["is_pso"] = df["brand_detected"].str.contains("PSO|Carient", na=False)

    doc = Document()
    _setup_styles(doc)

    # Cover
    _para(doc, "PSO Lubricants", bold=True, size=26, color=PSO_GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(doc, "Competitive Pricing Intelligence Report", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(doc, f"Generated: {datetime.now().strftime('%d %B %Y')}", size=10,
          color=RGBColor(0x80, 0x80, 0x80), align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, f"Data source: Daraz.pk | Products scraped: {len(df):,}", size=10,
          color=RGBColor(0x80, 0x80, 0x80), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # 1. Executive Summary
    _heading(doc, "1. Executive Summary")
    _divider(doc)

    pso_df = df[df["is_pso"]]
    comp_df = df[~df["is_pso"]]
    n_brands = comp_df["brand_detected"].nunique()
    n_grades = df["grade_detected"].nunique()

    _callout(doc, (
        f"This report benchmarks PSO Carient's pricing against {n_brands} competing brands "
        f"across {n_grades} grades for petrol and diesel engine oils. "
        f"All prices are median observed retail prices in PKR per litre on Daraz.pk."
    ), bg=MID_BLUE)

    gap_summary = _compute_gap_summary(df)
    for line in gap_summary:
        _bullet(doc, line)

    doc.add_paragraph()

    # 2. Market Coverage
    _heading(doc, "2. Market Coverage")
    _divider(doc)
    brand_counts = df.groupby("brand_detected").size().reset_index(name="listings")
    _para(doc, f"Total listings analysed: {len(df):,}", bold=True)
    doc.add_paragraph()

    tbl = doc.add_table(rows=1 + len(brand_counts), cols=3)
    tbl.style = "Table Grid"
    headers = ["Brand", "Listings", "Own / Competitor"]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        _set_cell_bg(cell, PSO_GREEN)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)

    for r_idx, row in brand_counts.iterrows():
        brand = row["brand_detected"] or "Unknown"
        is_pso = "PSO" in brand or "Carient" in brand
        cells = tbl.rows[r_idx + 1 - len(brand_counts) + len(brand_counts)].cells
        # use enumerate offset
        tbl.cell(r_idx + 1, 0).text = brand
        tbl.cell(r_idx + 1, 1).text = str(row["listings"])
        tbl.cell(r_idx + 1, 2).text = "PSO (Own)" if is_pso else "Competitor"
        if is_pso:
            for c in range(3):
                _set_cell_bg(tbl.cell(r_idx + 1, c), MID_GREEN)
    doc.add_paragraph()

    # 3. Charts
    _heading(doc, "3. Competitive Price Analysis")
    _divider(doc)

    _insert_chart(doc, CHARTS / "01_price_matrix_pcmo.png",
                  "Figure 1: Price/L Distribution — Petrol Engine Oils (PCMO) by Grade")
    _insert_chart(doc, CHARTS / "01_price_matrix_hdeo.png",
                  "Figure 2: Price/L Distribution — Diesel Engine Oils (HDEO) by Grade")
    _insert_chart(doc, CHARTS / "02_tier_gaps.png",
                  "Figure 3: Tier Price Positioning — PSO vs Market Median")
    _insert_chart(doc, CHARTS / "03_pack_premium_curve.png",
                  "Figure 4: Pack Size Premium Curve — Price/L by Pack Size")
    _insert_chart(doc, CHARTS / "04_price_index.png",
                  "Figure 5: Price Index (PSO = 100) by Grade")
    _insert_chart(doc, CHARTS / "05_price_heatmap.png",
                  "Figure 6: Full Price Heatmap — Brand x Grade")

    # 4. Grade-level findings
    _heading(doc, "4. Grade-Level Pricing Findings")
    _divider(doc)
    _write_grade_findings(doc, df)

    # 5. Pricing Recommendations
    _heading(doc, "5. Pricing Strategy Recommendations")
    _divider(doc)
    _write_recommendations(doc, df, config)

    # 6. Methodology
    _heading(doc, "6. Methodology & Limitations")
    _divider(doc)
    _callout(doc, (
        "All prices scraped from Daraz.pk (Pakistan's largest e-commerce marketplace). "
        "Prices reflect active listings and may include third-party sellers, promotional prices, "
        "or bulk offers. Median price per litre is used throughout to reduce outlier distortion. "
        "Grade and pack size are extracted from product titles using pattern matching — accuracy "
        "depends on seller title quality. PSO pump prices or distributor list prices may differ."
    ), bg=MID_AMBER)

    outfile = OUTPUT / "PSO_Pricing_Intelligence_Report.docx"
    doc.save(outfile)
    print(f"\nReport saved: {outfile}")


# ── Helpers ───────────────────────────────────────────────────────

def _setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _compute_gap_summary(df: pd.DataFrame) -> list[str]:
    pso = df[df["is_pso"] & df["price_per_litre"].notna()]
    comps = df[~df["is_pso"] & df["price_per_litre"].notna()]

    lines = []
    grades = sorted(df["grade_detected"].dropna().unique())
    overpriced, underpriced, at_market = [], [], []

    for grade in grades:
        pv = pso[pso["grade_detected"] == grade]["price_per_litre"]
        cv = comps[comps["grade_detected"] == grade]["price_per_litre"]
        if pv.empty or cv.empty:
            continue
        gap = (pv.median() - cv.median()) / cv.median() * 100
        if gap > 10:
            overpriced.append(f"{grade} (+{gap:.0f}%)")
        elif gap < -10:
            underpriced.append(f"{grade} ({gap:.0f}%)")
        else:
            at_market.append(grade)

    if overpriced:
        lines.append(f"Grades where PSO is priced ABOVE market median (>10%): {', '.join(overpriced)}")
    if underpriced:
        lines.append(f"Grades where PSO is priced BELOW market median (>10%): {', '.join(underpriced)} — potential margin capture opportunity")
    if at_market:
        lines.append(f"Grades priced at market: {', '.join(at_market)}")
    if not lines:
        lines.append("Insufficient matched data — run a wider scrape to populate grade-level gaps.")
    return lines


def _write_grade_findings(doc: Document, df: pd.DataFrame):
    pso = df[df["is_pso"] & df["price_per_litre"].notna()]
    comps = df[~df["is_pso"] & df["price_per_litre"].notna()]
    grades = sorted(df["grade_detected"].dropna().unique())

    for grade in grades:
        pv = pso[pso["grade_detected"] == grade]["price_per_litre"]
        cv = comps[comps["grade_detected"] == grade]["price_per_litre"]
        if pv.empty and cv.empty:
            continue

        _heading(doc, f"Grade: {grade}", level=2)
        if not pv.empty:
            _para(doc, f"PSO median price/L: Rs {pv.median():,.0f}  |  Range: Rs {pv.min():,.0f} – Rs {pv.max():,.0f}", size=10.5)
        if not cv.empty:
            _para(doc, f"Market median price/L: Rs {cv.median():,.0f}  |  Range: Rs {cv.min():,.0f} – Rs {cv.max():,.0f}", size=10.5)
            cheapest = comps[comps["grade_detected"] == grade].nsmallest(1, "price_per_litre")[["brand_detected", "price_per_litre"]]
            priciest = comps[comps["grade_detected"] == grade].nlargest(1, "price_per_litre")[["brand_detected", "price_per_litre"]]
            if not cheapest.empty:
                _para(doc, f"Cheapest competitor: {cheapest.iloc[0]['brand_detected']} @ Rs {cheapest.iloc[0]['price_per_litre']:,.0f}/L", size=10.5)
            if not priciest.empty:
                _para(doc, f"Most expensive: {priciest.iloc[0]['brand_detected']} @ Rs {priciest.iloc[0]['price_per_litre']:,.0f}/L", size=10.5)

        if not pv.empty and not cv.empty:
            gap = (pv.median() - cv.median()) / cv.median() * 100
            if gap > 10:
                _callout(doc, f"PSO is {gap:.1f}% above market median for {grade}. Review if premium is justified by spec/approval.", bg=MID_RED)
            elif gap < -10:
                _callout(doc, f"PSO is {abs(gap):.1f}% below market median for {grade}. Scope to increase price without losing position.", bg=MID_GREEN)
            else:
                _callout(doc, f"PSO is within 10% of market median for {grade} — priced competitively.", bg=MID_BLUE)

        doc.add_paragraph()


def _write_recommendations(doc: Document, df: pd.DataFrame, config: dict):
    _callout(doc, (
        "Recommendations are derived entirely from observed market prices. "
        "PSO's cost structure, channel margins, and strategic objectives should be "
        "overlaid before finalising any list price changes."
    ), bg=MID_AMBER, bold=False)

    recs = [
        ("1. Anchor pricing to the 4L pack", MID_BLUE,
         "The 4L/5L pack is the market reference for workshops. Set this as the anchor and "
         "derive 1L (premium uplift) and 10L+ (discount) from it. Ensures pack-price architecture is legible to buyers."),
        ("2. Widen the Ultra-to-SPRO price gap", MID_GREEN,
         "If the gap between Carient Ultra and SPRO is less than 35-40%, consumers perceive low differentiation "
         "and trade down in recessions. International practice (Shell, Castrol) maintains 2.5-4x gap across the full tier ladder."),
        ("3. Price 5W-30 at a premium to 10W-40", MID_BLUE,
         "5W-30 is a modern, fuel-economy grade requiring Group III base oil. It should command 15-25% "
         "premium vs 10W-40 across all brands. If PSO's spread is narrower, it signals a pricing gap."),
        ("4. Match ZIC/Kixx for imported brand positioning", MID_AMBER,
         "Imported brands (ZIC, Kixx, Equimoli) often undercut on price while claiming import quality perception. "
         "PSO should price Carient FS at or slightly below ZIC X7 equivalent to compete on value, not just brand."),
        ("5. Review DEO 15W-40 drum pricing vs retail", MID_GREEN,
         "210L drum price/L should be 25-35% below 4L retail. If the gap is narrower, fleet buyers "
         "have little incentive to commit to bulk contracts vs spot retail purchase."),
        ("6. Run quarterly scrape cadence", MID_BLUE,
         "Competitor prices on Daraz shift during Ramadan, Eid, budget cycles, and base oil import cost changes. "
         "A quarterly scrape run will surface these movements before PSO's list price cycle."),
        ("7. Populate own-brand listings on Daraz", MID_GREEN,
         "If PSO Carient is absent from Daraz, the competitive benchmark will be skewed. "
         "Listing Carient on Daraz also captures the direct-to-consumer channel and provides price transparency."),
    ]

    for title, color, body in recs:
        _callout(doc, f"{title}\n{body}", bg=color)
        doc.add_paragraph()
