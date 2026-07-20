"""
PSO Retail — Premium Fuel (R95 vs PMG) Opportunity Report
Word document: national trend, margin case, region/city/station breakdown,
whitespace launch candidates, and uplift scenarios.

Usage:
    uv run python workspace/premium_fuel_report.py
"""
import sys, io, os
sys.path.insert(0, 'src')
sys.stdout.reconfigure(encoding='utf-8')

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml

from pso import ingest, premium_fuel_analyze as pf
from _pso_common import INPUT_PATH, get_period_label, out_path

# ── scenario parameters ───────────────────────────────────────────────────────
# Full potential = every whitespace (R95-less) PMG station reaches the R95
# penetration rate already observed at stations that DO carry R95 today.
CONSERV = 0.25   # realistic 12-month target
OPTIMAL = 0.55   # stretch 18-month target

# ── colours ───────────────────────────────────────────────────────────────────
PSO_BLUE  = RGBColor(0x00, 0x47, 0x9D)
PSO_GREEN = RGBColor(0x00, 0x8C, 0x4A)
PSO_RED   = RGBColor(0xC0, 0x00, 0x00)
PSO_ORANGE= RGBColor(0xE4, 0x6C, 0x0A)
LGREY     = RGBColor(0xF2, 0xF2, 0xF2)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

PLT_BLUE  = '#00479D'
PLT_ORANGE= '#E46C0A'
PLT_GREY  = '#BFBFBF'

# ── helpers (mirrors workspace/lubes_report.py style) ─────────────────────────
def ml(v):    return v / 1_000_000
def bn(v):    return v / 1_000_000_000
def pct(a, b): return (a - b) / abs(b) * 100 if b else 0
def chg_str(v): return f"+{v:.1f}%" if v >= 0 else f"{v:.1f}%"

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{hex_color}"/>'))

def cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16 if level == 1 else (13 if level == 2 else 11))
    r.font.color.rgb = PSO_BLUE if level <= 2 else RGBColor(0x40, 0x40, 0x40)
    if level <= 2:
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(
            '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="00479D"/></w:pBdr>'))
    return p

def add_body(doc, text, size=10, space_before=3, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for r in p.runs:
        r.font.size = Pt(size)
    return p

def img_buf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close(fig)
    return buf

# ── load data + analysis ──────────────────────────────────────────────────────
print("Loading data…")
df, _ = ingest.load(INPUT_PATH)
REPORT_DATE = get_period_label(df)
tables = pf.run_premium_fuel(df)

trend    = tables["premium_product_trend"]
decomp   = tables["premium_margin_decomp"]
by_region= tables["premium_by_region"]
by_city  = tables["premium_by_city"]
stn_mix  = tables["premium_station_mix"]
whitespace = tables["premium_whitespace_stations"]
growing  = tables["premium_growing_markets"]
declining= tables["premium_declining_markets"]
corp_seg = tables["premium_customer_segments"]

pmg_row = trend[trend["ProductCategory"] == "PMG"].iloc[0]
r95_row = trend[trend["ProductCategory"] == "R95"].iloc[0]
uplift_row = decomp.iloc[-1]  # "R95 minus PMG" row

# ── uplift scenario: convert whitespace PMG volume at observed R95 penetration ──
avg_r95_penetration = (stn_mix[stn_mix["R95_Vol_CY_ML"] > 0]["R95_Share_Pct"].mean()) / 100
full_potential_ml = (whitespace["PMG_Vol_CY_ML"] * avg_r95_penetration).sum()
conservative_ml = full_potential_ml * CONSERV
optimal_ml      = full_potential_ml * OPTIMAL
nmgn_per_ltr_r95 = r95_row["NMgn_per_Ltr_CY"]
conservative_nmgn_m = conservative_ml * 1e6 * nmgn_per_ltr_r95 / 1e6
optimal_nmgn_m      = optimal_ml * 1e6 * nmgn_per_ltr_r95 / 1e6

# ═══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("PSO Retail — Premium Fuel Opportunity")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = PSO_BLUE
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run(f"R95 (Premium) vs PMG (Regular) Petrol — Where to Launch/Expand  |  {REPORT_DATE}")
r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

# ── 1. Executive Summary ──────────────────────────────────────────────────────
add_heading(doc, "1. Executive Summary", level=1)
add_body(doc,
    f"PSO already operates a premium petrol variant — R95 — alongside regular PMG. Rather than "
    f"speculate about launching a brand-new high-end fuel, this report uses R95's existing footprint "
    f"as a live proof-of-concept: {r95_row['Rows']} of {int(pmg_row['Rows'])+int(r95_row['Rows'])} retail "
    f"petrol rows already carry R95, representing {r95_row['Vol_Share_Pct']:.1f}% of retail petrol volume "
    f"nationally ({ml(r95_row['Vol_CY_ML']*1e6):.1f}M litres CY).")
add_body(doc,
    f"R95 volume grew {chg_str(r95_row['Vol_Chg_Pct'])} vs SPLY — faster than PMG's "
    f"{chg_str(pmg_row['Vol_Chg_Pct'])} — while carrying PKR {uplift_row['NMgn_per_Ltr_CY']:.2f}/litre "
    f"more net margin than PMG. Demand for premium fuel is real, growing, and more profitable per litre; "
    f"the constraint is distribution, not demand. {len(whitespace)} active PMG stations currently carry "
    f"zero R95 — this is the direct launch/expansion target list (Section 5).")

# KPI table
kpi_tbl = doc.add_table(rows=2, cols=4)
kpi_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
kpi_heads = ["R95 Vol CY (ML)", "R95 Share of Petrol", "R95 Vol Growth (vs SPLY)", "R95 NMgn/Ltr Premium vs PMG"]
kpi_vals  = [f"{r95_row['Vol_CY_ML']:.1f}", f"{r95_row['Vol_Share_Pct']:.1f}%",
             chg_str(r95_row['Vol_Chg_Pct']), f"+PKR {uplift_row['NMgn_per_Ltr_CY']:.2f}"]
for ci, h in enumerate(kpi_heads):
    cell_text(kpi_tbl.cell(0, ci), h, bold=True, size=8.5, color=WHITE)
    set_cell_bg(kpi_tbl.cell(0, ci), "00479D")
for ci, v in enumerate(kpi_vals):
    cell_text(kpi_tbl.cell(1, ci), v, bold=True, size=13, color=PSO_BLUE)
    set_cell_bg(kpi_tbl.cell(1, ci), "F2F2F2")

# ── 2. National Trend Chart ────────────────────────────────────────────────────
add_heading(doc, "2. National Volume Trend — R95 vs PMG (CY vs SPLY)", level=1)
fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 3.6))
fig.patch.set_facecolor('white')

cats = ["PMG", "R95"]
cy_vals   = [pmg_row["Vol_CY_ML"], r95_row["Vol_CY_ML"]]
sply_vals = [pmg_row["Vol_SPLY_ML"], r95_row["Vol_SPLY_ML"]]
x = np.arange(len(cats)); bw = 0.35
ax1.bar(x - bw/2, cy_vals, bw, color=PLT_BLUE, label='CY', zorder=3)
ax1.bar(x + bw/2, sply_vals, bw, color=PLT_GREY, label='SPLY', zorder=3)
ax1.set_xticks(x); ax1.set_xticklabels(cats)
ax1.set_title('Volume (Million Litres) — CY vs SPLY', fontsize=10, fontweight='bold', color=PLT_BLUE)
ax1.legend(fontsize=8)
ax1.grid(axis='y', alpha=0.3, zorder=0)

nmgn_vals = [pmg_row["NMgn_per_Ltr_CY"], r95_row["NMgn_per_Ltr_CY"]]
bars = ax2.bar(cats, nmgn_vals, color=[PLT_BLUE, PLT_ORANGE], zorder=3)
ax2.set_title('Net Margin per Litre (PKR) — CY', fontsize=10, fontweight='bold', color=PLT_BLUE)
ax2.grid(axis='y', alpha=0.3, zorder=0)
for b, v in zip(bars, nmgn_vals):
    ax2.text(b.get_x() + b.get_width()/2, v, f"{v:.2f}", ha='center', va='bottom', fontsize=9)

doc.add_picture(img_buf(fig), width=Inches(6.8))

# ── 3. Margin Case ─────────────────────────────────────────────────────────────
add_heading(doc, "3. Unit Economics — Why R95 Is Worth Pushing", level=1)
add_body(doc,
    f"R95 carries PKR {uplift_row['PMgn_per_Ltr_CY']:.2f}/litre more primary margin and "
    f"PKR {uplift_row['NMgn_per_Ltr_CY']:.2f}/litre more net margin than PMG, despite "
    f"PKR {abs(uplift_row['Disc_per_Ltr_CY']):.2f}/litre {'less' if uplift_row['Disc_per_Ltr_CY'] < 0 else 'more'} "
    f"discounting. Every litre converted from PMG to R95 is a direct margin improvement, not just a mix shift.")

m_hdr = ["Product", "Vol CY (ML)", "PMgn/Ltr CY", "Disc/Ltr CY", "NMgn/Ltr CY", "NMgn/Ltr SPLY", "NMgn/Ltr Chg%"]
m_tbl = doc.add_table(rows=len(decomp) + 1, cols=len(m_hdr))
m_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(m_hdr):
    cell_text(m_tbl.cell(0, ci), h, bold=True, size=8.5, color=WHITE)
    set_cell_bg(m_tbl.cell(0, ci), "00479D")
for ri, (_, row) in enumerate(decomp.iterrows(), 1):
    vals = [
        row["Product"],
        f"{row['Vol_CY_ML']:.1f}" if pd.notna(row["Vol_CY_ML"]) else "—",
        f"{row['PMgn_per_Ltr_CY']:.2f}" if pd.notna(row["PMgn_per_Ltr_CY"]) else "—",
        f"{row['Disc_per_Ltr_CY']:.2f}" if pd.notna(row["Disc_per_Ltr_CY"]) else "—",
        f"{row['NMgn_per_Ltr_CY']:.2f}" if pd.notna(row["NMgn_per_Ltr_CY"]) else "—",
        f"{row['NMgn_per_Ltr_SPLY']:.2f}" if pd.notna(row["NMgn_per_Ltr_SPLY"]) else "—",
        f"{row['NMgn_per_Ltr_Chg_Pct']:.1f}%" if pd.notna(row["NMgn_per_Ltr_Chg_Pct"]) else "—",
    ]
    is_uplift = "minus" in str(row["Product"])
    bg = "FFF3E0" if is_uplift else ("F2F2F2" if ri % 2 == 0 else "FFFFFF")
    for ci, v in enumerate(vals):
        cell_text(m_tbl.cell(ri, ci), v, bold=is_uplift, size=8.5,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(m_tbl.cell(ri, ci), bg)

# ── 4. Region & City Breakdown ─────────────────────────────────────────────────
add_heading(doc, "4. Where R95 Already Sells — Region & City Penetration", level=1)
add_body(doc,
    "Regions and cities with the highest existing R95 share are the strongest proof-of-demand markets — "
    "premium fuel already works there. This is the base to defend and deepen, distinct from the "
    "whitespace expansion list in Section 5.")

r_hdr = ["Region", "PMG Vol (ML)", "R95 Vol (ML)", "Total Vol (ML)", "R95 Share%", "R95 Stations"]
r_tbl = doc.add_table(rows=len(by_region) + 1, cols=len(r_hdr))
r_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(r_hdr):
    cell_text(r_tbl.cell(0, ci), h, bold=True, size=8.5, color=WHITE)
    set_cell_bg(r_tbl.cell(0, ci), "00479D")
for ri, (_, row) in enumerate(by_region.iterrows(), 1):
    vals = [row["Sales office Region"], f"{row['PMG_Vol_CY_ML']:.1f}", f"{row['R95_Vol_CY_ML']:.1f}",
            f"{row['Total_Vol_CY_ML']:.1f}", f"{row['R95_Share_Pct']:.1f}%", str(int(row["R95_Stations"]))]
    bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
    for ci, v in enumerate(vals):
        cell_text(r_tbl.cell(ri, ci), v, size=8.5,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(r_tbl.cell(ri, ci), bg)

doc.add_paragraph()
add_body(doc, "Top 20 cities by R95 penetration (share of city's own petrol volume):", size=9.5, space_after=2)
top_cities = by_city.head(20)
c_hdr = ["City", "Region", "PMG Vol (ML)", "R95 Vol (ML)", "R95 Share%"]
c_tbl = doc.add_table(rows=len(top_cities) + 1, cols=len(c_hdr))
c_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(c_hdr):
    cell_text(c_tbl.cell(0, ci), h, bold=True, size=8, color=WHITE)
    set_cell_bg(c_tbl.cell(0, ci), "00479D")
for ri, (_, row) in enumerate(top_cities.iterrows(), 1):
    vals = [row["CityNorm"], row["Sales office Region"], f"{row['PMG_Vol_CY_ML']:.2f}",
            f"{row['R95_Vol_CY_ML']:.2f}", f"{row['R95_Share_Pct']:.1f}%"]
    bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
    for ci, v in enumerate(vals):
        cell_text(c_tbl.cell(ri, ci), v, size=8,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci <= 1 else WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(c_tbl.cell(ri, ci), bg)

# ── 5. Whitespace Launch Candidates ────────────────────────────────────────────
add_heading(doc, "5. Launch Candidates — Active PMG Stations With Zero R95", level=1)
add_body(doc,
    f"{len(whitespace)} of {len(stn_mix)} retail petrol stations ({len(whitespace)/len(stn_mix)*100:.0f}%) "
    f"sell PMG but carry no R95 at all. Ranked by PMG volume, the stations below represent the highest-"
    f"volume, highest-confidence candidates — proven fuel demand at the location, simply missing the "
    f"premium product line.")

top_ws = whitespace.head(30)
w_hdr = ["#", "Station", "City", "Region", "PMG Vol CY (ML)"]
w_tbl = doc.add_table(rows=len(top_ws) + 1, cols=len(w_hdr))
w_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(w_hdr):
    cell_text(w_tbl.cell(0, ci), h, bold=True, size=8, color=WHITE)
    set_cell_bg(w_tbl.cell(0, ci), "C0392B")
for ri, (_, row) in enumerate(top_ws.iterrows(), 1):
    vals = [str(ri), str(row["Name 1"])[:35], row["CityNorm"], row["Sales office Region"],
            f"{row['PMG_Vol_CY_ML']:.2f}"]
    bg = "FDEDEC" if ri % 2 == 0 else "FFFFFF"
    for ci, v in enumerate(vals):
        cell_text(w_tbl.cell(ri, ci), v, size=8,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci == 1 else WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(w_tbl.cell(ri, ci), bg)

# ── 6. Momentum — Growing / Declining R95 Markets ─────────────────────────────
add_heading(doc, "6. Momentum — Cities Where R95 Demand Is Building or Eroding", level=1)
g_hdr = ["City", "Region", "R95 Vol CY (ML)", "R95 Vol SPLY (ML)", "Chg%"]

add_body(doc, "Top 10 growing R95 cities (absolute volume gain vs SPLY):", size=9.5, space_after=2)
g_tbl = doc.add_table(rows=min(10, len(growing)) + 1, cols=len(g_hdr))
g_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(g_hdr):
    cell_text(g_tbl.cell(0, ci), h, bold=True, size=8, color=WHITE)
    set_cell_bg(g_tbl.cell(0, ci), "1E8449")
for ri, (_, row) in enumerate(growing.head(10).iterrows(), 1):
    vals = [row["CityNorm"], row["Sales office Region"], f"{row['Vol_CY_ML']:.3f}",
            f"{row['Vol_SPLY_ML']:.3f}", chg_str(row["Vol_Chg_Pct"])]
    bg = "EAFAF1" if ri % 2 == 0 else "FFFFFF"
    for ci, v in enumerate(vals):
        cell_text(g_tbl.cell(ri, ci), v, size=8,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci <= 1 else WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(g_tbl.cell(ri, ci), bg)

doc.add_paragraph()
add_body(doc, "Cities where R95 volume declined vs SPLY:", size=9.5, space_after=2)
d_tbl = doc.add_table(rows=min(10, max(len(declining),1)) + 1, cols=len(g_hdr))
d_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(g_hdr):
    cell_text(d_tbl.cell(0, ci), h, bold=True, size=8, color=WHITE)
    set_cell_bg(d_tbl.cell(0, ci), "C0392B")
if declining.empty:
    cell_text(d_tbl.cell(1, 0), "No cities with declining R95 volume vs SPLY.", size=8)
    for ci in range(1, len(g_hdr)):
        cell_text(d_tbl.cell(1, ci), "", size=8)
else:
    for ri, (_, row) in enumerate(declining.head(10).iterrows(), 1):
        vals = [row["CityNorm"], row["Sales office Region"], f"{row['Vol_CY_ML']:.3f}",
                f"{row['Vol_SPLY_ML']:.3f}", chg_str(row["Vol_Chg_Pct"])]
        bg = "FDEDEC" if ri % 2 == 0 else "FFFFFF"
        for ci, v in enumerate(vals):
            cell_text(d_tbl.cell(ri, ci), v, size=8,
                      align=WD_ALIGN_PARAGRAPH.LEFT if ci <= 1 else WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_bg(d_tbl.cell(ri, ci), bg)

# ── 7. Uplift Scenarios ────────────────────────────────────────────────────────
add_heading(doc, "7. Volume & Margin Uplift Scenarios", level=1)
add_body(doc,
    f"Stations that already sell R95 carry an average R95 penetration of {avg_r95_penetration*100:.1f}% of "
    f"their petrol volume. Applying that same penetration rate to the {len(whitespace)} whitespace PMG "
    f"stations' existing volume gives a full-potential R95 opportunity of {full_potential_ml:.1f}ML. "
    f"Conservative ({int(CONSERV*100)}% of full potential, realistic 12-month target) and optimal "
    f"({int(OPTIMAL*100)}%, stretch 18-month target) scenarios are shown below.")

u_hdr = ["Scenario", "R95 Vol Added (ML)", "Est. Net Margin Added (PKR M)"]
u_tbl = doc.add_table(rows=4, cols=len(u_hdr))
u_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(u_hdr):
    cell_text(u_tbl.cell(0, ci), h, bold=True, size=9, color=WHITE)
    set_cell_bg(u_tbl.cell(0, ci), "00479D")
scenario_rows = [
    ("Full Potential", full_potential_ml, full_potential_ml * 1e6 * nmgn_per_ltr_r95 / 1e6),
    ("Optimal (55%)", optimal_ml, optimal_nmgn_m),
    ("Conservative (25%)", conservative_ml, conservative_nmgn_m),
]
for ri, (label, vol, nmgn) in enumerate(scenario_rows, 1):
    vals = [label, f"{vol:.1f}", f"{nmgn:.1f}"]
    bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
    for ci, v in enumerate(vals):
        cell_text(u_tbl.cell(ri, ci), v, bold=(ri == 1), size=9,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(u_tbl.cell(ri, ci), bg)

# ── 8. Customer Segments ───────────────────────────────────────────────────────
if not corp_seg.empty:
    add_heading(doc, "8. R95 Customer Segments — Corporate Groups", level=1)
    add_body(doc, "Top corporate groups by R95 volume — a fleet/commercial-account lens on premium demand.")
    seg_hdr = ["Corporate Group", "R95 Vol CY (ML)", "R95 Stations"]
    seg_tbl = doc.add_table(rows=min(15, len(corp_seg)) + 1, cols=len(seg_hdr))
    seg_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(seg_hdr):
        cell_text(seg_tbl.cell(0, ci), h, bold=True, size=8.5, color=WHITE)
        set_cell_bg(seg_tbl.cell(0, ci), "00479D")
    for ri, (_, row) in enumerate(corp_seg.head(15).iterrows(), 1):
        vals = [str(row["Corporate Group"])[:40], f"{row['R95_Vol_CY_ML']:.3f}", str(int(row["R95_Stations"]))]
        bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, v in enumerate(vals):
            cell_text(seg_tbl.cell(ri, ci), v, size=8.5,
                      align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_bg(seg_tbl.cell(ri, ci), bg)

# ── 9. Recommendations ─────────────────────────────────────────────────────────
add_heading(doc, "9. Recommendations", level=1)
top_region = by_region.iloc[0]
add_body(doc,
    f"1. Prioritise the top {min(30, len(whitespace))} whitespace stations in Section 5 for R95 rollout — "
    f"they already have proven PMG throughput, so distribution is the only missing piece.")
add_body(doc,
    f"2. Defend and deepen R95 in {top_region['Sales office Region']} region, its strongest existing market "
    f"at {top_region['R95_Share_Pct']:.1f}% penetration — protect this base while expanding elsewhere.")
add_body(doc,
    f"3. Target the conservative scenario ({conservative_ml:.1f}ML, PKR {conservative_nmgn_m:.1f}M net "
    f"margin) as the 12-month plan; revisit the optimal scenario once early rollout cities prove the "
    f"penetration assumption holds.")
add_body(doc,
    "4. Use growing R95 cities (Section 6) as marketing/pricing test cases before wider rollout — "
    "they show organic demand momentum PSO can amplify rather than create from scratch.")
if not corp_seg.empty:
    add_body(doc,
        f"5. Engage top corporate/fleet accounts in Section 8 directly — commercial buyers are a faster "
        f"path to volume than retail walk-in conversion.")

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run(
    "Data: PSO Working File, Retail Business channel, domestic stations only. "
    "SPLY = Same Period Last Year (10-month like-for-like comparison, not full-year LY). "
    f"{REPORT_DATE}")
r.font.size = Pt(8); r.font.italic = True; r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

out = out_path('PSO_Premium_Fuel_Report', 'docx', df)
doc.save(out)
print(f"Saved -> {out}")
