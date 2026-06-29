"""
PSO Lubricants — Top 20 Cities Volume Profile (Word Document)
Two tables: (1) City overview, (2) Category breakdown
"""
import sys, os
sys.path.insert(0,'src')
sys.stdout.reconfigure(encoding='utf-8')

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from pso import ingest
from _pso_common import INPUT_PATH, get_period_label, out_path

# ── colours ───────────────────────────────────────────────────────────────────
W_BLUE   = RGBColor(0x0A,0x16,0x28)
W_TEAL   = RGBColor(0x14,0x8F,0x77)
W_GREEN  = RGBColor(0x1E,0x84,0x49)
W_RED    = RGBColor(0xC0,0x00,0x00)
W_ORANGE = RGBColor(0xE6,0x7E,0x22)
W_WHITE  = RGBColor(0xFF,0xFF,0xFF)
W_GREY   = RGBColor(0x7F,0x8C,0x8D)
W_DARK   = RGBColor(0x1B,0x2A,0x4A)
W_PURPLE = RGBColor(0x8E,0x44,0xAD)

CAT_ORDER  = ['DEO','PCMO','MCO','LOW GRADE','OTHERS']
CAT_COLORS = {
    'DEO':       RGBColor(0x2E,0x86,0xC1),
    'PCMO':      RGBColor(0x8E,0x44,0xAD),
    'MCO':       RGBColor(0xE6,0x7E,0x22),
    'LOW GRADE': RGBColor(0x1E,0x84,0x49),
    'OTHERS':    RGBColor(0x95,0xA5,0xA6),
}

def set_bg(cell, hex_str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{hex_str}"/>')
    tcPr.append(shd)

def ct(cell, text, bold=False, size=8, color=None,
       align=WD_ALIGN_PARAGRAPH.CENTER, italic=False):
    cell.text = ''
    p = cell.paragraphs[0]; p.alignment = align
    r = p.add_run(str(text))
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = 'Arial'
    if color: r.font.color.rgb = color

def chg_str(v):
    return f"+{v:.1f}%" if v >= 0 else f"{v:.1f}%"

def fmt_kl(v):
    kl = v / 1000
    if kl >= 100: return f"{kl:.1f}"
    if kl >= 10:  return f"{kl:.2f}"
    return f"{kl:.3f}"

# ── load ──────────────────────────────────────────────────────────────────────
print("Loading data…")
df, _ = ingest.load(INPUT_PATH)
REPORT_DATE = get_period_label(df)
retail    = df[df['IsRetail'] & ~df['IsInternational']].copy()
lubes_all = retail[retail['FuelSegment']=='Lubricants'].copy()

# ── compute top 20 cities ─────────────────────────────────────────────────────
top20_cities = (lubes_all.groupby('CityNorm')['SalesLtr_CY'].sum()
                .sort_values(ascending=False).head(20).index.tolist())

rows = []
for city in top20_cities:
    cl  = lubes_all[lubes_all['CityNorm']==city]
    ca  = retail[retail['CityNorm']==city]

    vol_cy   = cl['SalesLtr_CY'].sum()
    vol_ly   = cl['SalesLtr_LY'].sum()
    vol_sply = cl['SalesLtr_SPLY'].sum()
    vol_chg  = (vol_cy - vol_sply) / vol_sply * 100 if vol_sply else 0

    # region (most common)
    region = (cl.groupby('Sales office Region')['SalesLtr_CY'].sum()
                .idxmax() if len(cl) else '')

    stn_total = cl.groupby('Customer Number')['SalesLtr_CY'].sum()
    n_stns    = ca['Customer Number'].nunique()
    n_active  = (stn_total > 0).sum()
    n_zero    = n_stns - n_active
    avg_stn   = vol_cy / n_stns if n_stns else 0
    med_stn   = stn_total[stn_total > 0].median() if n_active else 0
    top10_vol = stn_total.head(10).sum()
    top10_pct = top10_vol / vol_cy * 100 if vol_cy else 0

    # tiers
    p75     = stn_total[stn_total>0].quantile(0.75) if n_active else 0
    n_high  = (stn_total  > p75).sum()
    n_mid   = ((stn_total > med_stn) & (stn_total <= p75)).sum()
    n_low   = ((stn_total > 0) & (stn_total <= med_stn)).sum()

    # per category
    cat_data = {}
    for cat in CAT_ORDER:
        dc = cl[cl['LubeCategory']==cat]
        cv = dc['SalesLtr_CY'].sum()
        cl_stn = dc.groupby('Customer Number')['SalesLtr_CY'].sum()
        cat_data[cat] = dict(
            vol    = cv,
            vol_pct= cv/vol_cy*100 if vol_cy else 0,
            n_sell = (cl_stn > 0).sum(),
            avg_kl = cl_stn[cl_stn>0].mean()/1000 if (cl_stn>0).sum() else 0,
            vol_chg= (cv - dc['SalesLtr_SPLY'].sum()) / dc['SalesLtr_SPLY'].sum() * 100
                      if dc['SalesLtr_SPLY'].sum() else 0,
        )

    rows.append(dict(
        city=city, region=region,
        n_stns=n_stns, n_active=n_active, n_zero=n_zero,
        vol_cy=vol_cy, vol_ly=vol_ly, vol_chg=vol_chg,
        avg_stn=avg_stn, med_stn=med_stn,
        top10_pct=top10_pct,
        n_high=n_high, n_mid=n_mid, n_low=n_low,
        cat=cat_data,
    ))

# ── build Word document ───────────────────────────────────────────────────────
print("Building document…")
doc = Document()
sec = doc.sections[0]
# Landscape A4
sec.page_width   = Cm(29.7)
sec.page_height  = Cm(21.0)
sec.left_margin  = sec.right_margin = Cm(1.4)
sec.top_margin   = sec.bottom_margin = Cm(1.2)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(2)
r = p.add_run('PSO Retail Lubricants — Top 20 Cities Volume Profile')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = W_BLUE; r.font.name = 'Arial'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(8)
r2 = p2.add_run(f'Volume Edition  |  {REPORT_DATE}  |  Ranked by Total Lubricant Volume CY')
r2.font.size = Pt(9.5); r2.font.color.rgb = W_TEAL; r2.font.name = 'Arial'

# ══════════════════════════════════════════════════════════════════════════════
# TABLE 1: City Overview
# ══════════════════════════════════════════════════════════════════════════════
p3 = doc.add_paragraph()
r3 = p3.add_run('Table 1 — City Overview: Volume, Stations & Performance Tiers')
r3.bold = True; r3.font.size = Pt(11); r3.font.color.rgb = W_BLUE; r3.font.name = 'Arial'
p3.paragraph_format.space_after = Pt(3)

hdr1 = ['#','City','Region','Total\nStns',
        'Vol CY\n(KL)','Vol LY 12M\n(KL)','vs\nSPLY',
        'Avg KL\n/Stn','Median KL\n/Stn',
        'Active\nStns','Zero\nStns',
        'Top 10\nConc %',
        'High\nPerf','Mid\nPerf','Low\nPerf']

tbl1 = doc.add_table(rows=len(rows)+2, cols=len(hdr1))
tbl1.style = 'Table Grid'
tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER

# header row
for ci, h in enumerate(hdr1):
    set_bg(tbl1.cell(0, ci), '0A1628')
    ct(tbl1.cell(0, ci), h, bold=True, size=7.5, color=W_WHITE)

# sub-header
sub1 = ['','','','',
        'Litres/1000','Litres/1000','vs SPLY',
        'KL','KL',
        'count','count',
        '% of Vol',
        'stns','stns','stns']
for ci, s in enumerate(sub1):
    set_bg(tbl1.cell(1, ci), '1B2A4A')
    ct(tbl1.cell(1, ci), s, size=6.5, color=W_TEAL, italic=True)

for ri, d in enumerate(rows, 2):
    bg  = 'F4F8FF' if ri % 2 == 0 else 'FFFFFF'
    chg = d['vol_chg']
    chg_col = W_GREEN if chg >= 0 else W_RED

    vals = [
        str(ri-1),
        d['city'],
        d['region'],
        str(d['n_stns']),
        fmt_kl(d['vol_cy']),
        fmt_kl(d['vol_ly']),
        chg_str(chg),
        fmt_kl(d['avg_stn']),
        fmt_kl(d['med_stn']),
        str(d['n_active']),
        str(d['n_zero']),
        f"{d['top10_pct']:.1f}%",
        str(d['n_high']),
        str(d['n_mid']),
        str(d['n_low']),
    ]
    for ci, v in enumerate(vals):
        set_bg(tbl1.cell(ri, ci), bg)
        fc    = chg_col if ci == 6 else (W_RED if ci == 10 and d['n_zero'] > 0 else None)
        bold  = ci in [1, 4]
        align = WD_ALIGN_PARAGRAPH.LEFT if ci in [1, 2] else WD_ALIGN_PARAGRAPH.CENTER
        ct(tbl1.cell(ri, ci), v, bold=bold, size=8, color=fc, align=align)

# set column widths (landscape A4 usable ~26.9cm)
col_widths = [0.45, 3.0, 1.8, 1.0, 1.5, 1.5, 1.1, 1.5, 1.6, 1.0, 0.9, 1.2, 0.9, 0.9, 0.9]
for ci, w in enumerate(col_widths):
    for row in tbl1.rows:
        row.cells[ci].width = Cm(w)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE 2: Category Breakdown
# ══════════════════════════════════════════════════════════════════════════════
p4 = doc.add_paragraph()
r4 = p4.add_run('Table 2 — Category Volume Breakdown by City (KL & % Mix)')
r4.bold = True; r4.font.size = Pt(11); r4.font.color.rgb = W_BLUE; r4.font.name = 'Arial'
p4.paragraph_format.space_after  = Pt(3)
p4.paragraph_format.space_before = Pt(6)

# columns: City | for each cat: Vol(KL) | Mix% | YoY% | Stns
cat_hdr_main = ['#','City','Vol CY\n(KL)','vs\nSPLY']
for cat in CAT_ORDER:
    cat_hdr_main += [f'{cat}\nKL', f'{cat}\n%', f'{cat}\nYoY', f'{cat}\nStns']

n_cols2 = len(cat_hdr_main)
tbl2 = doc.add_table(rows=len(rows)+2, cols=n_cols2)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

# header
cat_col_colors = {
    'DEO':       '1A3A6E', 'PCMO':      '5B2C6F',
    'MCO':       '784212', 'LOW GRADE': '0B5E32',
    'OTHERS':    '5D6D7E',
}
for ci, h in enumerate(cat_hdr_main):
    # colour by category
    if ci < 4:
        bg_hex = '0A1628'
    else:
        cat_idx = (ci - 4) // 4
        bg_hex  = cat_col_colors[CAT_ORDER[cat_idx]]
    set_bg(tbl2.cell(0, ci), bg_hex)
    ct(tbl2.cell(0, ci), h, bold=True, size=7, color=W_WHITE)

# sub-header
sub2 = ['','','Litres/1000','vs SPLY']
for _ in CAT_ORDER:
    sub2 += ['KL','% of total','vs SPLY','# selling']
for ci, s in enumerate(sub2):
    if ci < 4:
        bg_hex = '1B2A4A'
    else:
        cat_idx = (ci - 4) // 4
        bg_hex  = cat_col_colors[CAT_ORDER[cat_idx]]
    set_bg(tbl2.cell(1, ci), bg_hex)
    ct(tbl2.cell(1, ci), s, size=6, color=W_TEAL, italic=True)

for ri, d in enumerate(rows, 2):
    bg  = 'F4F8FF' if ri % 2 == 0 else 'FFFFFF'
    chg = d['vol_chg']
    chg_col = W_GREEN if chg >= 0 else W_RED

    base_vals = [
        str(ri-1),
        d['city'],
        fmt_kl(d['vol_cy']),
        chg_str(chg),
    ]
    cat_vals = []
    for cat in CAT_ORDER:
        cs = d['cat'][cat]
        c_chg = cs['vol_chg']
        cat_vals += [
            fmt_kl(cs['vol']) if cs['vol'] > 0 else '—',
            f"{cs['vol_pct']:.1f}%" if cs['vol'] > 0 else '—',
            chg_str(c_chg) if cs['vol'] > 0 else '—',
            str(int(cs['n_sell'])) if cs['n_sell'] > 0 else '0',
        ]

    all_vals = base_vals + cat_vals
    for ci, v in enumerate(all_vals):
        set_bg(tbl2.cell(ri, ci), bg)
        # colour for chg columns
        fc = None
        if ci == 3:                      # overall YoY
            fc = W_GREEN if chg >= 0 else W_RED
        elif ci >= 4 and (ci - 4) % 4 == 2:  # category YoY cols
            cat_idx = (ci - 4) // 4
            cat_chg = d['cat'][CAT_ORDER[cat_idx]]['vol_chg']
            fc = W_GREEN if cat_chg >= 0 else W_RED
        bold  = ci in [1, 2]
        align = WD_ALIGN_PARAGRAPH.LEFT if ci == 1 else WD_ALIGN_PARAGRAPH.CENTER
        ct(tbl2.cell(ri, ci), v, bold=bold, size=7.5, color=fc, align=align)

# column widths for table 2
col_w2 = [0.45, 2.80, 1.30, 1.00]
for _ in CAT_ORDER:
    col_w2 += [1.20, 0.90, 0.85, 0.80]
for ci, w in enumerate(col_w2):
    for row in tbl2.rows:
        row.cells[ci].width = Cm(w)

# ── footnote ─────────────────────────────────────────────────────────────────
doc.add_paragraph()
fn = doc.add_paragraph()
fn.paragraph_format.space_before = Pt(4)
for part, bold in [
    ('Notes: ', True),
    ('Vol CY/LY = litres ÷ 1,000 (kilolitres). '
     'vs SPLY = current year vs same period last year (like-for-like 10M comparison). LY 12M shown for reference only. '
     'Avg/Median KL per Station = among all stations in city, not just active ones for avg. '
     'High Performers = above 75th percentile volume; Mid = 50th–75th; Low = below 50th (active only). '
     'Top 10 Conc % = share of city volume held by top 10 stations.  '
     'Data: Retail Business only, excl. International.  |  ', False),
    (REPORT_DATE, True),
]:
    r = fn.add_run(part); r.bold = bold
    r.font.size = Pt(7.5); r.font.color.rgb = W_GREY; r.font.name = 'Arial'

# ── save ─────────────────────────────────────────────────────────────────────
out = out_path('PSO_Lubes_Vol_Top20_Cities_Table', 'docx', df)
doc.save(out)
print(f"\nSaved: {out}")
