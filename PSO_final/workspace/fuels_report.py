"""
PSO Retail Fuels Business Report (Diesel + Petrol, incl. R95 vs PMG premium)
Word document with tables and embedded charts. Mirrors workspace/lubes_report.py's
structure/depth, adapted to Fuels — every figure is computed live from the loaded
data (no hardcoded facts), so it keeps working unmodified against a new data drop.
"""
import sys, io, os
sys.path.insert(0, 'src')
sys.stdout.reconfigure(encoding='utf-8')

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.ticker import FuncFormatter
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml

from pso import ingest, premium_fuel_analyze as pf
from _pso_common import INPUT_PATH, get_period_label, out_path

# ── colours ───────────────────────────────────────────────────────────────────
PSO_BLUE  = RGBColor(0x00, 0x47, 0x9D)
PSO_GREEN = RGBColor(0x00, 0x8C, 0x4A)
PSO_RED   = RGBColor(0xC0, 0x00, 0x00)
PSO_ORANGE= RGBColor(0xE4, 0x6C, 0x0A)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

PLT_BLUE  = '#00479D'
PLT_GREEN = '#008C4A'
PLT_RED   = '#C00000'
PLT_ORANGE= '#E46C0A'
PLT_GREY  = '#BFBFBF'

FUEL_CATS = ['Diesel', 'Petrol', 'Other Fuels', 'LPG']
REGIONS   = ['Central', 'South', 'North']

# ── helpers ───────────────────────────────────────────────────────────────────
def kl(v):   return v / 1_000
def ml(v):   return v / 1_000_000
def bn(v):   return v / 1_000_000_000
def pct(a, b): return (a - b) / abs(b) * 100 if b else 0
def chg_str(v): return f"+{v:.1f}%" if v >= 0 else f"{v:.1f}%"

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{hex_color}"/>'))

def cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    if level == 1:
        r.font.size = Pt(16); r.font.color.rgb = PSO_BLUE
    elif level == 2:
        r.font.size = Pt(13); r.font.color.rgb = PSO_BLUE
    else:
        r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    if level <= 2:
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(
            '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="00479D"/></w:pBdr>'))
    return p

def add_body(doc, text, size=10, space_before=3, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for r in p.runs:
        r.font.size = Pt(size)
    return p

def img_buf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close(fig)
    return buf

# ── load data ─────────────────────────────────────────────────────────────────
print("Loading data…")
df, _ = ingest.load(INPUT_PATH)
REPORT_DATE = get_period_label(df)
retail = df[df['IsRetail'] & ~df['IsInternational']].copy()
fuels  = retail[retail['FuelSegment'].isin(FUEL_CATS)].copy()

premium_tables = pf.run_premium_fuel(df)

# ── aggregate helpers ─────────────────────────────────────────────────────────
def agg_vol_margin(frame, by):
    return (frame.groupby(by, as_index=False)
            .agg(vol_cy=('SalesLtr_CY', 'sum'),
                 vol_sply=('SalesLtr_SPLY', 'sum'),
                 rev_cy=('SalesGRS_CY', 'sum'),
                 rev_sply=('SalesGRS_SPLY', 'sum'),
                 mgn_cy=('NetMargin_CY', 'sum'),
                 mgn_sply=('NetMargin_SPLY', 'sum'),
                 stns=('Customer Number', 'nunique'))
            .assign(vol_chg=lambda d: d.apply(lambda r: pct(r.vol_cy, r.vol_sply), axis=1),
                    rev_chg=lambda d: d.apply(lambda r: pct(r.rev_cy, r.rev_sply), axis=1),
                    mgn_pl_cy=lambda d: d.mgn_cy / d.vol_cy.replace(0, np.nan),
                    mgn_pl_sply=lambda d: d.mgn_sply / d.vol_sply.replace(0, np.nan),
                    vol_sh=lambda d: d.vol_cy / d.vol_cy.sum() * 100))

# National
vc = fuels['SalesLtr_CY'].sum()
vs = fuels['SalesLtr_SPLY'].sum()
rc = fuels['SalesGRS_CY'].sum()
rs = fuels['SalesGRS_SPLY'].sum()
mc = fuels['NetMargin_CY'].sum()
ms = fuels['NetMargin_SPLY'].sum()
n_stns = fuels['Customer Number'].nunique()
n_cities = fuels['CityNorm'].nunique()

cat_df = agg_vol_margin(fuels, 'FuelSegment').sort_values('vol_cy', ascending=False)
reg_df = agg_vol_margin(fuels, 'Sales office Region').sort_values('vol_cy', ascending=False)
city_df = agg_vol_margin(fuels, 'CityNorm').sort_values('vol_cy', ascending=False)
stn_df = (fuels.groupby(['Customer Number', 'Name 1', 'CityNorm', 'Sales office Region'], as_index=False)
          .agg(vol_cy=('SalesLtr_CY', 'sum'), vol_sply=('SalesLtr_SPLY', 'sum'), mgn_cy=('NetMargin_CY', 'sum'))
          .assign(vol_chg=lambda d: d.apply(lambda r: pct(r.vol_cy, r.vol_sply), axis=1),
                  mgn_pl_cy=lambda d: d.mgn_cy / d.vol_cy.replace(0, np.nan))
          .sort_values('vol_cy', ascending=False))

cat_lbl = cat_df.set_index('FuelSegment')
reg_lbl = reg_df.set_index('Sales office Region').reindex(REGIONS)

top1_city = city_df.iloc[0]
top1_stn  = stn_df.iloc[0]
top5_city_sh = city_df.head(5)['vol_cy'].sum() / vc * 100 if vc else 0
under_cities = city_df[(city_df['vol_cy'] > 100_000) & (city_df['vol_chg'] < 0)].sort_values('vol_chg')
bot_stns = stn_df[(kl(stn_df['vol_sply']) > 2) & (stn_df['vol_chg'] < -20)].sort_values('vol_chg').head(20)

premium_trend = premium_tables['premium_product_trend']
r95_row = premium_trend[premium_trend['ProductCategory'] == 'R95'].iloc[0]
pmg_row = premium_trend[premium_trend['ProductCategory'] == 'PMG'].iloc[0]
uplift_row = premium_tables['premium_margin_decomp'].iloc[-1]

# ── charts ────────────────────────────────────────────────────────────────────
print("Generating charts…")

# 1. Volume by category (grouped bar) + mix donut
fig1, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 4.2))
fig1.patch.set_facecolor('white')

x = np.arange(len(cat_df)); bw = 0.35
bars_cy = ax1.bar(x - bw/2, kl(cat_df['vol_cy']), bw, color=PLT_BLUE, label='CY', zorder=3)
ax1.bar(x + bw/2, kl(cat_df['vol_sply']), bw, color=PLT_GREY, label='SPLY', zorder=3)
ax1.set_xticks(x); ax1.set_xticklabels(cat_df['FuelSegment'].tolist(), fontsize=9)
ax1.set_ylabel('Volume (KL)', fontsize=9)
ax1.set_title('Volume by Category — CY vs SPLY', fontsize=10, fontweight='bold', color=PLT_BLUE, pad=8)
ax1.legend(fontsize=8)
ax1.yaxis.set_major_formatter(FuncFormatter(lambda v, _: f'{v:,.0f}'))
ax1.set_axisbelow(True); ax1.yaxis.grid(True, linestyle='--', alpha=0.5)
ax1.spines[['top', 'right']].set_visible(False)
for bar, row in zip(bars_cy, cat_df.itertuples()):
    c = PLT_GREEN if row.vol_chg >= 0 else PLT_RED
    ax1.annotate(chg_str(row.vol_chg), xy=(bar.get_x() + bar.get_width()/2, bar.get_height()),
                 xytext=(0, 4), textcoords='offset points', ha='center', fontsize=7.5, color=c, fontweight='bold')

colours = [PLT_BLUE, PLT_GREEN, PLT_ORANGE, PLT_RED]
wedges, texts, autotexts = ax2.pie(
    cat_df['vol_cy'], labels=None, autopct='%1.1f%%', startangle=90,
    colors=colours[:len(cat_df)], wedgeprops=dict(width=0.55, edgecolor='white'), pctdistance=0.75)
for at in autotexts:
    at.set_fontsize(8); at.set_color('white'); at.set_fontweight('bold')
ax2.legend(cat_df['FuelSegment'].tolist(), loc='lower center', bbox_to_anchor=(0.5, -0.12), ncol=2, fontsize=8)
ax2.set_title('Volume Mix — CY', fontsize=10, fontweight='bold', color=PLT_BLUE, pad=8)
plt.tight_layout()
chart1_buf = img_buf(fig1)

# 2. R95 vs PMG margin/litre
fig2, ax = plt.subplots(figsize=(8, 3.8))
fig2.patch.set_facecolor('white')
prods = ['PMG', 'R95']
nmgn_cy = [pmg_row['NMgn_per_Ltr_CY'], r95_row['NMgn_per_Ltr_CY']]
nmgn_sply = [pmg_row['NMgn_per_Ltr_SPLY'], r95_row['NMgn_per_Ltr_SPLY']]
x = np.arange(len(prods)); bw = 0.35
ax.bar(x - bw/2, nmgn_cy, bw, color=PLT_BLUE, label='CY')
ax.bar(x + bw/2, nmgn_sply, bw, color=PLT_GREY, label='SPLY')
ax.set_xticks(x); ax.set_xticklabels(prods, fontsize=9)
ax.set_ylabel('Net Margin / Litre (PKR)', fontsize=9)
ax.set_title('Net Margin per Litre — PMG vs R95', fontsize=10, fontweight='bold', color=PLT_BLUE, pad=8)
ax.legend(fontsize=8)
ax.yaxis.set_major_formatter(FuncFormatter(lambda v, _: f'PKR {v:,.0f}'))
ax.set_axisbelow(True); ax.yaxis.grid(True, linestyle='--', alpha=0.5)
ax.spines[['top', 'right']].set_visible(False)
plt.tight_layout()
chart2_buf = img_buf(fig2)

# 3. Regional comparison
fig3, axes = plt.subplots(1, 3, figsize=(11, 4))
fig3.patch.set_facecolor('white')
ax3a, ax3b, ax3c = axes
bars = ax3a.bar(REGIONS, kl(reg_lbl['vol_cy']), color=[PLT_BLUE, PLT_GREEN, PLT_ORANGE], zorder=3)
ax3a.bar(REGIONS, kl(reg_lbl['vol_sply']), color='none', edgecolor=PLT_GREY, linewidth=1.5, linestyle='--', zorder=2)
ax3a.set_title('Volume CY (KL)', fontsize=10, fontweight='bold', color=PLT_BLUE)
ax3a.yaxis.set_major_formatter(FuncFormatter(lambda v, _: f'{v:,.0f}'))
ax3a.set_axisbelow(True); ax3a.yaxis.grid(True, linestyle='--', alpha=0.5)
ax3a.spines[['top', 'right']].set_visible(False)
for bar, reg in zip(bars, REGIONS):
    chg = reg_lbl.loc[reg, 'vol_chg']
    c = PLT_GREEN if chg >= 0 else PLT_RED
    ax3a.annotate(chg_str(chg), xy=(bar.get_x() + bar.get_width()/2, bar.get_height()),
                  xytext=(0, 4), textcoords='offset points', ha='center', fontsize=8, color=c, fontweight='bold')

ax3b.bar(REGIONS, reg_lbl['mgn_pl_cy'], color=[PLT_BLUE, PLT_GREEN, PLT_ORANGE], zorder=3)
ax3b.axhline(mc/vc if vc else 0, color='black', linestyle='--', linewidth=1, label=f'National avg PKR {mc/vc:.0f}' if vc else '')
ax3b.set_title('Net Margin / Litre (PKR)', fontsize=10, fontweight='bold', color=PLT_BLUE)
ax3b.set_axisbelow(True); ax3b.yaxis.grid(True, linestyle='--', alpha=0.5)
ax3b.spines[['top', 'right']].set_visible(False)
ax3b.legend(fontsize=7)

ax3c.bar(REGIONS, reg_lbl['stns'], color=[PLT_BLUE, PLT_GREEN, PLT_ORANGE], zorder=3)
ax3c.set_title('Active Stations', fontsize=10, fontweight='bold', color=PLT_BLUE)
ax3c.set_axisbelow(True); ax3c.yaxis.grid(True, linestyle='--', alpha=0.5)
ax3c.spines[['top', 'right']].set_visible(False)
plt.tight_layout()
chart3_buf = img_buf(fig3)

# 4. Top 15 cities
fig4, ax4 = plt.subplots(figsize=(11, 5))
fig4.patch.set_facecolor('white')
top15 = city_df.head(15).copy()
ax4.barh(range(len(top15)), kl(top15['vol_cy']), color=PLT_BLUE, zorder=3)
ax4.barh(range(len(top15)), kl(top15['vol_sply']), color='none', edgecolor=PLT_GREY, linewidth=1.2, linestyle='--', zorder=2)
ax4.set_yticks(range(len(top15))); ax4.set_yticklabels(top15['CityNorm'].tolist(), fontsize=9)
ax4.invert_yaxis()
ax4.set_xlabel('Volume (KL)', fontsize=9)
ax4.set_title('Top 15 Cities — Fuels Volume (CY vs SPLY)', fontsize=10, fontweight='bold', color=PLT_BLUE, pad=8)
ax4.xaxis.set_major_formatter(FuncFormatter(lambda v, _: f'{v:,.0f}'))
ax4.set_axisbelow(True); ax4.xaxis.grid(True, linestyle='--', alpha=0.5)
ax4.spines[['top', 'right']].set_visible(False)
for bar, (_, row) in zip(ax4.patches, top15.iterrows()):
    c = PLT_GREEN if row.vol_chg >= 0 else PLT_RED
    ax4.annotate(chg_str(row.vol_chg), xy=(bar.get_width(), bar.get_y() + bar.get_height()/2),
                 xytext=(5, 0), textcoords='offset points', va='center', fontsize=8, color=c)
plt.tight_layout()
chart4_buf = img_buf(fig4)

# 5. Underperforming cities
fig5, ax5 = plt.subplots(figsize=(9, max(3, len(under_cities) * 0.35 + 1)))
fig5.patch.set_facecolor('white')
ax5.barh(range(len(under_cities)), under_cities['vol_chg'], color=PLT_RED, zorder=3)
ax5.set_yticks(range(len(under_cities))); ax5.set_yticklabels(under_cities['CityNorm'].tolist(), fontsize=8)
ax5.invert_yaxis()
ax5.axvline(0, color='black', linewidth=0.8)
ax5.set_xlabel('Volume Change %', fontsize=9)
ax5.set_title('Declining Cities (>100 KL threshold)', fontsize=10, fontweight='bold', color=PLT_RED, pad=8)
ax5.xaxis.set_major_formatter(FuncFormatter(lambda v, _: f'{v:+.1f}%'))
ax5.set_axisbelow(True); ax5.xaxis.grid(True, linestyle='--', alpha=0.4)
ax5.spines[['top', 'right']].set_visible(False)
plt.tight_layout()
chart5_buf = img_buf(fig5)

# 6. Station scatter (top 200)
top200 = stn_df[stn_df['vol_cy'] > 0].head(200).copy()
fig6, ax6 = plt.subplots(figsize=(10, 5))
fig6.patch.set_facecolor('white')
reg_colours = {'Central': PLT_BLUE, 'South': PLT_GREEN, 'North': PLT_ORANGE}
for reg, grp in top200.groupby('Sales office Region'):
    ax6.scatter(kl(grp['vol_cy']), grp['mgn_pl_cy'], color=reg_colours.get(reg, 'grey'),
                alpha=0.7, s=40, label=reg, zorder=3)
ax6.set_xlabel('Volume CY (KL)', fontsize=9)
ax6.set_ylabel('Net Margin / Litre (PKR)', fontsize=9)
ax6.set_title('Station Landscape — Volume vs Margin (Top 200 by Volume)', fontsize=10, fontweight='bold', color=PLT_BLUE, pad=8)
ax6.legend(fontsize=8)
ax6.set_axisbelow(True); ax6.grid(True, linestyle='--', alpha=0.4)
ax6.spines[['top', 'right']].set_visible(False)
if vc:
    ax6.axhline(mc / vc, color=PLT_GREY, linestyle='--', linewidth=1)
active_vol = stn_df[stn_df['vol_cy'] > 0]['vol_cy']
if len(active_vol):
    ax6.axvline(kl(active_vol.median()), color=PLT_GREY, linestyle='--', linewidth=1)
plt.tight_layout()
chart6_buf = img_buf(fig6)

# ── build Word document ───────────────────────────────────────────────────────
print("Building Word document…")
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21); sec.page_height = Cm(29.7)
sec.left_margin = sec.right_margin = Cm(1.8)
sec.top_margin = sec.bottom_margin = Cm(1.6)

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(40)
r = cover.add_run('PSO Retail Fuels Business')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = PSO_BLUE

cover2 = doc.add_paragraph()
cover2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = cover2.add_run('Performance Report — Diesel, Petrol & Premium Fuel (R95 vs PMG)')
r2.font.size = Pt(13); r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

cover3 = doc.add_paragraph()
cover3.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover3.paragraph_format.space_before = Pt(4)
r3 = cover3.add_run(REPORT_DATE)
r3.font.size = Pt(11); r3.font.color.rgb = PSO_BLUE; r3.bold = True

doc.add_paragraph()

kpi_tbl = doc.add_table(rows=2, cols=4)
kpi_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
kpi_headers = ['Total Volume CY', 'Revenue CY', 'Net Margin CY', 'Margin / Litre']
kpi_values = [f'{ml(vc):.3f} ML', f'PKR {bn(rc):.2f} Bn', f'PKR {bn(mc):.2f} Bn', f'PKR {mc/vc:.0f}' if vc else 'N/A']
kpi_changes = [chg_str(pct(vc, vs)), chg_str(pct(rc, rs)), chg_str(pct(mc, ms)),
               chg_str(pct(mc/vc if vc else 0, ms/vs if vs else 0))]
kpi_chg_col = [PSO_GREEN if pct(vc, vs) >= 0 else PSO_RED, PSO_GREEN if pct(rc, rs) >= 0 else PSO_RED,
               PSO_GREEN if pct(mc, ms) >= 0 else PSO_RED,
               PSO_GREEN if pct(mc/vc if vc else 0, ms/vs if vs else 0) >= 0 else PSO_RED]
for ci in range(4):
    h_cell = kpi_tbl.cell(0, ci); v_cell = kpi_tbl.cell(1, ci)
    set_cell_bg(h_cell, '00479D')
    cell_text(h_cell, kpi_headers[ci], bold=True, size=9, color=WHITE)
    set_cell_bg(v_cell, 'F2F2F2')
    v_cell.text = ''
    vp = v_cell.paragraphs[0]; vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr1 = vp.add_run(kpi_values[ci] + ' '); vr1.bold = True; vr1.font.size = Pt(11); vr1.font.color.rgb = PSO_BLUE
    vr2 = vp.add_run(kpi_changes[ci]); vr2.bold = True; vr2.font.size = Pt(9); vr2.font.color.rgb = kpi_chg_col[ci]

doc.add_paragraph()
add_body(doc,
    f'This report covers Diesel, Petrol, and premium/regular petrol (R95 vs PMG) sales within PSO\'s '
    f'Retail Business segment for {REPORT_DATE}. The network comprises {n_stns:,} active stations across '
    f'{n_cities} cities nationally. Unless otherwise noted, volumes are in kilolitres (KL) and all '
    f'financials are in PKR.', size=9, space_after=2)

doc.add_page_break()

# SECTION 1: TOTAL VOLUME OVERVIEW
add_heading(doc, '1.  Total Fuels Volume — National Overview', level=1)
add_body(doc,
    f'PSO\'s retail fuels network sold {ml(vc):.3f} million litres (ML) in CY, compared to {ml(vs):.3f} ML '
    f'in the same period last year (SPLY) — a change of {pct(vc, vs):.1f}%. Revenue moved {pct(rc, rs):.1f}% '
    f'to PKR {bn(rc):.2f} Bn. Net margin per litre moved from PKR {ms/vs:.0f} (SPLY) to PKR {mc/vc:.0f} (CY), '
    f'a change of {pct(mc/vc if vc else 0, ms/vs if vs else 0):.1f}%. {cat_lbl.index[0] if len(cat_lbl) else "Diesel"} '
    f'is the largest category at {cat_df.iloc[0]["vol_sh"]:.1f}% of fuels volume.', size=10)
doc.add_picture(chart1_buf, width=Inches(6.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading(doc, '1.1  Volume & Margin by Category', level=2)
hdr_row = ['Category', 'Vol CY (KL)', 'Vol SPLY (KL)', 'vs SPLY', 'Mix %', 'Mgn/L CY (PKR)', 'Mgn/L SPLY (PKR)']
tbl1 = doc.add_table(rows=len(cat_df) + 2, cols=len(hdr_row))
tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl1.style = 'Table Grid'
for ci, h in enumerate(hdr_row):
    set_cell_bg(tbl1.cell(0, ci), '00479D')
    cell_text(tbl1.cell(0, ci), h, bold=True, size=9, color=WHITE)
for ri, (_, row) in enumerate(cat_df.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    vals = [row['FuelSegment'], f'{kl(row.vol_cy):,.1f}', f'{kl(row.vol_sply):,.1f}', chg_str(row.vol_chg),
            f'{row.vol_sh:.1f}%', f'{row.mgn_pl_cy:.0f}' if not np.isnan(row.mgn_pl_cy) else 'N/A',
            f'{row.mgn_pl_sply:.0f}' if not np.isnan(row.mgn_pl_sply) else 'N/A']
    for ci, v in enumerate(vals):
        set_cell_bg(tbl1.cell(ri, ci), bg)
        chg_clr = (PSO_GREEN if row.vol_chg >= 0 else PSO_RED) if ci == 3 else None
        cell_text(tbl1.cell(ri, ci), v, size=9, color=chg_clr,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)
tot_row = len(cat_df) + 1
set_cell_bg(tbl1.cell(tot_row, 0), '1F3864')
cell_text(tbl1.cell(tot_row, 0), 'TOTAL', bold=True, size=9, color=WHITE)
for ci, v in enumerate([f'{kl(vc):,.1f}', f'{kl(vs):,.1f}', chg_str(pct(vc, vs)), '100.0%',
                         f'{mc/vc:.0f}' if vc else 'N/A', f'{ms/vs:.0f}' if vs else 'N/A'], 1):
    set_cell_bg(tbl1.cell(tot_row, ci), '1F3864')
    cell_text(tbl1.cell(tot_row, ci), v, bold=True, size=9, color=WHITE)
doc.add_paragraph()

# SECTION 1.2: PREMIUM FUEL (R95 vs PMG)
add_heading(doc, '1.2  Premium Fuel — R95 vs PMG', level=2)
add_body(doc,
    f'R95 (premium petrol) represents {r95_row["Vol_Share_Pct"]:.1f}% of retail petrol volume, growing '
    f'{chg_str(r95_row["Vol_Chg_Pct"])} vs SPLY against PMG\'s {chg_str(pmg_row["Vol_Chg_Pct"])}, and carries '
    f'PKR {uplift_row["NMgn_per_Ltr_CY"]:.2f}/litre more net margin than PMG. '
    f'{len(premium_tables["premium_whitespace_stations"])} active PMG stations currently sell zero R95 — '
    f'see workspace/fuels_stations_report.py and the dedicated Premium Fuel report for the full launch-site '
    f'analysis.', size=10)
doc.add_picture(chart2_buf, width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# SECTION 2: REGIONAL ANALYSIS
add_heading(doc, '2.  Regional Analysis', level=1)
add_body(doc,
    f'PSO\'s fuels sales are organised across three sales office regions: Central, South, and North. '
    f'{reg_lbl["vol_cy"].idxmax()} leads with {reg_lbl.loc[reg_lbl["vol_cy"].idxmax(), "vol_sh"]:.1f}% of '
    f'national fuels volume.', size=10)
doc.add_picture(chart3_buf, width=Inches(6.4))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading(doc, '2.1  Region-wise Summary', level=2)
r_hdr = ['Region', 'Vol CY (KL)', 'Vol SPLY (KL)', 'vs SPLY', 'Mix %', 'Stations', 'Mgn/L CY', 'Mgn/L SPLY']
tbl2 = doc.add_table(rows=len(reg_df) + 2, cols=len(r_hdr))
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = 'Table Grid'
for ci, h in enumerate(r_hdr):
    set_cell_bg(tbl2.cell(0, ci), '00479D')
    cell_text(tbl2.cell(0, ci), h, bold=True, size=9, color=WHITE)
for ri, (_, row) in enumerate(reg_df.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    vals = [row['Sales office Region'], f'{kl(row.vol_cy):,.1f}', f'{kl(row.vol_sply):,.1f}', chg_str(row.vol_chg),
            f'{row.vol_sh:.1f}%', str(int(row.stns)), f'PKR {row.mgn_pl_cy:.0f}', f'PKR {row.mgn_pl_sply:.0f}']
    for ci, v in enumerate(vals):
        set_cell_bg(tbl2.cell(ri, ci), bg)
        chg_clr = (PSO_GREEN if row.vol_chg >= 0 else PSO_RED) if ci == 3 else None
        cell_text(tbl2.cell(ri, ci), v, size=9, color=chg_clr,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)
tot_r = len(reg_df) + 1
set_cell_bg(tbl2.cell(tot_r, 0), '1F3864')
cell_text(tbl2.cell(tot_r, 0), 'NATIONAL', bold=True, size=9, color=WHITE)
for ci, v in enumerate([f'{kl(vc):,.1f}', f'{kl(vs):,.1f}', chg_str(pct(vc, vs)), '100.0%', str(n_stns),
                         f'PKR {mc/vc:.0f}' if vc else 'N/A', f'PKR {ms/vs:.0f}' if vs else 'N/A'], 1):
    set_cell_bg(tbl2.cell(tot_r, ci), '1F3864')
    cell_text(tbl2.cell(tot_r, ci), v, bold=True, size=9, color=WHITE)
doc.add_page_break()

# SECTION 3: CITY-LEVEL ANALYSIS
add_heading(doc, '3.  City-Level Performance', level=1)
add_body(doc,
    f'Fuels are sold across {n_cities} cities. {top1_city["CityNorm"]} leads nationally with '
    f'{kl(top1_city.vol_cy):,.1f} KL. The top 5 cities contribute {top5_city_sh:.0f}% of national fuels volume.',
    size=10)
doc.add_picture(chart4_buf, width=Inches(6.4))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading(doc, '3.1  Top 15 Cities by Volume', level=2)
c_hdr = ['#', 'City', 'Region', 'Vol CY (KL)', 'Vol SPLY (KL)', 'vs SPLY', 'Stns', 'Mgn/L (PKR)']
top15_full = city_df.head(15).copy()
city_reg = (fuels.groupby(['CityNorm', 'Sales office Region'])['SalesLtr_CY'].sum().reset_index()
            .sort_values('SalesLtr_CY', ascending=False).drop_duplicates('CityNorm')
            .set_index('CityNorm')['Sales office Region'])
tbl4 = doc.add_table(rows=len(top15_full) + 1, cols=len(c_hdr))
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl4.style = 'Table Grid'
for ci, h in enumerate(c_hdr):
    set_cell_bg(tbl4.cell(0, ci), '00479D')
    cell_text(tbl4.cell(0, ci), h, bold=True, size=8.5, color=WHITE)
for ri, (_, row) in enumerate(top15_full.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    reg_name = city_reg.get(row['CityNorm'], '')
    chg_c = PSO_GREEN if row.vol_chg >= 0 else PSO_RED
    vals_data = [(str(ri), None), (row['CityNorm'], None), (reg_name, None),
                 (f'{kl(row.vol_cy):,.1f}', None), (f'{kl(row.vol_sply):,.1f}', None),
                 (chg_str(row.vol_chg), chg_c), (str(int(row.stns)), None),
                 (f'{row.mgn_pl_cy:.0f}' if not np.isnan(row.mgn_pl_cy) else 'N/A', None)]
    for ci, (v, fc) in enumerate(vals_data):
        set_cell_bg(tbl4.cell(ri, ci), bg)
        cell_text(tbl4.cell(ri, ci), v, size=8.5, color=fc,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci in (1, 2) else WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_heading(doc, '3.2  Underperforming Cities (Declining Volume, >100 KL Threshold)', level=2)
add_body(doc,
    f'{len(under_cities)} cities with meaningful volume (>100 KL) recorded year-on-year declines. '
    f'{"" if under_cities.empty else under_cities.iloc[0]["CityNorm"] + " is the largest by absolute volume, "}'
    f'showing the steepest decline among cities above the threshold.', size=10)
if len(under_cities):
    doc.add_picture(chart5_buf, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    uc_hdr = ['City', 'Region', 'Vol CY (KL)', 'Vol SPLY (KL)', 'vs SPLY', 'Mgn/L (PKR)']
    tbl5 = doc.add_table(rows=min(len(under_cities), 20) + 1, cols=len(uc_hdr))
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl5.style = 'Table Grid'
    for ci, h in enumerate(uc_hdr):
        set_cell_bg(tbl5.cell(0, ci), 'C00000')
        cell_text(tbl5.cell(0, ci), h, bold=True, size=8.5, color=WHITE)
    for ri, (_, row) in enumerate(under_cities.head(20).iterrows(), 1):
        bg = 'FFEDED' if ri % 2 == 1 else 'FFF4F4'
        reg_name = city_reg.get(row['CityNorm'], '')
        vals_data = [row['CityNorm'], reg_name, f'{kl(row.vol_cy):,.1f}', f'{kl(row.vol_sply):,.1f}',
                     chg_str(row.vol_chg), f'{row.mgn_pl_cy:.0f}' if not np.isnan(row.mgn_pl_cy) else 'N/A']
        for ci, v in enumerate(vals_data):
            set_cell_bg(tbl5.cell(ri, ci), bg)
            chg_clr = PSO_RED if ci == 4 else None
            cell_text(tbl5.cell(ri, ci), v, size=8.5, color=chg_clr,
                      align=WD_ALIGN_PARAGRAPH.LEFT if ci in (0, 1) else WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# SECTION 4: STATION-LEVEL SEGMENTATION
add_heading(doc, '4.  Station-Level Segmentation', level=1)
active_vol_median = kl(stn_df[stn_df.vol_cy > 0]['vol_cy'].median()) if (stn_df['vol_cy'] > 0).any() else 0
add_body(doc,
    f'The fuels network has {n_stns:,} active stations. The top station ({top1_stn["Name 1"]}, '
    f'{top1_stn["CityNorm"]}) sold {kl(top1_stn.vol_cy):,.1f} KL while the median station sells '
    f'{active_vol_median:.1f} KL.', size=10)
doc.add_picture(chart6_buf, width=Inches(6.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading(doc, '4.1  Top 25 Stations by Volume', level=2)
top25 = stn_df.head(25).copy()
s_hdr = ['#', 'Station Name', 'City', 'Region', 'Vol CY (KL)', 'Vol SPLY (KL)', 'vs SPLY', 'Mgn/L (PKR)']
tbl6 = doc.add_table(rows=len(top25) + 1, cols=len(s_hdr))
tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl6.style = 'Table Grid'
for ci, h in enumerate(s_hdr):
    set_cell_bg(tbl6.cell(0, ci), '00479D')
    cell_text(tbl6.cell(0, ci), h, bold=True, size=8, color=WHITE)
for ri, (_, row) in enumerate(top25.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    chg_clr = PSO_GREEN if row.vol_chg >= 0 else PSO_RED
    vals_d = [(str(ri), None), (str(row['Name 1'])[:32], None), (row['CityNorm'], None),
              (row['Sales office Region'], None), (f'{kl(row.vol_cy):.1f}', None),
              (f'{kl(row.vol_sply):.1f}', None), (chg_str(row.vol_chg), chg_clr),
              (f'{row.mgn_pl_cy:.0f}' if not np.isnan(row.mgn_pl_cy) else 'N/A', None)]
    for ci, (v, fc) in enumerate(vals_d):
        set_cell_bg(tbl6.cell(ri, ci), bg)
        cell_text(tbl6.cell(ri, ci), v, size=8, color=fc,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci in (1, 2, 3) else WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_heading(doc, '4.2  Stations with Significant Decline (SPLY > 2 KL)', level=2)
add_body(doc,
    f'{len(bot_stns)} stations that previously sold meaningful volumes (>2 KL SPLY) have declined by more '
    f'than 20%. These represent the highest-risk accounts and should be prioritised for sales intervention.',
    size=10)
if len(bot_stns):
    b_hdr = ['Station Name', 'City', 'Region', 'Vol CY (KL)', 'Vol SPLY (KL)', 'vs SPLY', 'Mgn/L (PKR)']
    tbl8 = doc.add_table(rows=len(bot_stns) + 1, cols=len(b_hdr))
    tbl8.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl8.style = 'Table Grid'
    for ci, h in enumerate(b_hdr):
        set_cell_bg(tbl8.cell(0, ci), 'C00000')
        cell_text(tbl8.cell(0, ci), h, bold=True, size=8.5, color=WHITE)
    for ri, (_, row) in enumerate(bot_stns.iterrows(), 1):
        bg = 'FFEDED' if ri % 2 == 1 else 'FFF4F4'
        vals_d = [(str(row['Name 1'])[:32], None), (row['CityNorm'], None), (row['Sales office Region'], None),
                  (f'{kl(row.vol_cy):.2f}', None), (f'{kl(row.vol_sply):.2f}', None),
                  (chg_str(row.vol_chg), PSO_RED),
                  (f'{row.mgn_pl_cy:.0f}' if not np.isnan(row.mgn_pl_cy) else 'N/A', None)]
        for ci, (v, fc) in enumerate(vals_d):
            set_cell_bg(tbl8.cell(ri, ci), bg)
            cell_text(tbl8.cell(ri, ci), v, size=8.5, color=fc,
                      align=WD_ALIGN_PARAGRAPH.LEFT if ci in (0, 1, 2) else WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# SECTION 5: STRATEGIC IMPLICATIONS (data-driven, generated from the tables above)
add_heading(doc, '5.  Strategic Implications', level=1)
add_body(doc, 'Based on the data-driven analysis above, the following priorities are recommended.', size=10)

implications = []
implications.append((
    '5.1  Premium Fuel (R95) Expansion',
    f'R95 carries PKR {uplift_row["NMgn_per_Ltr_CY"]:.2f}/litre more net margin than PMG and is growing '
    f'{chg_str(r95_row["Vol_Chg_Pct"])} vs SPLY, faster than PMG. '
    f'{len(premium_tables["premium_whitespace_stations"])} active PMG stations sell zero R95 today — '
    f'the highest-volume of these are proven-demand, low-risk expansion targets. '
    f'See the dedicated Premium Fuel report for the full station-level whitespace list.'
))
worst_region = reg_lbl['mgn_pl_cy'].idxmin() if reg_lbl['mgn_pl_cy'].notna().any() else REGIONS[0]
best_region = reg_lbl['mgn_pl_cy'].idxmax() if reg_lbl['mgn_pl_cy'].notna().any() else REGIONS[0]
implications.append((
    '5.2  Regional Margin Gap',
    f'{best_region} region has the highest margin per litre (PKR {reg_lbl.loc[best_region, "mgn_pl_cy"]:.0f}), '
    f'while {worst_region} trails at PKR {reg_lbl.loc[worst_region, "mgn_pl_cy"]:.0f}/litre. '
    f'Closing even part of this gap in {worst_region} — PSO\'s '
    f'{"largest" if reg_lbl.loc[worst_region, "vol_sh"] == reg_lbl["vol_sh"].max() else ""} volume region — '
    f'would be a direct margin improvement without needing volume growth.'
))
if len(under_cities):
    implications.append((
        '5.3  Declining-City Recovery',
        f'{len(under_cities)} cities with >100 KL fuels volume are declining vs SPLY, led by '
        f'{under_cities.iloc[0]["CityNorm"]} ({chg_str(under_cities.iloc[0]["vol_chg"])}). '
        f'A structured recovery programme — root-cause analysis per city, competitor pricing checks, '
        f'distribution audits — should be prioritised for the largest-volume decliners first.'
    ))
if len(bot_stns):
    implications.append((
        '5.4  Station-Level Intervention',
        f'{len(bot_stns)} stations with meaningful SPLY volume have declined more than 20%. These accounts '
        f'should be flagged for immediate sales officer review — early intervention (product offers, credit '
        f'terms, visibility support) is the best lever to prevent permanent share loss to competitors.'
    ))

for title, body in implications:
    add_heading(doc, title, level=2)
    add_body(doc, body, size=10)

doc.add_page_break()

# APPENDIX
add_heading(doc, 'Appendix A — All Significant Cities (Top 50)', level=1)
add_body(doc, 'Cities ranked by CY volume. Includes all cities with >50 KL CY volume.', size=9)
top50 = city_df[kl(city_df['vol_cy']) > 50].copy()
app_hdr = ['#', 'City', 'Vol CY (KL)', 'Vol SPLY (KL)', 'vs SPLY', 'Stns', 'Mgn/L (PKR)']
tbl_app = doc.add_table(rows=len(top50) + 1, cols=len(app_hdr))
tbl_app.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_app.style = 'Table Grid'
for ci, h in enumerate(app_hdr):
    set_cell_bg(tbl_app.cell(0, ci), '00479D')
    cell_text(tbl_app.cell(0, ci), h, bold=True, size=8, color=WHITE)
for ri, (_, row) in enumerate(top50.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    chg_clr = PSO_GREEN if row.vol_chg >= 0 else PSO_RED
    vals = [(str(ri), None), (row['CityNorm'], None), (f'{kl(row.vol_cy):,.1f}', None),
            (f'{kl(row.vol_sply):,.1f}', None), (chg_str(row.vol_chg), chg_clr),
            (str(int(row.stns)), None), (f'{row.mgn_pl_cy:.0f}' if not np.isnan(row.mgn_pl_cy) else 'N/A', None)]
    for ci, (v, fc) in enumerate(vals):
        set_cell_bg(tbl_app.cell(ri, ci), bg)
        cell_text(tbl_app.cell(ri, ci), v, size=8, color=fc,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci == 1 else WD_ALIGN_PARAGRAPH.CENTER)

# ── save ──────────────────────────────────────────────────────────────────────
out = out_path('PSO_Fuels_Report', 'docx', df)
doc.save(out)
print(f"\nReport saved: {out}")
