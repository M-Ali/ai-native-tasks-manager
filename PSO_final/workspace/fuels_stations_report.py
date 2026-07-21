"""
PSO Retail Fuels — Station & Volume Analysis Word Report
Mirrors workspace/lubes_stations_report.py, adapted to Fuels — every finding is
computed live from the data (no hardcoded city facts).
"""
import sys, io, os
sys.path.insert(0, 'src')
sys.stdout.reconfigure(encoding='utf-8')

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from pso import ingest
from _pso_common import INPUT_PATH, get_period_label, out_path

BLUE, GREEN, RED, ORANGE, GREY = '#00479D', '#008C4A', '#C00000', '#E46C0A', '#BFBFBF'
W_BLUE  = RGBColor(0x00, 0x47, 0x9D)
W_GREEN = RGBColor(0x00, 0x8C, 0x4A)
W_RED   = RGBColor(0xC0, 0x00, 0x00)
W_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

CATS = ['Diesel', 'Petrol', 'Other Fuels', 'LPG']

def set_bg(cell, hex_col):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{hex_col}"/>'))

def ct(cell, text, bold=False, size=9.5, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]; p.alignment = align
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    if level == 1:
        r.font.size = Pt(15); r.font.color.rgb = W_BLUE
    else:
        r.font.size = Pt(12); r.font.color.rgb = W_BLUE
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="00479D"/></w:pBdr>'))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def body(doc, text, size=10, before=3, after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    for r in p.runs: r.font.size = Pt(size)

def img_buf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=160, bbox_inches='tight')
    buf.seek(0); plt.close(fig); return buf

def chg(v): return f"+{v:.1f}%" if v >= 0 else f"{v:.1f}%"

# ── load data ─────────────────────────────────────────────────────────────────
print("Loading data…")
df, _ = ingest.load(INPUT_PATH)
REPORT_DATE = get_period_label(df)
retail = df[df['IsRetail'] & ~df['IsInternational']].copy()
fuels  = retail[retail['FuelSegment'].isin(CATS)].copy()

cat = (fuels.groupby('FuelSegment')
       .agg(vol_cy=('SalesLtr_CY', 'sum'), vol_sply=('SalesLtr_SPLY', 'sum'),
            rev_cy=('SalesGRS_CY', 'sum'), mgn_cy=('NetMargin_CY', 'sum'),
            stns=('Customer Number', 'nunique'))
       .assign(vol_chg=lambda d: (d.vol_cy - d.vol_sply) / d.vol_sply.abs().replace(0, float('nan')) * 100,
               vol_sh=lambda d: d.vol_cy / d.vol_cy.sum() * 100,
               vol_ps=lambda d: d.vol_cy / d.stns,
               mgn_pl=lambda d: d.mgn_cy / d.vol_cy)
       .sort_values('vol_cy', ascending=False))

city_reg = (fuels.groupby(['CityNorm', 'Sales office Region'])['SalesLtr_CY']
            .sum().reset_index().sort_values('SalesLtr_CY', ascending=False)
            .drop_duplicates('CityNorm').set_index('CityNorm')['Sales office Region'])

city = (fuels.groupby('CityNorm')
        .agg(vol_cy=('SalesLtr_CY', 'sum'), vol_sply=('SalesLtr_SPLY', 'sum'),
             mgn_cy=('NetMargin_CY', 'sum'), stns=('Customer Number', 'nunique'))
        .assign(vol_chg=lambda d: (d.vol_cy - d.vol_sply) / d.vol_sply.abs().replace(0, float('nan')) * 100,
                vol_sh=lambda d: d.vol_cy / d.vol_cy.sum() * 100,
                vol_ps=lambda d: d.vol_cy / d.stns,
                mgn_pl=lambda d: d.mgn_cy / d.vol_cy)
        .sort_values('vol_cy', ascending=False)
        .head(15))
city['Region'] = city.index.map(city_reg)

# All-cities table for city-wide growth/decline stats (beyond top 15)
city_all = (fuels.groupby('CityNorm')
            .agg(vol_cy=('SalesLtr_CY', 'sum'), vol_sply=('SalesLtr_SPLY', 'sum'),
                 stns=('Customer Number', 'nunique'))
            .assign(vol_chg=lambda d: (d.vol_cy - d.vol_sply) / d.vol_sply.abs().replace(0, float('nan')) * 100,
                    vol_ps=lambda d: d.vol_cy / d.stns))

# ── CHART 1: Category — volume bars + stations dot ────────────────────────────
print("Generating charts…")
fig1, ax1 = plt.subplots(figsize=(9, 5))
fig1.patch.set_facecolor('white')
ax1r = ax1.twinx()

labels4 = cat.index.tolist()
x4 = np.arange(len(cat)); bw = 0.5
bar_cols = [BLUE, GREEN, ORANGE, RED][:len(cat)]

bars_cy = ax1.bar(x4 - 0.03, cat['vol_cy']/1000, bw, color=bar_cols, alpha=0.88, zorder=3, label='Vol CY')
ax1.bar(x4 - 0.03, cat['vol_sply']/1000, bw, color='none',
        edgecolor='#999999', linewidth=1.6, linestyle='--', zorder=2, label='Vol SPLY')

ax1r.plot(x4, cat['stns'], 'o--', color='#222222', linewidth=1.6,
          markersize=8, markerfacecolor='white', markeredgewidth=2.2, zorder=5, label='Stations')
for xi, (_, r) in zip(x4, cat.iterrows()):
    ax1r.annotate(f"{int(r.stns):,}", xy=(xi, r.stns), xytext=(0, 10), textcoords='offset points',
                  ha='center', fontsize=9, fontweight='bold', color='#222222')

for bar, (_, r) in zip(bars_cy, cat.iterrows()):
    c = GREEN if r.vol_chg >= 0 else RED
    ax1.annotate(chg(r.vol_chg), xy=(bar.get_x()+bar.get_width()/2, bar.get_height()),
                 xytext=(0, 6), textcoords='offset points', ha='center', fontsize=9, color=c, fontweight='bold')
    ax1.annotate(f"{r.vol_cy/1000:,.0f} KL", xy=(bar.get_x()+bar.get_width()/2, bar.get_height()/2),
                 ha='center', fontsize=8.5, color='white', fontweight='bold')

ax1.set_xticks(x4); ax1.set_xticklabels(labels4, fontsize=11)
ax1.set_ylabel('Volume CY (KL)', fontsize=10, color='#333333')
ax1r.set_ylabel('No. of Stations', fontsize=10, color='#222222')
ax1.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'{v:,.0f}'))
ax1r.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'{v:,.0f}'))
ax1.set_title('Fuels Volume & Station Count by Category', fontsize=12, fontweight='bold', color=BLUE, pad=10)
ax1.set_axisbelow(True); ax1.yaxis.grid(True, linestyle='--', alpha=0.4)
ax1.spines[['top', 'right']].set_visible(False)

h1, l1 = ax1.get_legend_handles_labels()
stn_line = plt.Line2D([0], [0], color='#222222', marker='o', linewidth=1.5,
                       markerfacecolor='white', markeredgewidth=2, label='Stations')
ax1.legend(h1 + [stn_line], l1 + ['Stations'], fontsize=9, loc='upper right')
plt.tight_layout()
chart1_buf = img_buf(fig1)

# ── CHART 2: Top 15 cities — stations bar + KL/station line ──────────────────
fig2, ax2 = plt.subplots(figsize=(12, 5.5))
fig2.patch.set_facecolor('white')
ax2r = ax2.twinx()

n_city = len(city)
city_names = [n[:20] for n in city.index.tolist()]
bar_c2 = [GREEN if c >= 0 else RED for c in city['vol_chg']]
bars2 = ax2.bar(range(n_city), city['stns'], 0.6, color=bar_c2, alpha=0.85, zorder=3)
for bar, n in zip(bars2, city['stns']):
    ax2.annotate(str(int(n)), xy=(bar.get_x()+bar.get_width()/2, bar.get_height()),
                 xytext=(0, 4), textcoords='offset points', ha='center', fontsize=8.5,
                 fontweight='bold', color='#222222')

ax2r.plot(range(n_city), city['vol_ps']/1000, 's-', color=BLUE, linewidth=2, markersize=7, zorder=4, label='KL / Station')
for xi, vps in enumerate(city['vol_ps']/1000):
    ax2r.annotate(f'{vps:.1f}', xy=(xi, vps), xytext=(0, 9), textcoords='offset points',
                  ha='center', fontsize=8, color=BLUE, fontweight='bold')

ax2.set_xticks(range(n_city)); ax2.set_xticklabels(city_names, rotation=38, ha='right', fontsize=9)
ax2.set_ylabel('No. of Stations', fontsize=10)
ax2r.set_ylabel('Volume per Station (KL)', fontsize=10, color=BLUE)
ax2r.tick_params(axis='y', labelcolor=BLUE)
ax2.set_title('Top 15 Cities — Active Fuel Stations & Volume per Station', fontsize=12, fontweight='bold', color=BLUE, pad=10)
ax2.set_axisbelow(True); ax2.yaxis.grid(True, linestyle='--', alpha=0.4)
ax2.spines[['top', 'right']].set_visible(False)

g_patch = mpatches.Patch(color=GREEN, label='Growing city')
r_patch = mpatches.Patch(color=RED, label='Declining city')
b_line = plt.Line2D([0], [0], color=BLUE, marker='s', linewidth=2, label='KL / Station')
ax2.legend(handles=[g_patch, r_patch, b_line], fontsize=9, loc='upper right')
plt.tight_layout()
chart2_buf = img_buf(fig2)

# ── CHART 3: Donut — volume & station share per category ─────────────────────
fig3, (ax3a, ax3b) = plt.subplots(1, 2, figsize=(10, 4.5))
fig3.patch.set_facecolor('white')
pie_cols = [BLUE, GREEN, ORANGE, RED][:len(cat)]

wedges1, _, autos1 = ax3a.pie(cat['vol_cy'], labels=None, autopct='%1.1f%%', colors=pie_cols, startangle=90,
                               wedgeprops=dict(width=0.55, edgecolor='white'), pctdistance=0.75)
for at in autos1: at.set_fontsize(9); at.set_color('white'); at.set_fontweight('bold')
ax3a.legend(labels4, loc='lower center', bbox_to_anchor=(0.5, -0.12), ncol=2, fontsize=9)
ax3a.set_title('Volume Mix (CY)', fontsize=11, fontweight='bold', color=BLUE, pad=8)

wedges2, _, autos2 = ax3b.pie(cat['stns'], labels=None, autopct='%1.1f%%', colors=pie_cols, startangle=90,
                               wedgeprops=dict(width=0.55, edgecolor='white'), pctdistance=0.75)
for at in autos2: at.set_fontsize(9); at.set_color('white'); at.set_fontweight('bold')
ax3b.legend(labels4, loc='lower center', bbox_to_anchor=(0.5, -0.12), ncol=2, fontsize=9)
ax3b.set_title('Station Count Mix', fontsize=11, fontweight='bold', color=BLUE, pad=8)

fig3.suptitle('Volume vs Station Distribution by Category', fontsize=11, color='#444444', y=1.01)
plt.tight_layout()
chart3_buf = img_buf(fig3)

# ── BUILD WORD DOC ────────────────────────────────────────────────────────────
print("Building Word document…")
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21); sec.page_height = Cm(29.7)
sec.left_margin = sec.right_margin = Cm(1.8)
sec.top_margin = sec.bottom_margin = Cm(1.6)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('PSO Retail Fuels')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = W_BLUE

p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Station Coverage & Volume Analysis by Category and City')
r2.font.size = Pt(13); r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(2)
r3 = p3.add_run(REPORT_DATE)
r3.bold = True; r3.font.size = Pt(11); r3.font.color.rgb = W_BLUE

doc.add_paragraph()

kpi_tbl = doc.add_table(rows=2, cols=4)
kpi_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tot_vol = fuels['SalesLtr_CY'].sum()
tot_vol_sply = fuels['SalesLtr_SPLY'].sum()
tot_stns = fuels['Customer Number'].nunique()
tot_cities = fuels['CityNorm'].nunique()
kpi_heads = ['Total Volume CY', 'Stations Active', 'Cities Covered', 'Avg KL / Station']
kpi_vals = [f"{tot_vol/1e6:.3f} ML", f"{tot_stns:,}", str(tot_cities), f"{tot_vol/tot_stns/1000:.2f} KL" if tot_stns else "N/A"]
kpi_sub = [f"{(tot_vol-tot_vol_sply)/tot_vol_sply*100:+.1f}% vs SPLY" if tot_vol_sply else "N/A",
           'across all regions', 'nationwide', 'per active outlet']
for ci in range(4):
    set_bg(kpi_tbl.cell(0, ci), '00479D')
    ct(kpi_tbl.cell(0, ci), kpi_heads[ci], bold=True, size=9, color=W_WHITE)
    set_bg(kpi_tbl.cell(1, ci), 'F2F2F2')
    kpi_tbl.cell(1, ci).text = ''
    vp = kpi_tbl.cell(1, ci).paragraphs[0]; vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run(kpi_vals[ci]); vr.bold = True; vr.font.size = Pt(11); vr.font.color.rgb = W_BLUE
    vp.add_run('\n')
    sr = vp.add_run(kpi_sub[ci]); sr.font.size = Pt(8); sr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()

# SECTION 1 — BY CATEGORY
heading(doc, '1.  Volume & Station Coverage by Fuel Category', level=1)
top_cat = cat.iloc[0]
most_stns_cat = cat['stns'].idxmax()
body(doc,
    f'PSO sells fuels through its retail network across {len(cat)} categories. '
    f'{top_cat.name} accounts for {top_cat.vol_sh:.0f}% of total volume ({chg(top_cat.vol_chg)} vs SPLY). '
    f'{most_stns_cat} has the largest station footprint ({int(cat.loc[most_stns_cat, "stns"]):,} stations) '
    f'selling {cat.loc[most_stns_cat, "vol_ps"]/1000:.2f} KL per station on average.', size=10)

doc.add_picture(chart1_buf, width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

heading(doc, '1.1  Category Summary Table', level=2)
col_h = ['Category', 'Vol CY (KL)', 'Vol SPLY (KL)', 'vs SPLY', 'Mix %', 'Stations', 'KL / Station', 'Margin / L (PKR)']
tbl1 = doc.add_table(rows=len(cat) + 2, cols=len(col_h))
tbl1.style = 'Table Grid'; tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(col_h):
    set_bg(tbl1.cell(0, ci), '00479D')
    ct(tbl1.cell(0, ci), h, bold=True, size=9, color=W_WHITE)
for ri, (cat_name, r) in enumerate(cat.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    chg_c = W_GREEN if r.vol_chg >= 0 else W_RED
    vals = [cat_name, f'{r.vol_cy/1000:,.1f}', f'{r.vol_sply/1000:,.1f}', chg(r.vol_chg),
            f'{r.vol_sh:.1f}%', f'{int(r.stns):,}', f'{r.vol_ps/1000:.2f}', f'{r.mgn_pl:.0f}']
    for ci, v in enumerate(vals):
        set_bg(tbl1.cell(ri, ci), bg)
        fc = chg_c if ci == 3 else None
        al = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
        ct(tbl1.cell(ri, ci), v, size=9, color=fc, align=al)
tot_r = len(cat) + 1
set_bg(tbl1.cell(tot_r, 0), '1F3864')
ct(tbl1.cell(tot_r, 0), 'TOTAL', bold=True, size=9, color=W_WHITE, align=WD_ALIGN_PARAGRAPH.LEFT)
tot_vals = [f'{tot_vol/1000:,.1f}', f'{tot_vol_sply/1000:,.1f}',
            chg((tot_vol-tot_vol_sply)/tot_vol_sply*100) if tot_vol_sply else 'N/A', '100.0%',
            f'{tot_stns:,}', f'{tot_vol/tot_stns/1000:.2f}' if tot_stns else 'N/A',
            f'{fuels["NetMargin_CY"].sum()/tot_vol:.0f}' if tot_vol else 'N/A']
for ci, v in enumerate(tot_vals, 1):
    set_bg(tbl1.cell(tot_r, ci), '1F3864')
    ct(tbl1.cell(tot_r, ci), v, bold=True, size=9, color=W_WHITE)

doc.add_paragraph()
doc.add_picture(chart3_buf, width=Inches(5.6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

vol_share_gap = cat.assign(stn_sh=lambda d: d.stns / d.stns.sum() * 100)
vol_share_gap['gap'] = vol_share_gap['vol_sh'] - vol_share_gap['stn_sh']
under_indexed = vol_share_gap.sort_values('gap').iloc[0]
over_indexed = vol_share_gap.sort_values('gap').iloc[-1]
body(doc,
    f'The two donuts above compare volume share against station-count share by category. '
    f'{over_indexed.name} over-indexes — {over_indexed.stn_sh:.0f}% of stations deliver '
    f'{over_indexed.vol_sh:.0f}% of volume — while {under_indexed.name} under-indexes at '
    f'{under_indexed.stn_sh:.0f}% of stations for only {under_indexed.vol_sh:.0f}% of volume, '
    f'pointing to under-utilised distribution capacity in that category.',
    size=9, before=5)

doc.add_page_break()

# SECTION 2 — BY CITY
heading(doc, '2.  Top 15 Cities — Active Stations & Productivity', level=1)
body(doc,
    'The top 15 cities by fuels volume collectively account for a majority of national fuels sales. '
    'This section examines how many stations are active in each city and how productively they sell '
    '(volume per station). High productivity per station indicates strong demand or superior product '
    'placement; low productivity may signal distribution gaps or competitive pressure.', size=10)

doc.add_picture(chart2_buf, width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

heading(doc, '2.1  Top 15 Cities — Detailed Table', level=2)
col_h2 = ['#', 'City', 'Region', 'Stations', 'Vol CY (KL)', 'Vol SPLY (KL)', 'vs SPLY', 'Mix %', 'KL / Station', 'Mgn/L (PKR)']
tbl2 = doc.add_table(rows=len(city) + 1, cols=len(col_h2))
tbl2.style = 'Table Grid'; tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(col_h2):
    set_bg(tbl2.cell(0, ci), '00479D')
    ct(tbl2.cell(0, ci), h, bold=True, size=8.5, color=W_WHITE)
for ri, (city_name, r) in enumerate(city.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    chg_c = W_GREEN if r.vol_chg >= 0 else W_RED
    vals = [str(ri), city_name, r.get('Region', ''), f'{int(r.stns):,}', f'{r.vol_cy/1000:,.1f}',
            f'{r.vol_sply/1000:,.1f}', chg(r.vol_chg), f'{r.vol_sh:.1f}%', f'{r.vol_ps/1000:.2f}', f'{r.mgn_pl:.0f}']
    for ci, v in enumerate(vals):
        set_bg(tbl2.cell(ri, ci), bg)
        fc = chg_c if ci == 6 else None
        al = WD_ALIGN_PARAGRAPH.LEFT if ci in [1, 2] else WD_ALIGN_PARAGRAPH.CENTER
        ct(tbl2.cell(ri, ci), v, size=8.5, color=fc, align=al)

doc.add_paragraph()

# ── Key Findings (dynamically computed) ───────────────────────────────────────
heading(doc, '2.2  Key Findings', level=2)

city_by_stns = city.sort_values('stns', ascending=False)
city_by_vps = city.sort_values('vol_ps', ascending=False)
growers = city_all[city_all['vol_chg'] > 0].sort_values('vol_chg', ascending=False).head(3)
decliners_top15 = city[city['vol_chg'] < 0].sort_values('vol_chg')

findings = []
top2_stns = city_by_stns.head(2)
findings.append((
    f'Largest Networks — {" & ".join(top2_stns.index.tolist())}',
    f'{top2_stns.index[0]} ({int(top2_stns.iloc[0].stns)} stations) and {top2_stns.index[1]} '
    f'({int(top2_stns.iloc[1].stns)} stations) have the largest fuels networks. Their productivity per '
    f'station is {top2_stns.iloc[0].vol_ps/1000:.2f} KL and {top2_stns.iloc[1].vol_ps/1000:.2f} KL respectively. '
    f'{top2_stns.index[0]} is {chg(top2_stns.iloc[0].vol_chg)} vs SPLY, {top2_stns.index[1]} is '
    f'{chg(top2_stns.iloc[1].vol_chg)}.'
))
if len(city_by_vps):
    top_vps = city_by_vps.iloc[0]
    second_vps = city_by_vps.iloc[1] if len(city_by_vps) > 1 else top_vps
    findings.append((
        f'Highest Productivity — {city_by_vps.index[0]}',
        f'{city_by_vps.index[0]} stands out with {int(top_vps.stns)} stations yet '
        f'{top_vps.vol_ps/1000:.2f} KL per station — {"the highest in the top 15" if top_vps.vol_ps >= second_vps.vol_ps else ""}. '
        f'Volume there is {chg(top_vps.vol_chg)} vs SPLY. '
        f'{"This market is high-value but fragile and warrants priority sales attention." if top_vps.vol_chg < 0 else "This is a strong, well-placed network worth studying for best practice."}'
    ))
if len(growers):
    findings.append((
        f'High-Growth Cities — {", ".join(growers.index.tolist())}',
        f'{", ".join(f"{c} ({chg(r.vol_chg)})" for c, r in growers.iterrows())} are the fastest-growing '
        f'fuels markets nationally. These should receive priority stock allocation and new dealer '
        f'appointments to sustain growth.'
    ))
if len(decliners_top15):
    findings.append((
        f'Declining Cities in Top 15 — {", ".join(decliners_top15.index.tolist())}',
        f'{", ".join(f"{c} ({chg(r.vol_chg)}, {int(r.stns)} stations)" for c, r in decliners_top15.iterrows())} '
        f'are the only top-15 cities in decline. Given their large networks, even a small per-station '
        f'improvement would add significant volume.'
    ))

for title, text in findings:
    fp = doc.add_paragraph()
    fp.paragraph_format.space_before = Pt(6)
    fp.paragraph_format.left_indent = Inches(0.1)
    fr = fp.add_run(f'► {title}')
    fr.bold = True; fr.font.size = Pt(10); fr.font.color.rgb = W_BLUE

    fb = doc.add_paragraph(text)
    fb.paragraph_format.left_indent = Inches(0.25)
    fb.paragraph_format.space_before = Pt(2)
    fb.paragraph_format.space_after = Pt(5)
    for r in fb.runs: r.font.size = Pt(10)

doc.add_page_break()

# SECTION 3 — STATION EFFICIENCY RANKING
heading(doc, '3.  City Efficiency Ranking — Volume per Station', level=1)
body(doc,
    'Volume per station (KL/station) is a measure of network productivity. The table below ranks all '
    'top-15 cities by this metric, helping identify markets where the existing network is being fully '
    'utilised versus markets where growth can come from adding new outlets.', size=10)

eff = city.copy().reset_index()
eff['rank_vps'] = eff['vol_ps'].rank(ascending=False).astype(int)
eff = eff.sort_values('vol_ps', ascending=False)

eff_h = ['Efficiency Rank', 'City', 'Stations', 'KL / Station', 'Total Vol CY (KL)', 'YoY Chg', 'Interpretation']
tbl3 = doc.add_table(rows=len(eff) + 1, cols=len(eff_h))
tbl3.style = 'Table Grid'; tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(eff_h):
    set_bg(tbl3.cell(0, ci), '00479D')
    ct(tbl3.cell(0, ci), h, bold=True, size=8.5, color=W_WHITE)

def interpret(vps, stns, vol_chg):
    if vps > 8:
        return 'Very high productivity — consider network expansion'
    elif vps > 5:
        return 'High productivity — strong demand per outlet'
    elif vol_chg < 0:
        return 'Below average & declining — priority recovery'
    elif stns > 100:
        return 'Large network — focus on per-station uplift'
    else:
        return 'Moderate — growth through network + uplift'

for ri, (_, r) in enumerate(eff.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    if ri <= 3:
        bg = 'DDEEFF'
    vps = r.vol_ps / 1000
    chg_c = W_GREEN if r.vol_chg >= 0 else W_RED
    vals = [str(ri), r['CityNorm'], f'{int(r.stns):,}', f'{vps:.2f}', f'{r.vol_cy/1000:,.1f}',
            chg(r.vol_chg), interpret(vps, r.stns, r.vol_chg)]
    for ci, v in enumerate(vals):
        set_bg(tbl3.cell(ri, ci), bg)
        fc = chg_c if ci == 5 else None
        al = WD_ALIGN_PARAGRAPH.LEFT if ci in [1, 6] else WD_ALIGN_PARAGRAPH.CENTER
        bold = (ri <= 3 and ci == 3)
        ct(tbl3.cell(ri, ci), v, size=8.5, color=fc, align=al, bold=bold)

doc.add_paragraph()
body(doc,
    'Cities highlighted in blue (top 3 by KL/station) represent the most productive markets. These '
    'should be benchmarks for the broader network improvement programme. Cities with large station '
    'counts but moderate productivity offer the biggest absolute volume upside if per-station '
    'productivity can be improved even marginally.', size=10)

out = out_path('PSO_Fuels_Stations_Report', 'docx', df)
doc.save(out)
print(f"\nSaved: {out}")
