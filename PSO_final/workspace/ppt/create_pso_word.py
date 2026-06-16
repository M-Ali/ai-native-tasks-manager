"""
PSO OMC Analytics — Word Document Generator
10 Months FY26 | Pakistan State Oil
"""

import json
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AI_JSON = Path("D:/Personal/PSO_final/reports/PSO_AI_10M_FY26_2026-06-14.json")
OUTPUT  = Path(f"D:/Personal/PSO_final/reports/PSO_Report_10M_FY26.docx")

# ── Colours ────────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x2A, 0x4A)
GOLD   = RGBColor(0xC9, 0xA0, 0x30)
RED    = RGBColor(0xB2, 0x22, 0x22)
GREEN  = RGBColor(0x1A, 0x7A, 0x44)
GREY   = RGBColor(0x6B, 0x7D, 0x8A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LTBLUE = RGBColor(0xF5, 0xF7, 0xFA)

def hex_to_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)

def set_cell_border(cell, side='left', color='C9A030', sz='18'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    border = OxmlElement(f'w:{side}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), sz)
    border.set(qn('w:space'), '0')
    border.set(qn('w:color'), color)
    tcBorders.append(border)

# ── Helpers ────────────────────────────────────────────────────────────────────
def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if color:
        run.font.color.rgb = color
    return p

def para(doc, text='', bold=False, color=None, size=10, align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.size  = Pt(size)
        run.font.bold  = bold
        if color:
            run.font.color.rgb = color
    return p

def bullet(doc, text, color=None, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        if color:
            run.font.color.rgb = color
        p.add_run(text)
    else:
        run = p.add_run(text)
        if color:
            run.font.color.rgb = color
    return p

def divider(doc, color='C9A030'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '12')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def kpi_row(doc, label, cy, ly, chg, chg_color):
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    cells = tbl.rows[0].cells
    for c in cells:
        set_cell_bg(c, '#F5F7FA')
    cells[0].text = label
    cells[1].text = str(cy)
    cells[2].text = str(ly)
    cells[3].text = str(chg)
    for i, c in enumerate(cells):
        for p in c.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                if i == 3:
                    run.font.color.rgb = chg_color
                    run.font.bold = True
    return tbl

def section_heading(doc, text, color=NAVY):
    divider(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.color.rgb = color
    return p

# ── Document Setup ─────────────────────────────────────────────────────────────
def build(doc, ai):
    # ── Page Setup ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ── Cover ────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PSO OMC ANALYTICS")
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Performance & Strategic Insight Report")
    r2.font.size  = Pt(16)
    r2.font.color.rgb = GOLD

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(8)
    r3 = p3.add_run("10 Months FY26  |  Pakistan State Oil  |  June 2026")
    r3.font.size  = Pt(10)
    r3.font.color.rgb = GREY

    divider(doc)
    para(doc, "CONFIDENTIAL — For Management Use Only", bold=True, color=GREY, size=9,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=2)
    para(doc, f"Data: 27,344 transactions  |  21,334 Retail rows  |  3 Regions  |  450+ Cities",
         color=GREY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

    doc.add_page_break()

    # ── 1. EXECUTIVE SUMMARY ─────────────────────────────────────────────────
    section_heading(doc, "1. Executive Summary")

    # Headline KPIs
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = 'Table Grid'
    headers = ["GRS YoY", "Volume YoY", "Net Margin YoY", "Lube NMgn/Ltr YoY"]
    values  = ["+23.6%", "-1.85%", "-4.86%", "-5.67%"]
    colors  = ["C9A030", "B22222", "B22222", "B22222"]
    for i, cell in enumerate(tbl.rows[0].cells):
        set_cell_bg(cell, '#1B2A4A')
        p = cell.paragraphs[0]
        run = p.add_run(headers[i])
        run.font.size  = Pt(8)
        run.font.bold  = True
        run.font.color.rgb = GOLD
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, cell in enumerate(tbl.rows[1].cells):
        set_cell_bg(cell, '#F5F7FA')
        p = cell.paragraphs[0]
        run = p.add_run(values[i])
        run.font.size  = Pt(14)
        run.font.bold  = True
        run.font.color.rgb = hex_to_rgb(colors[i])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, '', space_before=0, space_after=8)

    para(doc, "Revenue growth of +23.6% is entirely driven by price inflation — real demand is contracting. "
         "Across all retail fuel segments, net margin per litre is declining while discounting continues. "
         "The lubricants business is growing in volume but losing per-litre profitability due to a structural shift toward LOW GRADE products.",
         size=10, space_after=6)

    para(doc, "Top 3 Concerns:", bold=True, color=NAVY, size=10, space_after=3)
    bullet(doc, " Central Petrol NMgn declined PKR 1.23B YoY (-10.5%) — largest single-region margin loss in the portfolio", bold_prefix="1.")
    bullet(doc, " LOW GRADE lubes now 54.2% of lube volume at PKR 116/ltr — crowding out PKR 348-501/ltr premium categories", bold_prefix="2.")
    bullet(doc, " R95 premium petrol NMgn/ltr collapsed 14-27% across all regions with zero discounts applied", bold_prefix="3.")
    para(doc, '', space_after=4)

    # ── 2. PORTFOLIO OVERVIEW ────────────────────────────────────────────────
    section_heading(doc, "2. Portfolio Overview — All Channels")

    para(doc, "PSO's consolidated portfolio shows 10-channel diversification with growth concentrated in non-retail segments. "
         "Retail Business remains the largest single channel at 59.9% of GRS but is losing margin.",
         size=10, space_after=6)

    channel_tbl = doc.add_table(rows=8, cols=4)
    channel_tbl.style = 'Table Grid'
    ch_headers = ["Channel", "GRS Chg %", "Vol Chg %", "Signal"]
    rows_data = [
        ("Chemicals", "+424.7%", "+338.2%", "STRONG GROWTH"),
        ("Power Projects", "+113.0%", "+29.0%", "STRONG GROWTH"),
        ("Marine", "+55.6%", "+16.7%", "GROWING"),
        ("PR Strategic Accts.", "+52.5%", "+18.7%", "GROWING"),
        ("Retail Business", "+23.6%", "-1.85%", "VOLUME EROSION"),
        ("LNG", "-11.8%", "-12.9%", "DECLINING"),
        ("Agency", "-9.8%", "-48.4%", "DECLINING"),
    ]
    for i, cell in enumerate(channel_tbl.rows[0].cells):
        set_cell_bg(cell, '#1B2A4A')
        run = cell.paragraphs[0].add_run(ch_headers[i])
        run.font.size = Pt(8); run.font.bold = True
        run.font.color.rgb = GOLD
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    signal_colors = {'STRONG GROWTH': '1A7A44', 'GROWING': '2E5B9A',
                     'VOLUME EROSION': 'B22222', 'DECLINING': 'B22222'}
    for r_i, (ch, grs, vol, sig) in enumerate(rows_data):
        row = channel_tbl.rows[r_i + 1]
        bg = 'F5F7FA' if r_i % 2 == 0 else 'FFFFFF'
        for c in row.cells:
            set_cell_bg(c, bg)
        row.cells[0].paragraphs[0].add_run(ch).font.size = Pt(9)
        for i, val in enumerate([grs, vol]):
            run = row.cells[i+1].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            row.cells[i+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig_run = row.cells[3].paragraphs[0].add_run(sig)
        sig_run.font.size = Pt(8); sig_run.font.bold = True
        sig_run.font.color.rgb = hex_to_rgb(signal_colors.get(sig, '333333'))
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, '', space_after=8)

    # ── 3. DIESEL ANALYSIS ───────────────────────────────────────────────────
    section_heading(doc, "3. Diesel Analysis")

    para(doc, "Total diesel volume: 2,603 ML (−2.58% YoY). Revenue grew +24.2% — a pure price effect. "
         "North is the only growing region; Central and South are both losing volume.", size=10, space_after=6)

    diesel_tbl = doc.add_table(rows=4, cols=5)
    diesel_tbl.style = 'Table Grid'
    d_headers = ["Region", "Vol CY (ML)", "Vol Chg %", "NMgn/Ltr CY", "NMgn/Ltr LY"]
    d_rows = [
        ("Central", "1,156.1", "-4.71%", "6.408", "6.148"),
        ("North",   "941.9",  "+1.68%", "6.187", "6.668"),
        ("South",   "505.1",  "-5.16%", "7.466", "7.131"),
    ]
    for i, cell in enumerate(diesel_tbl.rows[0].cells):
        set_cell_bg(cell, '#1B2A4A')
        run = cell.paragraphs[0].add_run(d_headers[i])
        run.font.size = Pt(8); run.font.bold = True
        run.font.color.rgb = GOLD
    for r_i, (reg, vol, chg, cy, ly) in enumerate(d_rows):
        row = diesel_tbl.rows[r_i + 1]
        set_cell_bg(row.cells[0], 'F5F7FA')
        row.cells[0].paragraphs[0].add_run(reg).font.bold = True
        for i, val in enumerate([vol, chg, cy, ly]):
            run = row.cells[i+1].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if i == 1:  # chg col
                run.font.color.rgb = GREEN if val.startswith('+') else RED
                run.font.bold = True
    para(doc, '', space_after=6)

    bullet(doc, "North diesel is the only growing region (+1.68%) but pays PKR 1.683/ltr in discounts — the highest nationally. "
           "South achieves 20% better NMgn/ltr at PKR 0.40/ltr discount.", bold_prefix="Key insight: ")
    bullet(doc, "Okara -20.0%, Rawalpindi -16.7%, Bahawalpur -13.8% are the worst city-level NMgn/ltr performers.", bold_prefix="City pressure: ")
    para(doc, '', space_after=4)

    # ── 4. PETROL ANALYSIS ───────────────────────────────────────────────────
    section_heading(doc, "4. Petrol Analysis")

    para(doc, "Total petrol volume: 3,469 ML (−1.33% YoY). All three regions are simultaneously losing "
         "both volume share and NMgn/ltr — a portfolio-wide deterioration with no exception.", size=10, space_after=6)

    petrol_tbl = doc.add_table(rows=4, cols=5)
    petrol_tbl.style = 'Table Grid'
    p_headers = ["Region", "Vol CY (ML)", "Vol Chg %", "NMgn/Ltr CY", "NMgn/Ltr LY"]
    p_rows = [
        ("Central", "1,557.9", "-2.5%",  "6.723", "7.319"),
        ("North",   "1,109.9", "+2.2%",  "6.622", "7.471"),
        ("South",   "801.5",   "-3.7%",  "7.973", "8.119"),
    ]
    for i, cell in enumerate(petrol_tbl.rows[0].cells):
        set_cell_bg(cell, '#1B2A4A')
        run = cell.paragraphs[0].add_run(p_headers[i])
        run.font.size = Pt(8); run.font.bold = True
        run.font.color.rgb = GOLD
    for r_i, (reg, vol, chg, cy, ly) in enumerate(p_rows):
        row = petrol_tbl.rows[r_i + 1]
        set_cell_bg(row.cells[0], 'F5F7FA')
        row.cells[0].paragraphs[0].add_run(reg).font.bold = True
        for i, val in enumerate([vol, chg, cy, ly]):
            run = row.cells[i+1].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if i == 1:
                run.font.color.rgb = GREEN if val.startswith('+') else RED
                run.font.bold = True
    para(doc, '', space_after=6)

    para(doc, "R95 Premium Petrol — Critical Signal:", bold=True, color=RED, size=10, space_after=3)
    r95_tbl = doc.add_table(rows=4, cols=4)
    r95_tbl.style = 'Table Grid'
    r95_h = ["Region", "R95 NMgn/Ltr CY", "R95 NMgn/Ltr LY", "Change"]
    r95_r = [
        ("Central", "9.487",  "12.493", "-24.1%"),
        ("North",   "7.966",  "10.885", "-26.8%"),
        ("South",   "17.360", "20.102", "-13.6%"),
    ]
    for i, cell in enumerate(r95_tbl.rows[0].cells):
        set_cell_bg(cell, '#B22222')
        run = cell.paragraphs[0].add_run(r95_h[i])
        run.font.size = Pt(8); run.font.bold = True
        run.font.color.rgb = WHITE
    for r_i, row_data in enumerate(r95_r):
        row = r95_tbl.rows[r_i + 1]
        for j, val in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if j == 3:
                run.font.color.rgb = RED; run.font.bold = True
    para(doc, '', space_after=4)
    bullet(doc, " Zero discounts applied to R95 in any region. The collapse is entirely primary-margin driven. "
           "Requires immediate cost/pricing review.", bold_prefix="WARNING:")

    para(doc, "Central Punjab Cluster — Volume Decline:", bold=True, color=RED, size=10, space_before=8, space_after=3)
    cluster_data = [
        ("Sahiwal",     "35.4", "46.3", "-23.5%"),
        ("Kasur",       "35.4", "42.5", "-16.7%"),
        ("Sheikhupura", "35.0", "41.0", "-14.8%"),
        ("Gujranwala",  "90.9", "102.0","-10.9%"),
        ("Rawalpindi",  "122.8","136.1", "-9.8%"),
    ]
    cl_tbl = doc.add_table(rows=6, cols=4)
    cl_tbl.style = 'Table Grid'
    cl_h = ["City", "Vol CY (ML)", "Vol LY (ML)", "Vol Chg %"]
    for i, cell in enumerate(cl_tbl.rows[0].cells):
        set_cell_bg(cell, '#B22222')
        run = cell.paragraphs[0].add_run(cl_h[i])
        run.font.size = Pt(8); run.font.bold = True
        run.font.color.rgb = WHITE
    for r_i, row_data in enumerate(cluster_data):
        row = cl_tbl.rows[r_i + 1]
        for j, val in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if j == 3:
                run.font.color.rgb = RED; run.font.bold = True
    para(doc, '', space_after=4)

    # ── 5. LUBRICANTS — MIX PROBLEM ──────────────────────────────────────────
    section_heading(doc, "5. Lubricants — The Mix Shift Problem")

    para(doc, "Total lube volume grew +7.6% YoY (12.4 → 13.3 ML). Revenue grew +16.6%. "
         "Yet NMgn/ltr FELL from 239.5 to 225.9 PKR/ltr (−5.67%). "
         "The cause: LOW GRADE volume surged +24.4% and now accounts for 54.2% of total lube volume, "
         "crowding out high-margin DEO, PCMO, and MCO.", size=10, space_after=6)

    lube_tbl = doc.add_table(rows=6, cols=5)
    lube_tbl.style = 'Table Grid'
    l_headers = ["Category", "Vol CY (ML)", "Vol Chg %", "NMgn/Ltr CY (PKR)", "Mix Share CY"]
    l_rows = [
        ("LOW GRADE", "7.204", "+24.4%", "116.0", "54.2%"),
        ("DEO",       "3.895",  "-6.7%", "348.1", "29.3%"),
        ("MCO",       "1.218", "-11.1%", "268.9",  "9.2%"),
        ("PCMO",      "0.954",  "-4.7%", "500.9",  "7.2%"),
        ("OTHERS",    "0.027",  "+3.8%",     "–",  "0.2%"),
    ]
    for i, cell in enumerate(lube_tbl.rows[0].cells):
        set_cell_bg(cell, '#1B2A4A')
        run = cell.paragraphs[0].add_run(l_headers[i])
        run.font.size = Pt(8); run.font.bold = True
        run.font.color.rgb = GOLD
    for r_i, (cat, vol, chg, mgn, mix) in enumerate(l_rows):
        row = lube_tbl.rows[r_i + 1]
        bg = 'FFF0F0' if cat == 'LOW GRADE' else ('F5F7FA' if r_i % 2 == 0 else 'FFFFFF')
        for c in row.cells:
            set_cell_bg(c, bg)
        cat_run = row.cells[0].paragraphs[0].add_run(cat)
        cat_run.font.bold = True
        cat_run.font.size = Pt(9)
        if cat == 'LOW GRADE':
            cat_run.font.color.rgb = RED
        for i, val in enumerate([vol, chg, mgn, mix]):
            run = row.cells[i+1].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            row.cells[i+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 1:
                run.font.color.rgb = GREEN if val.startswith('+') else RED
                run.font.bold = True
    para(doc, '', space_after=6)

    para(doc, "The Margin Gap:", bold=True, color=NAVY, size=10, space_after=3)
    bullet(doc, "PCMO vs LOW GRADE: PKR 500.9 vs PKR 116.0 — a PKR 384.9/ltr difference (4.3× more profitable)")
    bullet(doc, "Each 1 ML of volume shifted from DEO to LOW GRADE costs PSO approximately PKR 232M in net margin")
    bullet(doc, "Zero discounts across all lube categories — this is 100% a product mix problem, not a pricing problem")
    para(doc, '', space_after=4)

    # ── 6. LUBRICANTS — GEOGRAPHY ────────────────────────────────────────────
    section_heading(doc, "6. Lubricants — Geographic Analysis")

    para(doc, "Central region (6.3 ML, 47.4% of total lube volume) has the worst mix — LOW GRADE dominant. "
         "North shows the most disciplined mix (NMgn/ltr decline only −1.5% vs −9.2% in Central). "
         "South's PCMO segment delivers PKR 518/ltr — the highest margin pocket in the entire dataset.",
         size=10, space_after=6)

    reg_lube = doc.add_table(rows=4, cols=4)
    reg_lube.style = 'Table Grid'
    rl_h = ["Region", "Vol CY (ML)", "NMgn/Ltr CY", "Signal"]
    rl_r = [
        ("Central", "6.345", "PKR 198 (↓9.2%)", "LOW GRADE dominant — worst mix"),
        ("North",   "3.376", "PKR 261 (↓1.5%)", "Most disciplined mix"),
        ("South",   "3.580", "PKR 241 (↓4.4%)", "PCMO PKR 518/ltr — best margin pocket"),
    ]
    for i, cell in enumerate(reg_lube.rows[0].cells):
        set_cell_bg(cell, '#1B2A4A')
        run = cell.paragraphs[0].add_run(rl_h[i])
        run.font.size = Pt(8); run.font.bold = True
        run.font.color.rgb = GOLD
    for r_i, row_data in enumerate(rl_r):
        row = reg_lube.rows[r_i + 1]
        set_cell_bg(row.cells[0], 'F5F7FA')
        for j, val in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if j == 0:
                run.font.bold = True
    para(doc, '', space_after=6)

    bullet(doc, "Islamabad's balanced lube mix (DEO ≈ PCMO ≈ LOW GRADE thirds) is the benchmark model for large cities.")
    bullet(doc, "Lahore and Faisalabad have LOW GRADE at 3x the rate of DEO — both are priority cities for mix improvement.")
    bullet(doc, "South PCMO is the most underleveraged opportunity: PKR 518/ltr at only 0.342 ML total volume.")
    para(doc, '', space_after=4)

    # ── 7. REGIONAL PERFORMANCE GAPS ─────────────────────────────────────────
    section_heading(doc, "7. Regional Performance Gaps")

    para(doc, "The three regions operate with fundamentally different discount disciplines and achieve "
         "proportionally different NMgn/ltr outcomes. South's low-discount model is the most financially efficient.",
         size=10, space_after=6)

    matrix_tbl = doc.add_table(rows=10, cols=6)
    matrix_tbl.style = 'Table Grid'
    m_h = ["Region", "Segment", "Vol CY (ML)", "Vol Chg %", "NMgn/Ltr CY", "Signal"]
    m_r = [
        ("Central","Petrol","1,557.9","-2.5%", "6.72","CONCERN"),
        ("Central","Diesel","1,156.1","-4.7%", "6.41","CONCERN"),
        ("Central","Lubes",    "6.3","+10.6%","PKR 198","MIX ISSUE"),
        ("North",  "Petrol","1,109.9","+2.2%", "6.62","MARGIN RISK"),
        ("North",  "Diesel",  "941.9","+1.7%", "6.19","MARGIN RISK"),
        ("North",  "Lubes",     "3.4", "+4.9%","PKR 261","HEALTHY"),
        ("South",  "Petrol",  "801.5","-3.7%", "7.97","CONCERN"),
        ("South",  "Diesel",  "505.1","-5.2%", "7.47","CONCERN"),
        ("South",  "Lubes",     "3.6", "+5.1%","PKR 241","MIX ISSUE"),
    ]
    for i, cell in enumerate(matrix_tbl.rows[0].cells):
        set_cell_bg(cell, '#1B2A4A')
        run = cell.paragraphs[0].add_run(m_h[i])
        run.font.size = Pt(8); run.font.bold = True
        run.font.color.rgb = GOLD
    sig_c = {"CONCERN":"FFF0F0","MARGIN RISK":"FEF9EC","HEALTHY":"E8F8EE","MIX ISSUE":"FEF9EC"}
    sig_fc= {"CONCERN":RED,"MARGIN RISK":hex_to_rgb("B8860B"),"HEALTHY":GREEN,"MIX ISSUE":hex_to_rgb("B8860B")}
    for r_i, row_data in enumerate(m_r):
        row = matrix_tbl.rows[r_i + 1]
        bg = 'F5F7FA' if r_i % 2 == 0 else 'FFFFFF'
        for c in row.cells:
            set_cell_bg(c, bg)
        for j, val in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if j == 3:
                run.font.color.rgb = GREEN if val.startswith('+') else RED
                run.font.bold = True
            if j == 5:
                run.font.color.rgb = sig_fc.get(val, hex_to_rgb("333333"))
                run.font.bold = True
    para(doc, '', space_after=6)

    bullet(doc, "NMgn/ltr gap between South Petrol (7.97) and North Petrol (6.62) = PKR 1.35/ltr. "
           "Replicating South discipline in North = ~PKR 1.5B additional annual margin at North's volume.")
    bullet(doc, "North diesel discount spend: PKR 1.683/ltr × 941.9 ML = PKR 1.58B. "
           "Cutting to South levels (PKR 0.40/ltr) saves PKR 1.18B — volume elasticity is the key unknown.")
    para(doc, '', space_after=4)

    # ── 8. CITY CONCENTRATION ────────────────────────────────────────────────
    section_heading(doc, "8. City Concentration — Pareto Analysis")

    para(doc, "PSO's retail volume is highly concentrated. 17 cities account for 50.6% of total retail volume. "
         "The top 2 cities alone (Karachi + Lahore) = 19.7% of national retail. "
         "This concentration creates both risk (when key cities decline) and opportunity (targeted interventions have outsized impact).",
         size=10, space_after=6)

    city_tbl = doc.add_table(rows=13, cols=3)
    city_tbl.style = 'Table Grid'
    city_h = ["City", "Retail Vol CY (ML)", "Status"]
    city_r = [
        ("Karachi",     "644.9", "Stable (+0.7%)"),
        ("Lahore",      "555.4", "Growing (+2.9%)"),
        ("Multan",      "197.9", "Stable"),
        ("Rawalpindi",  "188.2", "DECLINING (-11.2%)"),
        ("Faisalabad",  "188.1", "Stable"),
        ("Islamabad",   "186.1", "Stable"),
        ("Peshawar",    "182.5", "Growing (+2.9%)"),
        ("Gujranwala",  "130.4", "Stable"),
        ("Bahawalpur",  "112.6", "DECLINING (-15.7%)"),
        ("Gujrat",      "105.9", "Growing (+7.6%)"),
        ("Sialkot",     "100.1", "Stable"),
        ("R.Y. Khan",    "89.3", "Stable"),
    ]
    for i, cell in enumerate(city_tbl.rows[0].cells):
        set_cell_bg(cell, '#1B2A4A')
        run = cell.paragraphs[0].add_run(city_h[i])
        run.font.size = Pt(8); run.font.bold = True
        run.font.color.rgb = GOLD
    for r_i, (city, vol, status) in enumerate(city_r):
        row = city_tbl.rows[r_i + 1]
        bg = 'F5F7FA' if r_i % 2 == 0 else 'FFFFFF'
        for c in row.cells:
            set_cell_bg(c, bg)
        row.cells[0].paragraphs[0].add_run(city).font.size = Pt(9)
        vol_run = row.cells[1].paragraphs[0].add_run(vol)
        vol_run.font.size = Pt(9)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        st_run = row.cells[2].paragraphs[0].add_run(status)
        st_run.font.size = Pt(9)
        if 'DECLINING' in status:
            st_run.font.color.rgb = RED; st_run.font.bold = True
        elif 'Growing' in status:
            st_run.font.color.rgb = GREEN; st_run.font.bold = True
    para(doc, '', space_after=4)

    # ── 9. STRATEGIC IMPLICATIONS ────────────────────────────────────────────
    section_heading(doc, "9. Strategic Implications")

    para(doc, "THE REAL PROBLEM", bold=True, color=RED, size=11, space_after=3)
    bullet(doc, " Revenue growth (+23.6%) is entirely price inflation — real volume is contracting", bold_prefix="•")
    bullet(doc, " Lubricant issue is a mix shift, not volume loss: LOW GRADE +24.4% while DEO/PCMO/MCO all decline", bold_prefix="•")
    bullet(doc, " R95 margin collapsed 24-27% with zero discounts — cost or pricing structure failure", bold_prefix="•")
    bullet(doc, " Central Punjab: 5 cities lost ~57 ML petrol volume — likely outlet loss or competitor entry", bold_prefix="•")
    para(doc, '', space_after=4)

    para(doc, "THE OPPORTUNITY", bold=True, color=hex_to_rgb("B8860B"), size=11, space_after=3)
    bullet(doc, " South's low-discount model generates highest NMgn/ltr (Petrol 7.97, Diesel 7.47) — fully replicable", bold_prefix="•")
    bullet(doc, " Islamabad's balanced lube mix (DEO ≈ PCMO ≈ LOW GRADE) can be exported to Lahore and Faisalabad", bold_prefix="•")
    bullet(doc, " PCMO: PKR 501/ltr margin at only 0.95 ML volume — 4.3x more profitable than LOW GRADE", bold_prefix="•")
    bullet(doc, " South PCMO at PKR 518/ltr is the highest single margin pocket in the entire dataset", bold_prefix="•")
    para(doc, '', space_after=4)

    para(doc, "THE HIDDEN RISK", bold=True, color=RED, size=11, space_after=3)
    bullet(doc, " North is the only volume-growing region (Petrol +2.2%, Diesel +1.7%)", bold_prefix="•")
    bullet(doc, " But North NMgn/ltr is declining fastest: Diesel 6.668→6.187, Petrol 7.471→6.622", bold_prefix="•")
    bullet(doc, " North diesel discount PKR 1.683/ltr — highest nationally — for only +1.7% volume growth", bold_prefix="•")
    bullet(doc, " Risk: Volume growth is being bought at margin cost and may not be sustainable", bold_prefix="•")
    para(doc, '', space_after=4)

    # ── 10. RECOMMENDATIONS ──────────────────────────────────────────────────
    section_heading(doc, "10. Prioritized Recommendations")

    para(doc, "PRIORITY 1 — This Quarter (Immediate)", bold=True, color=RED, size=11, space_after=4)

    p1_items = [
        ("Stop LOW GRADE Lube Mix Erosion",
         "LOW GRADE is 54.2% of lube volume at PKR 116/ltr. Each ML shifted from DEO to LOW GRADE costs PSO ~PKR 232M in net margin.",
         "Set LOW GRADE volume ceiling at 50% of lube mix. Redesign sales incentives to reward DEO/PCMO/MCO sales at 2x weight vs LOW GRADE.",
         "LOW GRADE share held below 50%; premium lube share trend reversed within 2 months."),
        ("Arrest Central Petrol NMgn Collapse",
         "Central Petrol NMgn fell PKR 1.23B YoY (−10.5%). NMgn/ltr: 7.319 → 6.723. Single largest margin loss in the portfolio.",
         "Audit Central petrol dealer discount structure. Benchmark against South model (PKR 0.28/ltr discount, PKR 7.97/ltr NMgn).",
         "Central Petrol NMgn/ltr directional recovery toward PKR 7.00 by FY-end."),
        ("Fix R95 Pricing/Cost Structure",
         "R95 NMgn/ltr has collapsed 14-27% YoY across all regions with ZERO discounts applied. Primary margin destruction.",
         "Immediate cost-of-goods and ex-depot pricing review for R95. Volume growing (+7-11%) — pricing correction is commercially feasible.",
         "R95 NMgn/ltr recovery: Central >12, North >10, South >20 PKR/ltr."),
    ]
    for i, (title, problem, action, kpi) in enumerate(p1_items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(f"{i}. {title}")
        run.font.size = Pt(10); run.font.bold = True; run.font.color.rgb = NAVY
        bullet(doc, f" {problem}", bold_prefix="Problem:")
        bullet(doc, f" {action}", bold_prefix="Action:")
        bullet(doc, f" {kpi}", bold_prefix="KPI:")
    para(doc, '', space_after=6)

    para(doc, "PRIORITY 2 — This Half-Year", bold=True, color=hex_to_rgb("B8860B"), size=11, space_after=4)
    p2_items = [
        ("Recover Central Punjab Cluster Volume",
         "Rawalpindi, Gujranwala, Sheikhupura, Kasur, Sahiwal lost ~57 ML petrol and ~100 ML combined fuel volume. Geographic clustering suggests structural causes.",
         "Outlet-level audit in all 5 cities. Map competitor entry. Verify active PSO outlet count vs LY."),
        ("Review North Diesel Discount ROI",
         "North diesel discount = PKR 1.683/ltr (highest nationally) for only +1.7% volume growth.",
         "Pilot discount reduction from PKR 1.683 to PKR 1.00/ltr in 3 North cities. Measure volume elasticity over 2 months."),
        ("Protect South Margin Premium",
         "South has highest NMgn/ltr across all fuel segments (Petrol 7.97, Diesel 7.47, PCMO 518).",
         "Establish formal South regional pricing protection policy. No discount escalations without Central approval."),
    ]
    for title, problem, action in p2_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(f"• {title}")
        run.font.size = Pt(10); run.font.bold = True; run.font.color.rgb = NAVY
        bullet(doc, f" {problem}", bold_prefix="Context:")
        bullet(doc, f" {action}", bold_prefix="Action:")
    para(doc, '', space_after=6)

    para(doc, "PRIORITY 3 — Next Financial Year", bold=True, color=GREEN, size=11, space_after=4)
    p3_items = [
        ("Upgrade Lube Premium Mix in Lahore, Faisalabad, Multan",
         "Export Islamabad's balanced DEO/PCMO/LOW GRADE mix model. Trade marketing, mechanic engagement, consumer education. KPI: DEO share from ~17% → 25%+ by FY27."),
        ("Expand South PCMO Programme",
         "South PCMO NMgn/ltr = PKR 518 at only 0.342 ML volume — most underleveraged high-value combination in the dataset. Dedicated PCMO distributor programme in Karachi. KPI: South PCMO volume +30% next FY."),
        ("Understand North Growth Sustainability",
         "Only region growing fuel volumes but losing NMgn/ltr fastest. Determine if volume growth is structural before replicating in Central/South. KPI: NMgn/ltr stabilisation while holding volume positive."),
        ("Build Monthly City-Level Analytics",
         "Pipeline is ready. Add city × segment P&L as monthly management report. Set AI_PROVIDER=anthropic with API key for automated runs."),
    ]
    for title, desc in p3_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(f"• {title}")
        run.font.size = Pt(10); run.font.bold = True; run.font.color.rgb = NAVY
        bullet(doc, desc)
    para(doc, '', space_after=8)

    # ── Footer ───────────────────────────────────────────────────────────────
    divider(doc)
    para(doc, f"PSO OMC Analytics  |  10M FY26  |  Generated: {date.today().strftime('%d %B %Y')}  |  CONFIDENTIAL",
         color=GREY, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)


if __name__ == "__main__":
    doc = Document()
    # Remove default Normal style spacing
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    ai = json.loads(AI_JSON.read_text(encoding='utf-8'))
    build(doc, ai)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
