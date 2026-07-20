"""
PSO Retail Lubricants — Station & Volume Analysis Word Report
"""
import sys, io, os
sys.path.insert(0, 'src')
sys.stdout.reconfigure(encoding='utf-8')

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from pso import ingest
from _pso_common import INPUT_PATH, get_period_label, out_path

# ── palette ───────────────────────────────────────────────────────────────────
BLUE   = '#00479D'; DBLUE  = '#1F3864'
GREEN  = '#008C4A'; RED    = '#C00000'
ORANGE = '#E46C0A'; GREY   = '#BFBFBF'
LGREY  = '#F2F2F2'

W_BLUE   = RGBColor(0x00,0x47,0x9D)
W_DBLUE  = RGBColor(0x1F,0x38,0x64)
W_GREEN  = RGBColor(0x00,0x8C,0x4A)
W_RED    = RGBColor(0xC0,0x00,0x00)
W_WHITE  = RGBColor(0xFF,0xFF,0xFF)

# ── helpers ───────────────────────────────────────────────────────────────────
def set_bg(cell, hex_col):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{hex_col}"/>')
    tcPr.append(shd)

def ct(cell, text, bold=False, size=9.5, color=None,
       align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]; p.alignment = align
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    if level == 1:
        r.font.size = Pt(15); r.font.color.rgb = W_BLUE
    else:
        r.font.size = Pt(12); r.font.color.rgb = W_BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="00479D"/></w:pBdr>')
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)

def body(doc, text, size=10, before=3, after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    for r in p.runs: r.font.size = Pt(size)

def img_buf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=160, bbox_inches='tight')
    buf.seek(0); plt.close(fig); return buf

def chg(v): return f"+{v:.1f}%" if v >= 0 else f"{v:.1f}%"

# ── load data ─────────────────────────────────────────────────────────────────
print("Loading data…")
df, _ = ingest.load(INPUT_PATH)
REPORT_DATE = get_period_label(df)
retail = df[df['IsRetail'] & ~df['IsInternational']].copy()
lubes  = retail[retail['FuelSegment'] == 'Lubricants'].copy()

# ── aggregations ──────────────────────────────────────────────────────────────
# By category
cat = (lubes.groupby('LubeCategory')
       .agg(vol_cy=('SalesLtr_CY','sum'),
            vol_sply=('SalesLtr_SPLY','sum'),
            rev_cy=('SalesGRS_CY','sum'), mgn_cy=('NetMargin_CY','sum'),
            stns=('Customer Number','nunique'))
       .assign(vol_chg=lambda d:(d.vol_cy-d.vol_sply)/d.vol_sply.abs().replace(0,float('nan'))*100,
               vol_sh=lambda d:d.vol_cy/d.vol_cy.sum()*100,
               vol_ps=lambda d:d.vol_cy/d.stns,
               mgn_pl=lambda d:d.mgn_cy/d.vol_cy)
       .sort_values('vol_cy', ascending=False))

# By city (top 15)
city_reg = (lubes.groupby(['CityNorm','Sales office Region'])['SalesLtr_CY']
            .sum().reset_index().sort_values('SalesLtr_CY',ascending=False)
            .drop_duplicates('CityNorm').set_index('CityNorm')['Sales office Region'])

city = (lubes.groupby('CityNorm')
        .agg(vol_cy=('SalesLtr_CY','sum'),
             vol_sply=('SalesLtr_SPLY','sum'),
             mgn_cy=('NetMargin_CY','sum'), stns=('Customer Number','nunique'))
        .assign(vol_chg=lambda d:(d.vol_cy-d.vol_sply)/d.vol_sply.abs().replace(0,float('nan'))*100,
                vol_sh=lambda d:d.vol_cy/d.vol_cy.sum()*100,
                vol_ps=lambda d:d.vol_cy/d.stns,
                mgn_pl=lambda d:d.mgn_cy/d.vol_cy)
        .sort_values('vol_cy', ascending=False)
        .head(15))
city['Region'] = city.index.map(city_reg)

# ── CHART 1: Category — volume bars + stations dot ────────────────────────────
print("Generating charts…")
fig1, ax1 = plt.subplots(figsize=(9, 5))
fig1.patch.set_facecolor('white')
ax1r = ax1.twinx()

cat4 = cat[cat.index.isin(['LOW GRADE','DEO','MCO','PCMO'])].copy()
labels4 = ['Low Grade','DEO','MCO','PCMO']
x4 = np.arange(len(cat4)); bw = 0.5
bar_cols = [BLUE, GREEN, ORANGE, RED]

bars_cy = ax1.bar(x4 - 0.03, cat4['vol_cy']/1000, bw, color=bar_cols, alpha=0.88, zorder=3, label='Vol CY')
ax1.bar(x4 - 0.03, cat4['vol_sply']/1000, bw, color='none',
        edgecolor='#999999', linewidth=1.6, linestyle='--', zorder=2, label='Vol SPLY')

ax1r.plot(x4, cat4['stns'], 'o--', color='#222222', linewidth=1.6,
          markersize=8, markerfacecolor='white', markeredgewidth=2.2, zorder=5, label='Stations')
for xi, (_, r) in zip(x4, cat4.iterrows()):
    ax1r.annotate(f"{int(r.stns):,}", xy=(xi, r.stns),
                  xytext=(0, 10), textcoords='offset points',
                  ha='center', fontsize=9, fontweight='bold', color='#222222')

for bar, (_, r) in zip(bars_cy, cat4.iterrows()):
    c = GREEN if r.vol_chg >= 0 else RED
    ax1.annotate(chg(r.vol_chg),
                 xy=(bar.get_x()+bar.get_width()/2, bar.get_height()),
                 xytext=(0, 6), textcoords='offset points',
                 ha='center', fontsize=9, color=c, fontweight='bold')
    ax1.annotate(f"{r.vol_cy/1000:,.0f} KL",
                 xy=(bar.get_x()+bar.get_width()/2, bar.get_height()/2),
                 ha='center', fontsize=8.5, color='white', fontweight='bold')

ax1.set_xticks(x4); ax1.set_xticklabels(labels4, fontsize=11)
ax1.set_ylabel('Volume CY (KL)', fontsize=10, color='#333333')
ax1r.set_ylabel('No. of Stations', fontsize=10, color='#222222')
ax1.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: f'{v:,.0f}'))
ax1r.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: f'{v:,.0f}'))
ax1.set_title('Lubricant Volume & Station Count by Category', fontsize=12, fontweight='bold',
              color=BLUE, pad=10)
ax1.set_axisbelow(True); ax1.yaxis.grid(True, linestyle='--', alpha=0.4)
ax1.spines[['top','right']].set_visible(False)

h1,l1 = ax1.get_legend_handles_labels()
stn_line = plt.Line2D([0],[0],color='#222222',marker='o',linewidth=1.5,
                       markerfacecolor='white',markeredgewidth=2,label='Stations')
ax1.legend(h1+[stn_line], l1+['Stations'], fontsize=9, loc='upper right')

plt.tight_layout()
chart1_buf = img_buf(fig1)

# ── CHART 2: Top 15 cities — stations bar + KL/station line ──────────────────
fig2, ax2 = plt.subplots(figsize=(12, 5.5))
fig2.patch.set_facecolor('white')
ax2r = ax2.twinx()

city_names = [n[:20] for n in city.index.tolist()]
bar_c2 = [GREEN if c >= 0 else RED for c in city['vol_chg']]

bars2 = ax2.bar(range(15), city['stns'], 0.6, color=bar_c2, alpha=0.85, zorder=3)
for bar, n, v, vps in zip(bars2, city['stns'], city['vol_cy'], city['vol_ps']):
    ax2.annotate(str(int(n)),
                 xy=(bar.get_x()+bar.get_width()/2, bar.get_height()),
                 xytext=(0,4), textcoords='offset points',
                 ha='center', fontsize=8.5, fontweight='bold', color='#222222')

ax2r.plot(range(15), city['vol_ps']/1000, 's-', color=BLUE, linewidth=2,
          markersize=7, zorder=4, label='KL / Station')
for xi, vps in enumerate(city['vol_ps']/1000):
    offset = 9 if xi != 12 else -14   # KOHLU label goes below to avoid overlap
    ax2r.annotate(f'{vps:.1f}',
                  xy=(xi, vps), xytext=(0, offset), textcoords='offset points',
                  ha='center', fontsize=8, color=BLUE, fontweight='bold')

ax2.set_xticks(range(15))
ax2.set_xticklabels(city_names, rotation=38, ha='right', fontsize=9)
ax2.set_ylabel('No. of Stations', fontsize=10)
ax2r.set_ylabel('Volume per Station (KL)', fontsize=10, color=BLUE)
ax2r.tick_params(axis='y', labelcolor=BLUE)
ax2.set_title('Top 15 Cities — Active Lubricant Stations & Volume per Station', fontsize=12,
              fontweight='bold', color=BLUE, pad=10)
ax2.set_axisbelow(True); ax2.yaxis.grid(True, linestyle='--', alpha=0.4)
ax2.spines[['top','right']].set_visible(False)

g_patch = mpatches.Patch(color=GREEN, label='Growing city')
r_patch = mpatches.Patch(color=RED,   label='Declining city')
b_line  = plt.Line2D([0],[0], color=BLUE, marker='s', linewidth=2, label='KL / Station')
ax2.legend(handles=[g_patch, r_patch, b_line], fontsize=9, loc='upper right')

plt.tight_layout()
chart2_buf = img_buf(fig2)

# ── CHART 3: Donut — stations share per category ─────────────────────────────
fig3, (ax3a, ax3b) = plt.subplots(1, 2, figsize=(10, 4.5))
fig3.patch.set_facecolor('white')

donut_cats = cat4.copy()
pie_cols = [BLUE, GREEN, ORANGE, RED]

# Volume share donut
wedges1, _, autos1 = ax3a.pie(
    donut_cats['vol_cy'], labels=None, autopct='%1.1f%%',
    colors=pie_cols, startangle=90,
    wedgeprops=dict(width=0.55, edgecolor='white'), pctdistance=0.75)
for at in autos1: at.set_fontsize(9); at.set_color('white'); at.set_fontweight('bold')
ax3a.legend(labels4, loc='lower center', bbox_to_anchor=(0.5,-0.12), ncol=2, fontsize=9)
ax3a.set_title('Volume Mix (CY)', fontsize=11, fontweight='bold', color=BLUE, pad=8)

# Stations share donut
wedges2, _, autos2 = ax3b.pie(
    donut_cats['stns'], labels=None, autopct='%1.1f%%',
    colors=pie_cols, startangle=90,
    wedgeprops=dict(width=0.55, edgecolor='white'), pctdistance=0.75)
for at in autos2: at.set_fontsize(9); at.set_color('white'); at.set_fontweight('bold')
ax3b.legend(labels4, loc='lower center', bbox_to_anchor=(0.5,-0.12), ncol=2, fontsize=9)
ax3b.set_title('Station Count Mix', fontsize=11, fontweight='bold', color=BLUE, pad=8)

fig3.suptitle('Volume vs Station Distribution by Category', fontsize=11, color='#444444', y=1.01)
plt.tight_layout()
chart3_buf = img_buf(fig3)

# ── BUILD WORD DOC ────────────────────────────────────────────────────────────
print("Building Word document…")
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21); sec.page_height = Cm(29.7)
sec.left_margin = sec.right_margin = Cm(1.8)
sec.top_margin  = sec.bottom_margin = Cm(1.6)

# ── Title ─────────────────────────────────────────────────────────────────────
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('PSO Retail Lubricants')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = W_BLUE

p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Station Coverage & Volume Analysis by Category and City')
r2.font.size = Pt(13); r2.font.color.rgb = RGBColor(0x55,0x55,0x55)

p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(2)
r3 = p3.add_run(REPORT_DATE)
r3.bold = True; r3.font.size = Pt(11); r3.font.color.rgb = W_BLUE

doc.add_paragraph()

# KPI strip
kpi_tbl = doc.add_table(rows=2, cols=4)
kpi_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tot_vol = lubes['SalesLtr_CY'].sum()
tot_vol_sply = lubes['SalesLtr_SPLY'].sum()
tot_stns = lubes['Customer Number'].nunique()
tot_cities = lubes['CityNorm'].nunique()
kpi_heads = ['Total Volume CY','Stations Active','Cities Covered','Avg KL / Station']
kpi_vals  = [f"{tot_vol/1e6:.3f} ML", f"{tot_stns:,}", str(tot_cities),
             f"{tot_vol/tot_stns/1000:.2f} KL"]
kpi_sub   = [f"{(tot_vol-tot_vol_sply)/tot_vol_sply*100:+.1f}% vs SPLY",
             'across all regions', 'nationwide', 'per active outlet']
for ci in range(4):
    set_bg(kpi_tbl.cell(0,ci), '00479D')
    ct(kpi_tbl.cell(0,ci), kpi_heads[ci], bold=True, size=9, color=W_WHITE)
    set_bg(kpi_tbl.cell(1,ci), 'F2F2F2')
    kpi_tbl.cell(1,ci).text = ''
    vp = kpi_tbl.cell(1,ci).paragraphs[0]; vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run(kpi_vals[ci]); vr.bold=True; vr.font.size=Pt(11); vr.font.color.rgb=W_BLUE
    vp.add_run('\n')
    sr = vp.add_run(kpi_sub[ci]); sr.font.size=Pt(8); sr.font.color.rgb=RGBColor(0x70,0x70,0x70)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — BY CATEGORY
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, '1.  Volume & Station Coverage by Lubricant Category', level=1)

body(doc,
    'PSO sells five lubricant categories through its retail network. Low Grade now accounts for '
    '54% of total volume and is growing rapidly (+24.4%), while premium categories — DEO, MCO, '
    'and PCMO — are all declining. Importantly, DEO has the largest station footprint (3,576 '
    'stations) yet sells only 1.09 KL per station versus Low Grade at 2.17 KL per station, '
    'signalling lower sell-through rates for premium products across the network.', size=10)

doc.add_picture(chart1_buf, width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Category table
heading(doc, '1.1  Category Summary Table', level=2)

col_h = ['Category','Vol CY (KL)','Vol SPLY (KL)','vs SPLY','Mix %','Stations',
         'KL / Station','Margin / L (PKR)']
all_cats = cat[~cat.index.isin(['INDUSTRIAL GRADE'])].copy()
tbl1 = doc.add_table(rows=len(all_cats)+2, cols=len(col_h))
tbl1.style = 'Table Grid'; tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER

for ci, h in enumerate(col_h):
    set_bg(tbl1.cell(0,ci), '00479D')
    ct(tbl1.cell(0,ci), h, bold=True, size=9, color=W_WHITE)

for ri, (cat_name, r) in enumerate(all_cats.iterrows(), 1):
    bg = 'F2F2F2' if ri%2==0 else 'FFFFFF'
    chg_c = W_GREEN if r.vol_chg >= 0 else W_RED
    vals = [cat_name, f'{r.vol_cy/1000:,.1f}', f'{r.vol_sply/1000:,.1f}',
            chg(r.vol_chg), f'{r.vol_sh:.1f}%', f'{int(r.stns):,}',
            f'{r.vol_ps/1000:.2f}', f'{r.mgn_pl:.0f}']
    for ci, v in enumerate(vals):
        set_bg(tbl1.cell(ri,ci), bg)
        fc = chg_c if ci==3 else None
        al = WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER
        ct(tbl1.cell(ri,ci), v, size=9, color=fc, align=al)

# Total row
tot_r = len(all_cats)+1
set_bg(tbl1.cell(tot_r,0), '1F3864')
ct(tbl1.cell(tot_r,0), 'TOTAL', bold=True, size=9, color=W_WHITE,
   align=WD_ALIGN_PARAGRAPH.LEFT)
tot_vals = [f'{tot_vol/1000:,.1f}', f'{tot_vol_sply/1000:,.1f}',
            chg((tot_vol-tot_vol_sply)/tot_vol_sply*100), '100.0%',
            f'{tot_stns:,}',
            f'{tot_vol/tot_stns/1000:.2f}',
            f'{lubes["NetMargin_CY"].sum()/tot_vol:.0f}']
for ci, v in enumerate(tot_vals, 1):
    set_bg(tbl1.cell(tot_r,ci), '1F3864')
    ct(tbl1.cell(tot_r,ci), v, bold=True, size=9, color=W_WHITE)

doc.add_paragraph()

# Volume vs station donut
doc.add_picture(chart3_buf, width=Inches(5.6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

body(doc,
    'The two donuts above reveal a structural mismatch: DEO commands 28% of stations '
    'but only 29% of volume — broadly in line. However, PCMO and MCO together represent '
    '45% of stations yet generate only 16% of volume, pointing to significant under-utilisation '
    'of premium product stocking across the network. Low Grade, by contrast, punches above its '
    'station weight at 24% of stations delivering 54% of volume.',
    size=9, before=5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — BY CITY
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, '2.  Top 15 Cities — Active Stations & Productivity', level=1)

body(doc,
    'The top 15 cities by lubricant volume collectively account for the majority of national '
    'lubricant sales. This section examines how many stations are active in each city and how '
    'productively they sell (volume per station). High productivity per station indicates strong '
    'demand or superior product placement; low productivity may signal distribution gaps or '
    'competitive pressure.',
    size=10)

doc.add_picture(chart2_buf, width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

heading(doc, '2.1  Top 15 Cities — Detailed Table', level=2)

col_h2 = ['#','City','Region','Stations','Vol CY (KL)','Vol SPLY (KL)',
           'vs SPLY','Mix %','KL / Station','Mgn/L (PKR)']
tbl2 = doc.add_table(rows=len(city)+1, cols=len(col_h2))
tbl2.style = 'Table Grid'; tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

for ci, h in enumerate(col_h2):
    set_bg(tbl2.cell(0,ci), '00479D')
    ct(tbl2.cell(0,ci), h, bold=True, size=8.5, color=W_WHITE)

for ri, (city_name, r) in enumerate(city.iterrows(), 1):
    bg = 'F2F2F2' if ri%2==0 else 'FFFFFF'
    chg_c = W_GREEN if r.vol_chg >= 0 else W_RED
    vals = [str(ri), city_name, r.get('Region',''), f'{int(r.stns):,}',
            f'{r.vol_cy/1000:,.1f}', f'{r.vol_sply/1000:,.1f}',
            chg(r.vol_chg), f'{r.vol_sh:.1f}%',
            f'{r.vol_ps/1000:.2f}', f'{r.mgn_pl:.0f}']
    for ci, v in enumerate(vals):
        set_bg(tbl2.cell(ri,ci), bg)
        fc = chg_c if ci==6 else None
        al = WD_ALIGN_PARAGRAPH.LEFT if ci in [1,2] else WD_ALIGN_PARAGRAPH.CENTER
        ct(tbl2.cell(ri,ci), v, size=8.5, color=fc, align=al)

doc.add_paragraph()

# ── Key Findings boxes ────────────────────────────────────────────────────────
heading(doc, '2.2  Key Findings', level=2)

findings = [
    ('Largest Networks — Lahore & Karachi',
     f'Lahore (198 stations) and Karachi (195 stations) have the largest lubricant networks. '
     f'However, their productivity per station (4.32 KL and 6.14 KL respectively) is moderate, '
     f'suggesting room to grow volume per outlet through better stocking and merchandising. '
     f'Both cities are growing — Karachi +1.8%, Lahore +6.4%.'),
    ('Highest Productivity — KOHLU MARI BUGTI',
     f'KOHLU MARI BUGTI stands out with only 18 stations yet 12.38 KL per station — nearly double '
     f'the next highest (Islamabad at 6.40 KL/station). However, total volume is declining (-10.6%). '
     f'This market is high-value but fragile and requires priority sales attention to reverse the trend.'),
    ('Islamabad — Quality Network',
     f'Islamabad has 48 stations generating 6.40 KL/station with +2.2% growth. This reflects a '
     f'well-placed, productive network. The opportunity here is to add stations rather than push '
     f'volume per existing outlet.'),
    ('High-Growth Cities — Rahim Yar Khan, Sargodha, Hyderabad',
     f'Three cities are growing at 14–16%: Rahim Yar Khan (+15.7%), Hyderabad (+14.0%), and '
     f'Sargodha (outside top 15 at +17.9%). These markets have momentum and should receive '
     f'priority stock allocation and new dealer appointments to sustain growth.'),
    ('Declining Cities — Rawalpindi & KOHLU MARI BUGTI',
     f'Rawalpindi (-0.4%, 83 stations) and KOHLU MARI BUGTI (-10.6%, 18 stations) are the only '
     f'two top-15 cities in decline. Rawalpindi\'s decline is marginal but given its large '
     f'network, even a small per-station improvement would add significant volume. '
     f'KOHLU requires a targeted recovery plan.'),
]

for title, text in findings:
    # Finding header (indented bold)
    fp = doc.add_paragraph()
    fp.paragraph_format.space_before = Pt(6)
    fp.paragraph_format.left_indent  = Inches(0.1)
    fr = fp.add_run(f'► {title}')
    fr.bold = True; fr.font.size = Pt(10); fr.font.color.rgb = W_BLUE

    fb = doc.add_paragraph(text)
    fb.paragraph_format.left_indent  = Inches(0.25)
    fb.paragraph_format.space_before = Pt(2)
    fb.paragraph_format.space_after  = Pt(5)
    for r in fb.runs: r.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — STATION EFFICIENCY RANKING
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, '3.  City Efficiency Ranking — Volume per Station', level=1)

body(doc,
    'Volume per station (KL/station) is a measure of network productivity. '
    'The table below ranks all top-15 cities by this metric, '
    'helping identify markets where the existing network is being fully utilised '
    'versus markets where growth can come from adding new outlets.',
    size=10)

eff = city.copy().reset_index()
eff['rank_vps'] = eff['vol_ps'].rank(ascending=False).astype(int)
eff = eff.sort_values('vol_ps', ascending=False)

eff_h = ['Efficiency Rank','City','Stations','KL / Station','Total Vol CY (KL)','YoY Chg','Interpretation']
tbl3 = doc.add_table(rows=len(eff)+1, cols=len(eff_h))
tbl3.style = 'Table Grid'; tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER

for ci, h in enumerate(eff_h):
    set_bg(tbl3.cell(0,ci), '00479D')
    ct(tbl3.cell(0,ci), h, bold=True, size=8.5, color=W_WHITE)

def interpret(vps, stns, vol_chg):
    if vps > 8:
        return 'Very high productivity — consider network expansion'
    elif vps > 5:
        return 'High productivity — strong demand per outlet'
    elif vol_chg < 0:
        return 'Below average & declining — priority recovery'
    elif stns > 100:
        return 'Large network — focus on per-station uplift'
    else:
        return 'Moderate — growth through network + uplift'

for ri, (_, r) in enumerate(eff.iterrows(), 1):
    bg = 'F2F2F2' if ri%2==0 else 'FFFFFF'
    # highlight top 3 in light blue
    if ri <= 3:
        bg = 'DDEEFF'
    vps = r.vol_ps/1000
    chg_c = W_GREEN if r.vol_chg >= 0 else W_RED
    vals = [str(ri), r['CityNorm'], f'{int(r.stns):,}',
            f'{vps:.2f}', f'{r.vol_cy/1000:,.1f}',
            chg(r.vol_chg), interpret(vps, r.stns, r.vol_chg)]
    for ci, v in enumerate(vals):
        set_bg(tbl3.cell(ri,ci), bg)
        fc = chg_c if ci==5 else None
        al = WD_ALIGN_PARAGRAPH.LEFT if ci in [1,6] else WD_ALIGN_PARAGRAPH.CENTER
        bold = (ri<=3 and ci==3)
        ct(tbl3.cell(ri,ci), v, size=8.5, color=fc, align=al, bold=bold)

doc.add_paragraph()

body(doc,
    'Cities highlighted in blue (top 3 by KL/station) represent the most productive markets. '
    'These should be benchmarks for the broader network improvement programme. '
    'Cities with large station counts but moderate productivity (Lahore, Faisalabad, Multan) '
    'offer the biggest absolute volume upside if per-station productivity can be improved '
    'even marginally.',
    size=10)

# ── Save ──────────────────────────────────────────────────────────────────────
out = out_path('PSO_Lubes_Stations_Report', 'docx', df)
doc.save(out)
print(f"\nSaved: {out}")
