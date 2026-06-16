"""
PSO Retail Lubricants Business Report
Word document with tables and embedded charts.
"""
import sys, io, os
sys.path.insert(0, 'src')
sys.stdout.reconfigure(encoding='utf-8')

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.ticker import FuncFormatter
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from lxml import etree
from pso import ingest
from _pso_common import INPUT_PATH, get_period_label, out_path

# ── colours ───────────────────────────────────────────────────────────────────
PSO_BLUE  = RGBColor(0x00, 0x47, 0x9D)
PSO_GREEN = RGBColor(0x00, 0x8C, 0x4A)
PSO_RED   = RGBColor(0xC0, 0x00, 0x00)
PSO_ORANGE= RGBColor(0xE4, 0x6C, 0x0A)
LGREY     = RGBColor(0xF2, 0xF2, 0xF2)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
DARK      = RGBColor(0x1F, 0x1F, 0x1F)

PLT_BLUE  = '#00479D'
PLT_GREEN = '#008C4A'
PLT_RED   = '#C00000'
PLT_ORANGE= '#E46C0A'
PLT_LGREY = '#F2F2F2'
PLT_GREY  = '#BFBFBF'

# ── helpers ───────────────────────────────────────────────────────────────────
def kl(v):   return v / 1_000
def ml(v):   return v / 1_000_000
def bn(v):   return v / 1_000_000_000
def pct(a,b): return (a-b)/abs(b)*100 if b else 0

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = PSO_BLUE
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = PSO_BLUE
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    # bottom border line
    if level <= 2:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml('<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                         '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="00479D"/>'
                         '</w:pBdr>')
        pPr.append(pBdr)
    return p

def add_body(doc, text, size=10, space_before=3, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for run in p.runs:
        run.font.size = Pt(size)
    return p

def chg_str(v):
    return f"+{v:.1f}%" if v >= 0 else f"{v:.1f}%"

def img_buf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close(fig)
    return buf

# ── load data ─────────────────────────────────────────────────────────────────
print("Loading data…")
df, _ = ingest.load(INPUT_PATH)
REPORT_DATE = get_period_label(df)
retail = df[df['IsRetail'] & ~df['IsInternational']].copy()
lubes  = retail[retail['FuelSegment'] == 'Lubricants'].copy()

# ── aggregate helpers ─────────────────────────────────────────────────────────
def agg_vol_margin(frame, by):
    return (frame.groupby(by, as_index=False)
            .agg(vol_cy=('SalesLtr_CY','sum'), vol_ly=('SalesLtr_LY','sum'),
                 rev_cy=('SalesGRS_CY','sum'), rev_ly=('SalesGRS_LY','sum'),
                 mgn_cy=('NetMargin_CY','sum'), mgn_ly=('NetMargin_LY','sum'),
                 stns=('Customer Number','nunique'))
            .assign(vol_chg=lambda d: d.apply(lambda r: pct(r.vol_cy, r.vol_ly), axis=1),
                    rev_chg=lambda d: d.apply(lambda r: pct(r.rev_cy, r.rev_ly), axis=1),
                    mgn_pl_cy=lambda d: d.mgn_cy / d.vol_cy.replace(0, np.nan),
                    mgn_pl_ly=lambda d: d.mgn_ly / d.vol_ly.replace(0, np.nan),
                    vol_sh=lambda d: d.vol_cy / d.vol_cy.sum() * 100))

# National
vc = lubes['SalesLtr_CY'].sum()
vl = lubes['SalesLtr_LY'].sum()
rc = lubes['SalesGRS_CY'].sum()
rl = lubes['SalesGRS_LY'].sum()
mc = lubes['NetMargin_CY'].sum()
ml_ = lubes['NetMargin_LY'].sum()
n_stns = lubes['Customer Number'].nunique()
n_cities = lubes['CityNorm'].nunique()

# By category
cat_df = agg_vol_margin(lubes, 'LubeCategory').sort_values('vol_cy', ascending=False)

# By region
reg_df = agg_vol_margin(lubes, 'Sales office Region').sort_values('vol_cy', ascending=False)

# By city (all)
city_df = agg_vol_margin(lubes, 'CityNorm').sort_values('vol_cy', ascending=False)

# By station
stn_df = (lubes.groupby(['Customer Number','Name 1','CityNorm','Sales office Region'], as_index=False)
          .agg(vol_cy=('SalesLtr_CY','sum'), vol_ly=('SalesLtr_LY','sum'),
               mgn_cy=('NetMargin_CY','sum'))
          .assign(vol_chg=lambda d: d.apply(lambda r: pct(r.vol_cy, r.vol_ly), axis=1),
                  mgn_pl_cy=lambda d: d.mgn_cy / d.vol_cy.replace(0, np.nan))
          .sort_values('vol_cy', ascending=False))

# ── charts ────────────────────────────────────────────────────────────────────
print("Generating charts…")

# 1. Volume overview: grouped bar (national CY vs LY) + category donut
fig1, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 4.2))
fig1.patch.set_facecolor('white')

# Grouped bar – categories
cats = cat_df[cat_df['LubeCategory'].isin(['LOW GRADE','DEO','MCO','PCMO'])]
x = np.arange(len(cats))
bw = 0.35
bars_cy = ax1.bar(x - bw/2, kl(cats['vol_cy']), bw, color=PLT_BLUE, label='CY', zorder=3)
bars_ly = ax1.bar(x + bw/2, kl(cats['vol_ly']), bw, color=PLT_GREY, label='LY', zorder=3)
ax1.set_xticks(x)
ax1.set_xticklabels(['Low Grade','DEO','MCO','PCMO'], fontsize=9)
ax1.set_ylabel('Volume (KL)', fontsize=9)
ax1.set_title('Volume by Category — CY vs LY', fontsize=10, fontweight='bold', color=PLT_BLUE, pad=8)
ax1.legend(fontsize=8)
ax1.yaxis.set_major_formatter(FuncFormatter(lambda v,_: f'{v:,.0f}'))
ax1.set_axisbelow(True); ax1.yaxis.grid(True, linestyle='--', alpha=0.5)
ax1.spines[['top','right']].set_visible(False)
for bar, row in zip(bars_cy, cats.itertuples()):
    c = PLT_GREEN if row.vol_chg >= 0 else PLT_RED
    ax1.annotate(chg_str(row.vol_chg), xy=(bar.get_x()+bar.get_width()/2, bar.get_height()),
                 xytext=(0,4), textcoords='offset points', ha='center', fontsize=7.5, color=c, fontweight='bold')

# Donut – share
donut_data = cat_df[~cat_df['LubeCategory'].isin(['INDUSTRIAL GRADE','OTHERS'])]
colours = [PLT_BLUE, PLT_GREEN, PLT_ORANGE, PLT_RED]
wedges, texts, autotexts = ax2.pie(
    donut_data['vol_cy'], labels=None,
    autopct='%1.1f%%', startangle=90,
    colors=colours[:len(donut_data)],
    wedgeprops=dict(width=0.55, edgecolor='white'), pctdistance=0.75)
for at in autotexts:
    at.set_fontsize(8); at.set_color('white'); at.set_fontweight('bold')
ax2.legend(donut_data['LubeCategory'].tolist(), loc='lower center',
           bbox_to_anchor=(0.5, -0.12), ncol=2, fontsize=8)
ax2.set_title('Volume Mix — CY', fontsize=10, fontweight='bold', color=PLT_BLUE, pad=8)

plt.tight_layout()
chart1_buf = img_buf(fig1)

# 2. Margin/Litre by category
fig2, ax = plt.subplots(figsize=(8, 3.8))
fig2.patch.set_facecolor('white')
cats4 = cat_df[cat_df['LubeCategory'].isin(['LOW GRADE','DEO','MCO','PCMO'])].copy()
x = np.arange(len(cats4)); bw = 0.35
ax.bar(x - bw/2, cats4['mgn_pl_cy'], bw, color=PLT_BLUE, label='CY')
ax.bar(x + bw/2, cats4['mgn_pl_ly'], bw, color=PLT_GREY, label='LY')
ax.set_xticks(x)
ax.set_xticklabels(['Low Grade','DEO','MCO','PCMO'], fontsize=9)
ax.set_ylabel('Net Margin / Litre (PKR)', fontsize=9)
ax.set_title('Net Margin per Litre by Category', fontsize=10, fontweight='bold', color=PLT_BLUE, pad=8)
ax.legend(fontsize=8)
ax.yaxis.set_major_formatter(FuncFormatter(lambda v,_: f'PKR {v:,.0f}'))
ax.set_axisbelow(True); ax.yaxis.grid(True, linestyle='--', alpha=0.5)
ax.spines[['top','right']].set_visible(False)
plt.tight_layout()
chart2_buf = img_buf(fig2)

# 3. Regional comparison bar
fig3, axes = plt.subplots(1, 3, figsize=(11, 4))
fig3.patch.set_facecolor('white')
reg_order = ['Central','South','North']
reg_plot = reg_df.set_index('Sales office Region').reindex(reg_order)

# Volume
ax3a = axes[0]
bars = ax3a.bar(reg_order, kl(reg_plot['vol_cy']), color=[PLT_BLUE,PLT_GREEN,PLT_ORANGE], zorder=3)
ax3a.bar(reg_order, kl(reg_plot['vol_ly']), color='none', edgecolor=PLT_GREY, linewidth=1.5, linestyle='--', zorder=2)
ax3a.set_title('Volume CY (KL)', fontsize=10, fontweight='bold', color=PLT_BLUE)
ax3a.yaxis.set_major_formatter(FuncFormatter(lambda v,_: f'{v:,.0f}'))
ax3a.set_axisbelow(True); ax3a.yaxis.grid(True, linestyle='--', alpha=0.5)
ax3a.spines[['top','right']].set_visible(False)
for bar, reg in zip(bars, reg_order):
    chg = reg_plot.loc[reg,'vol_chg']
    c = PLT_GREEN if chg >= 0 else PLT_RED
    ax3a.annotate(chg_str(chg), xy=(bar.get_x()+bar.get_width()/2, bar.get_height()),
                  xytext=(0,4), textcoords='offset points', ha='center', fontsize=8, color=c, fontweight='bold')

# Margin/L
ax3b = axes[1]
ax3b.bar(reg_order, reg_plot['mgn_pl_cy'], color=[PLT_BLUE,PLT_GREEN,PLT_ORANGE], zorder=3)
ax3b.axhline(mc/vc, color='black', linestyle='--', linewidth=1, label=f'National avg PKR {mc/vc:.0f}')
ax3b.set_title('Net Margin / Litre (PKR)', fontsize=10, fontweight='bold', color=PLT_BLUE)
ax3b.yaxis.set_major_formatter(FuncFormatter(lambda v,_: f'{v:.0f}'))
ax3b.set_axisbelow(True); ax3b.yaxis.grid(True, linestyle='--', alpha=0.5)
ax3b.spines[['top','right']].set_visible(False)
ax3b.legend(fontsize=7)

# Stations
ax3c = axes[2]
ax3c.bar(reg_order, reg_plot['stns'], color=[PLT_BLUE,PLT_GREEN,PLT_ORANGE], zorder=3)
ax3c.set_title('Active Stations', fontsize=10, fontweight='bold', color=PLT_BLUE)
ax3c.set_axisbelow(True); ax3c.yaxis.grid(True, linestyle='--', alpha=0.5)
ax3c.spines[['top','right']].set_visible(False)
for ax_ in axes:
    ax_.set_xlabel('')

plt.tight_layout()
chart3_buf = img_buf(fig3)

# 4. Top 15 cities bar
fig4, ax4 = plt.subplots(figsize=(11, 5))
fig4.patch.set_facecolor('white')
top15 = city_df.head(15).copy()
colours4 = [PLT_GREEN if c >= 0 else PLT_RED for c in top15['vol_chg']]
bars4 = ax4.barh(range(len(top15)), kl(top15['vol_cy']), color=PLT_BLUE, zorder=3)
# LY ghost
ax4.barh(range(len(top15)), kl(top15['vol_ly']), color='none',
         edgecolor=PLT_GREY, linewidth=1.2, linestyle='--', zorder=2)
ax4.set_yticks(range(len(top15)))
ax4.set_yticklabels(top15['CityNorm'].tolist(), fontsize=9)
ax4.invert_yaxis()
ax4.set_xlabel('Volume (KL)', fontsize=9)
ax4.set_title('Top 15 Cities — Lubricants Volume (CY vs LY)', fontsize=10, fontweight='bold', color=PLT_BLUE, pad=8)
ax4.xaxis.set_major_formatter(FuncFormatter(lambda v,_: f'{v:,.0f}'))
ax4.set_axisbelow(True); ax4.xaxis.grid(True, linestyle='--', alpha=0.5)
ax4.spines[['top','right']].set_visible(False)
for bar, (_, row) in zip(bars4, top15.iterrows()):
    c = PLT_GREEN if row.vol_chg >= 0 else PLT_RED
    ax4.annotate(chg_str(row.vol_chg),
                 xy=(bar.get_width(), bar.get_y()+bar.get_height()/2),
                 xytext=(5, 0), textcoords='offset points', va='center', fontsize=8, color=c)
plt.tight_layout()
chart4_buf = img_buf(fig4)

# 5. Underperformers: cities with vol_cy > 100KL but declining
under = city_df[(city_df['vol_cy'] > 100_000) & (city_df['vol_chg'] < 0)].sort_values('vol_chg')
fig5, ax5 = plt.subplots(figsize=(9, max(3, len(under)*0.45+1)))
fig5.patch.set_facecolor('white')
ax5.barh(range(len(under)), under['vol_chg'], color=PLT_RED, zorder=3)
ax5.set_yticks(range(len(under)))
ax5.set_yticklabels(under['CityNorm'].tolist(), fontsize=9)
ax5.invert_yaxis()
ax5.axvline(0, color='black', linewidth=0.8)
ax5.set_xlabel('Volume Change %', fontsize=9)
ax5.set_title('Declining Cities (>100KL threshold)', fontsize=10, fontweight='bold', color=PLT_RED, pad=8)
ax5.xaxis.set_major_formatter(FuncFormatter(lambda v,_: f'{v:+.1f}%'))
ax5.set_axisbelow(True); ax5.xaxis.grid(True, linestyle='--', alpha=0.4)
ax5.spines[['top','right']].set_visible(False)
plt.tight_layout()
chart5_buf = img_buf(fig5)

# 6. Station scatter: volume vs margin/L (top 200 stations)
top200 = stn_df[stn_df['vol_cy'] > 0].head(200).copy()
fig6, ax6 = plt.subplots(figsize=(10, 5))
fig6.patch.set_facecolor('white')
reg_colours = {'Central': PLT_BLUE, 'South': PLT_GREEN, 'North': PLT_ORANGE}
for reg, grp in top200.groupby('Sales office Region'):
    ax6.scatter(kl(grp['vol_cy']), grp['mgn_pl_cy'],
                color=reg_colours.get(reg, 'grey'), alpha=0.7, s=40, label=reg, zorder=3)
ax6.set_xlabel('Volume CY (KL)', fontsize=9)
ax6.set_ylabel('Net Margin / Litre (PKR)', fontsize=9)
ax6.set_title('Station Landscape — Volume vs Margin (Top 200 by Volume)', fontsize=10, fontweight='bold', color=PLT_BLUE, pad=8)
ax6.legend(fontsize=8)
ax6.set_axisbelow(True); ax6.grid(True, linestyle='--', alpha=0.4)
ax6.spines[['top','right']].set_visible(False)
# median lines
ax6.axhline(mc/vc, color=PLT_GREY, linestyle='--', linewidth=1)
ax6.axvline(kl(stn_df[stn_df['vol_cy']>0]['vol_cy'].median()), color=PLT_GREY, linestyle='--', linewidth=1)
plt.tight_layout()
chart6_buf = img_buf(fig6)

# ── build Word document ───────────────────────────────────────────────────────
print("Building Word document…")
doc = Document()

# Page setup: A4, narrow margins
sec = doc.sections[0]
sec.page_width   = Cm(21)
sec.page_height  = Cm(29.7)
sec.left_margin  = sec.right_margin = Cm(1.8)
sec.top_margin   = sec.bottom_margin = Cm(1.6)

# ── COVER / TITLE ─────────────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(40)
r = cover.add_run('PSO Retail Lubricants Business')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = PSO_BLUE

cover2 = doc.add_paragraph()
cover2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = cover2.add_run('Performance Report — Volume, Margin & Growth Analysis')
r2.font.size = Pt(13); r2.font.color.rgb = RGBColor(0x60,0x60,0x60)

cover3 = doc.add_paragraph()
cover3.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover3.paragraph_format.space_before = Pt(4)
r3 = cover3.add_run(REPORT_DATE)
r3.font.size = Pt(11); r3.font.color.rgb = PSO_BLUE; r3.bold = True

doc.add_paragraph()

# KPI summary bar (4-cell table)
kpi_tbl = doc.add_table(rows=2, cols=4)
kpi_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
kpi_headers = ['Total Volume CY','Revenue CY','Net Margin CY','Margin / Litre']
kpi_values  = [f'{ml(vc):.3f} ML', f'PKR {bn(rc):.2f} Bn', f'PKR {bn(mc):.2f} Bn', f'PKR {mc/vc:.0f}']
kpi_changes = [chg_str(pct(vc,vl)), chg_str(pct(rc,rl)), chg_str(pct(mc,ml_)), chg_str(pct(mc/vc, ml_/vl))]
kpi_chg_col = [PSO_GREEN if pct(vc,vl)>=0 else PSO_RED,
               PSO_GREEN if pct(rc,rl)>=0 else PSO_RED,
               PSO_GREEN if pct(mc,ml_)>=0 else PSO_RED,
               PSO_GREEN if pct(mc/vc,ml_/vl)>=0 else PSO_RED]

for ci in range(4):
    h_cell = kpi_tbl.cell(0, ci)
    v_cell = kpi_tbl.cell(1, ci)
    set_cell_bg(h_cell, '00479D')
    cell_text(h_cell, kpi_headers[ci], bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(v_cell, 'F2F2F2')
    v_cell.text = ''
    vp = v_cell.paragraphs[0]; vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr1 = vp.add_run(kpi_values[ci]+' ')
    vr1.bold = True; vr1.font.size = Pt(11); vr1.font.color.rgb = PSO_BLUE
    vr2 = vp.add_run(kpi_changes[ci])
    vr2.bold = True; vr2.font.size = Pt(9); vr2.font.color.rgb = kpi_chg_col[ci]

doc.add_paragraph()
add_body(doc, f'This report covers lubricant sales within PSO\'s Retail Business segment for {REPORT_DATE}. '
              f'The network comprises {n_stns:,} active stations across {n_cities} cities nationally. '
              f'Unless otherwise noted, volumes are in kilolitres (KL) and all financials are in PKR.',
         size=9, space_after=2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1: TOTAL VOLUME OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1.  Total Lubricants Volume — National Overview', level=1)

add_body(doc,
    f'PSO\'s retail lubricant network sold {ml(vc):.3f} million litres (ML) in CY, compared to '
    f'{ml(vl):.3f} ML in the prior year — a growth of {pct(vc,vl):.1f}%. This volume growth '
    f'outpaces the overall retail portfolio, reflecting the continued penetration of PSO-branded '
    f'lubricants at petrol stations. Revenue grew by {pct(rc,rl):.1f}% to PKR {bn(rc):.2f} Bn, '
    f'partly driven by higher realised prices. However, net margin per litre compressed from '
    f'PKR {ml_/vl:.0f} to PKR {mc/vc:.0f} ({pct(mc/vc, ml_/vl):.1f}%), primarily because '
    f'the fastest-growing segment — Low Grade lubricants — carries significantly lower margins '
    f'than premium categories (DEO, PCMO).', size=10)

doc.add_picture(chart1_buf, width=Inches(6.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Category table
add_heading(doc, '1.1  Volume & Margin by Lubricant Category', level=2)

add_body(doc,
    'Low Grade lubricants now account for more than half of total volume (54.2%), growing at '
    '+24.4% YoY. While this drives headline volume growth, it dilutes the overall margin per litre '
    'because Low Grade earns only PKR 116/litre versus PKR 501/litre for PCMO. '
    'Premium categories (DEO, MCO, PCMO) collectively declined by ~6–11%, suggesting a potential '
    'shift in the consumer mix toward cheaper alternatives or a gap in premium product availability.', size=10)

# Table
hdr_row = ['Category', 'Vol CY (KL)', 'Vol LY (KL)', 'YoY Chg', 'Mix %', 'Mgn/L CY (PKR)', 'Mgn/L LY (PKR)']
tbl1 = doc.add_table(rows=len(cat_df)+2, cols=len(hdr_row))
tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl1.style = 'Table Grid'

for ci, h in enumerate(hdr_row):
    set_cell_bg(tbl1.cell(0, ci), '00479D')
    cell_text(tbl1.cell(0, ci), h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

for ri, (_, row) in enumerate(cat_df.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    vals = [row['LubeCategory'], f'{kl(row.vol_cy):,.1f}', f'{kl(row.vol_ly):,.1f}',
            chg_str(row.vol_chg), f'{row.vol_sh:.1f}%',
            f'{row.mgn_pl_cy:.0f}' if not np.isnan(row.mgn_pl_cy) else 'N/A',
            f'{row.mgn_pl_ly:.0f}' if not np.isnan(row.mgn_pl_ly) else 'N/A']
    for ci, v in enumerate(vals):
        set_cell_bg(tbl1.cell(ri, ci), bg)
        chg_clr = None
        if ci == 3:
            chg_clr = PSO_GREEN if row.vol_chg >= 0 else PSO_RED
        cell_text(tbl1.cell(ri, ci), v, size=9,
                  color=chg_clr, align=WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT)

# Total row
tot_row = len(cat_df) + 1
set_cell_bg(tbl1.cell(tot_row, 0), '1F3864')
cell_text(tbl1.cell(tot_row, 0), 'TOTAL', bold=True, size=9, color=WHITE)
for ci, v in enumerate([f'{kl(vc):,.1f}', f'{kl(vl):,.1f}', chg_str(pct(vc,vl)), '100.0%',
                         f'{mc/vc:.0f}', f'{ml_/vl:.0f}'], 1):
    set_cell_bg(tbl1.cell(tot_row, ci), '1F3864')
    chg_clr = WHITE
    cell_text(tbl1.cell(tot_row, ci), v, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

doc.add_picture(chart2_buf, width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_body(doc,
    'Key observation: The margin compression evident at the national level is structurally linked '
    'to category mix. If Low Grade continues to outgrow premium categories at the current rate, '
    'overall margin per litre will continue to erode even as absolute margin PKR values remain '
    'relatively stable.', size=9, space_before=6)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2: REGIONAL ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2.  Regional Analysis', level=1)

add_body(doc,
    f'PSO\'s lubricant sales are organised across three sales office regions: Central, South, and North. '
    f'Central dominates with {reg_df.set_index("Sales office Region").loc["Central","vol_sh"]:.1f}% '
    f'of national volume and the highest absolute growth ({chg_str(reg_df.set_index("Sales office Region").loc["Central","vol_chg"])} YoY), '
    f'driven by the dense urban markets of Lahore, Faisalabad, Multan, and Gujranwala. '
    f'South and North contribute approximately 27% and 26% respectively, both posting moderate '
    f'growth of around +5%.', size=10)

doc.add_picture(chart3_buf, width=Inches(6.4))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Regional table
add_heading(doc, '2.1  Region-wise Summary', level=2)

r_hdr = ['Region', 'Vol CY (KL)', 'Vol LY (KL)', 'YoY Chg', 'Mix %', 'Stations', 'Mgn/L CY', 'Mgn/L LY']
tbl2 = doc.add_table(rows=len(reg_df)+2, cols=len(r_hdr))
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = 'Table Grid'

for ci, h in enumerate(r_hdr):
    set_cell_bg(tbl2.cell(0, ci), '00479D')
    cell_text(tbl2.cell(0, ci), h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

for ri, (_, row) in enumerate(reg_df.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    vals = [row['Sales office Region'], f'{kl(row.vol_cy):,.1f}', f'{kl(row.vol_ly):,.1f}',
            chg_str(row.vol_chg), f'{row.vol_sh:.1f}%', str(int(row.stns)),
            f'PKR {row.mgn_pl_cy:.0f}', f'PKR {row.mgn_pl_ly:.0f}']
    for ci, v in enumerate(vals):
        set_cell_bg(tbl2.cell(ri, ci), bg)
        chg_clr = None
        if ci == 3:
            chg_clr = PSO_GREEN if row.vol_chg >= 0 else PSO_RED
        cell_text(tbl2.cell(ri, ci), v, size=9, color=chg_clr,
                  align=WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT)

tot_r = len(reg_df) + 1
set_cell_bg(tbl2.cell(tot_r, 0), '1F3864')
cell_text(tbl2.cell(tot_r, 0), 'NATIONAL', bold=True, size=9, color=WHITE)
for ci, v in enumerate([f'{kl(vc):,.1f}', f'{kl(vl):,.1f}', chg_str(pct(vc,vl)), '100.0%',
                         str(n_stns), f'PKR {mc/vc:.0f}', f'PKR {ml_/vl:.0f}'], 1):
    set_cell_bg(tbl2.cell(tot_r, ci), '1F3864')
    cell_text(tbl2.cell(tot_r, ci), v, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

add_body(doc,
    'North Region has the highest margin per litre (PKR {:.0f}) despite the smallest volume share, '
    'suggesting a healthier product mix weighted toward premium categories. Central\'s lower margin '
    'per litre (PKR {:.0f}) despite the largest volume is consistent with its higher Low Grade '
    'penetration. South sits in between at PKR {:.0f}/litre.'.format(
        reg_df.set_index('Sales office Region').loc['North','mgn_pl_cy'],
        reg_df.set_index('Sales office Region').loc['Central','mgn_pl_cy'],
        reg_df.set_index('Sales office Region').loc['South','mgn_pl_cy']), size=10)

# Category x Region breakdown
add_heading(doc, '2.2  Category Mix by Region', level=2)

add_body(doc, 'The table below shows volume by lubricant category for each region, '
              'highlighting where premium categories remain strong or are declining.', size=10)

cat_reg = (lubes.groupby(['Sales office Region','LubeCategory'])
           .agg(vol_cy=('SalesLtr_CY','sum'))
           .reset_index()
           .pivot(index='Sales office Region', columns='LubeCategory', values='vol_cy')
           .fillna(0))
main_cats = ['LOW GRADE','DEO','MCO','PCMO']
cat_reg = cat_reg[[c for c in main_cats if c in cat_reg.columns]]

cr_hdr = ['Region'] + [c.title() for c in cat_reg.columns] + ['Total']
tbl3 = doc.add_table(rows=len(cat_reg)+2, cols=len(cr_hdr))
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl3.style = 'Table Grid'
for ci, h in enumerate(cr_hdr):
    set_cell_bg(tbl3.cell(0, ci), '00479D')
    cell_text(tbl3.cell(0, ci), h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

for ri, (reg_name, row) in enumerate(cat_reg.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    cell_text(tbl3.cell(ri, 0), reg_name, size=9)
    set_cell_bg(tbl3.cell(ri, 0), bg)
    for ci, col in enumerate(cat_reg.columns, 1):
        set_cell_bg(tbl3.cell(ri, ci), bg)
        cell_text(tbl3.cell(ri, ci), f'{kl(row[col]):,.1f}', size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(tbl3.cell(ri, len(cat_reg.columns)+1), bg)
    cell_text(tbl3.cell(ri, len(cat_reg.columns)+1), f'{kl(row.sum()):,.1f}',
              bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

tot_r3 = len(cat_reg) + 1
set_cell_bg(tbl3.cell(tot_r3, 0), '1F3864')
cell_text(tbl3.cell(tot_r3, 0), 'TOTAL', bold=True, size=9, color=WHITE)
for ci, col in enumerate(cat_reg.columns, 1):
    set_cell_bg(tbl3.cell(tot_r3, ci), '1F3864')
    cell_text(tbl3.cell(tot_r3, ci), f'{kl(cat_reg[col].sum()):,.1f}', bold=True, size=9, color=WHITE,
              align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(tbl3.cell(tot_r3, len(cat_reg.columns)+1), '1F3864')
cell_text(tbl3.cell(tot_r3, len(cat_reg.columns)+1), f'{kl(cat_reg.values.sum()):,.1f}',
          bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3: CITY-LEVEL ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3.  City-Level Performance', level=1)

add_body(doc,
    f'Lubricants are sold across {n_cities} cities. The analysis below identifies the top performers, '
    f'high-growth cities, and markets where volumes are declining.', size=10)

# 3.1 Top 15
add_heading(doc, '3.1  Top 15 Cities by Volume', level=2)
add_body(doc,
    'Karachi leads nationally with 1,196 KL, followed by Lahore (856 KL) and Faisalabad (435 KL). '
    'Together, the top 5 cities contribute approximately '
    f'{city_df.head(5)["vol_cy"].sum()/vc*100:.0f}% of national lubricant volume. '
    'Notable high-growth cities among the top 20 include Sargodha (+17.9%), Rahim Yar Khan (+15.7%), '
    'and Hyderabad (+14.0%).', size=10)

doc.add_picture(chart4_buf, width=Inches(6.4))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Top 15 city table
c_hdr = ['#','City','Region','Vol CY (KL)','Vol LY (KL)','YoY Chg','Stns','Mgn/L (PKR)']
top15_full = city_df.head(15).copy()
# Need region for each city
city_reg = (lubes.groupby(['CityNorm','Sales office Region'])['SalesLtr_CY']
            .sum().reset_index()
            .sort_values('SalesLtr_CY', ascending=False)
            .drop_duplicates('CityNorm')
            .set_index('CityNorm')['Sales office Region'])

tbl4 = doc.add_table(rows=len(top15_full)+1, cols=len(c_hdr))
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl4.style = 'Table Grid'

for ci, h in enumerate(c_hdr):
    set_cell_bg(tbl4.cell(0, ci), '00479D')
    cell_text(tbl4.cell(0, ci), h, bold=True, size=8.5, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

for ri, (_, row) in enumerate(top15_full.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    reg_name = city_reg.get(row['CityNorm'], '')
    chg_c = PSO_GREEN if row.vol_chg >= 0 else PSO_RED
    vals_data = [(str(ri), None), (row['CityNorm'], None), (reg_name, None),
                 (f'{kl(row.vol_cy):,.1f}', None), (f'{kl(row.vol_ly):,.1f}', None),
                 (chg_str(row.vol_chg), chg_c), (str(int(row.stns)), None),
                 (f'{row.mgn_pl_cy:.0f}' if not np.isnan(row.mgn_pl_cy) else 'N/A', None)]
    for ci, (v, fc) in enumerate(vals_data):
        set_cell_bg(tbl4.cell(ri, ci), bg)
        cell_text(tbl4.cell(ri, ci), v, size=8.5, color=fc,
                  align=WD_ALIGN_PARAGRAPH.CENTER if ci not in [1,2] else WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()

# 3.2 Underperforming cities
add_heading(doc, '3.2  Underperforming Cities (Declining Volume, >100 KL Threshold)', level=2)

under_cities = city_df[(city_df['vol_cy'] > 100_000) & (city_df['vol_chg'] < 0)].sort_values('vol_chg').copy()

add_body(doc,
    f'{len(under_cities)} cities with meaningful volume (>100 KL) recorded year-on-year declines. '
    'KOHLU MARI BUGTI is the largest by absolute volume but shows a significant -10.6% decline, '
    'which may reflect market-specific factors (competition, access, product availability). '
    'Other declining markets include Rawalpindi (-0.4%), Kasur (-2.3%), and D.I. Khan (if applicable). '
    'These require targeted sales attention to arrest volume loss.', size=10)

doc.add_picture(chart5_buf, width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

if len(under_cities) > 0:
    uc_hdr = ['City','Region','Vol CY (KL)','Vol LY (KL)','YoY Chg','Mgn/L (PKR)']
    tbl5 = doc.add_table(rows=min(len(under_cities),20)+1, cols=len(uc_hdr))
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl5.style = 'Table Grid'
    for ci, h in enumerate(uc_hdr):
        set_cell_bg(tbl5.cell(0, ci), 'C00000')
        cell_text(tbl5.cell(0, ci), h, bold=True, size=8.5, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, (_, row) in enumerate(under_cities.head(20).iterrows(), 1):
        bg = 'FFEDED' if ri % 2 == 1 else 'FFF4F4'
        reg_name = city_reg.get(row['CityNorm'], '')
        vals_data = [row['CityNorm'], reg_name, f'{kl(row.vol_cy):,.1f}',
                     f'{kl(row.vol_ly):,.1f}', chg_str(row.vol_chg),
                     f'{row.mgn_pl_cy:.0f}' if not np.isnan(row.mgn_pl_cy) else 'N/A']
        for ci, v in enumerate(vals_data):
            set_cell_bg(tbl5.cell(ri, ci), bg)
            chg_clr = PSO_RED if ci == 4 else None
            cell_text(tbl5.cell(ri, ci), v, size=8.5, color=chg_clr,
                      align=WD_ALIGN_PARAGRAPH.CENTER if ci not in [0,1] else WD_ALIGN_PARAGRAPH.LEFT)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4: STATION-LEVEL SEGMENTATION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4.  Station-Level Segmentation', level=1)

add_body(doc,
    f'The lubricant network has {n_stns:,} active stations. Station performance varies widely — '
    f'the top station (KAMAL PET SERVICE, Kohlu) sold 38.6 KL while the median station sells '
    f'{kl(stn_df[stn_df.vol_cy>0]["vol_cy"].median()):.1f} KL. '
    f'The chart below shows the volume vs margin landscape for the top 200 stations, '
    f'colour-coded by region.', size=10)

doc.add_picture(chart6_buf, width=Inches(6.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Top 25 stations table
add_heading(doc, '4.1  Top 25 Stations by Volume', level=2)

top25 = stn_df.head(25).copy()
s_hdr = ['#','Station Name','City','Region','Vol CY (KL)','Vol LY (KL)','YoY Chg','Mgn/L (PKR)']
tbl6 = doc.add_table(rows=len(top25)+1, cols=len(s_hdr))
tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl6.style = 'Table Grid'

for ci, h in enumerate(s_hdr):
    set_cell_bg(tbl6.cell(0, ci), '00479D')
    cell_text(tbl6.cell(0, ci), h, bold=True, size=8, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

for ri, (_, row) in enumerate(top25.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    chg_clr = PSO_GREEN if row.vol_chg >= 0 else PSO_RED
    vals_d = [(str(ri),None), (row['Name 1'][:32],None), (row['CityNorm'],None),
              (row['Sales office Region'],None), (f'{kl(row.vol_cy):.1f}',None),
              (f'{kl(row.vol_ly):.1f}',None), (chg_str(row.vol_chg), chg_clr),
              (f'{row.mgn_pl_cy:.0f}' if not np.isnan(row.mgn_pl_cy) else 'N/A', None)]
    for ci, (v, fc) in enumerate(vals_d):
        set_cell_bg(tbl6.cell(ri, ci), bg)
        cell_text(tbl6.cell(ri, ci), v, size=8, color=fc,
                  align=WD_ALIGN_PARAGRAPH.CENTER if ci not in [1,2,3] else WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()

# Station volume distribution
add_heading(doc, '4.2  Station Volume Distribution', level=2)

vol_bins = [0,1,5,10,20,50,float('inf')]
bin_labels = ['<1 KL','1–5 KL','5–10 KL','10–20 KL','20–50 KL','>50 KL']
stn_active = stn_df[stn_df['vol_cy'] > 0].copy()
stn_active['vol_bin'] = pd.cut(kl(stn_active['vol_cy']), bins=vol_bins, labels=bin_labels, right=False)
dist_tbl = (stn_active.groupby('vol_bin', observed=True)
            .agg(stns=('Customer Number','count'), vol=('vol_cy','sum'))
            .assign(vol_sh=lambda d: d.vol/d.vol.sum()*100,
                    stn_sh=lambda d: d.stns/d.stns.sum()*100)
            .reset_index())

add_body(doc,
    'The majority of stations sell very low lubricant volumes. The right tail (>20 KL stations) '
    'is a small fraction of the network but contributes disproportionately to total sales — '
    'a classic long-tail distribution.', size=10)

dist_hdr = ['Volume Bracket','Stations','% of Stations','Volume CY (KL)','% of Total Volume']
tbl7 = doc.add_table(rows=len(dist_tbl)+1, cols=5)
tbl7.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl7.style = 'Table Grid'
for ci, h in enumerate(dist_hdr):
    set_cell_bg(tbl7.cell(0, ci), '00479D')
    cell_text(tbl7.cell(0, ci), h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
for ri, (_, row) in enumerate(dist_tbl.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    vals = [str(row['vol_bin']), str(int(row.stns)), f'{row.stn_sh:.1f}%',
            f'{kl(row.vol):,.1f}', f'{row.vol_sh:.1f}%']
    for ci, v in enumerate(vals):
        set_cell_bg(tbl7.cell(ri, ci), bg)
        cell_text(tbl7.cell(ri, ci), v, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# Bottom 20 stations (with meaningful LY but declining)
add_heading(doc, '4.3  Stations with Significant Decline (LY > 2 KL)', level=2)

bot_stns = (stn_df[(kl(stn_df['vol_ly']) > 2) & (stn_df['vol_chg'] < -20)]
            .sort_values('vol_chg').head(20).copy())

add_body(doc,
    f'{len(bot_stns)} stations that previously sold meaningful volumes (>2 KL LY) have '
    'declined by more than 20%. These represent the highest-risk accounts and should be '
    'prioritised for sales intervention.', size=10)

if len(bot_stns):
    b_hdr = ['Station Name','City','Region','Vol CY (KL)','Vol LY (KL)','YoY Chg','Mgn/L (PKR)']
    tbl8 = doc.add_table(rows=len(bot_stns)+1, cols=len(b_hdr))
    tbl8.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl8.style = 'Table Grid'
    for ci, h in enumerate(b_hdr):
        set_cell_bg(tbl8.cell(0, ci), 'C00000')
        cell_text(tbl8.cell(0, ci), h, bold=True, size=8.5, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, (_, row) in enumerate(bot_stns.iterrows(), 1):
        bg = 'FFEDED' if ri % 2 == 1 else 'FFF4F4'
        vals_d = [(row['Name 1'][:32],None), (row['CityNorm'],None), (row['Sales office Region'],None),
                  (f'{kl(row.vol_cy):.2f}',None), (f'{kl(row.vol_ly):.2f}',None),
                  (chg_str(row.vol_chg), PSO_RED),
                  (f'{row.mgn_pl_cy:.0f}' if not np.isnan(row.mgn_pl_cy) else 'N/A', None)]
        for ci, (v, fc) in enumerate(vals_d):
            set_cell_bg(tbl8.cell(ri, ci), bg)
            cell_text(tbl8.cell(ri, ci), v, size=8.5, color=fc,
                      align=WD_ALIGN_PARAGRAPH.CENTER if ci not in [0,1,2] else WD_ALIGN_PARAGRAPH.LEFT)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5: MARGIN DEEP-DIVE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5.  Margin Analysis', level=1)

add_body(doc,
    f'While total net margin grew by {pct(mc,ml_):.1f}% to PKR {bn(mc):.3f} Bn, '
    f'the per-litre net margin declined from PKR {ml_/vl:.0f} to PKR {mc/vc:.0f} — a compression '
    f'of PKR {ml_/vl - mc/vc:.0f}/litre ({pct(mc/vc, ml_/vl):.1f}%). '
    'The compression is driven by: '
    '(1) Low Grade lubricants growing 24.4% with only PKR 116/L margin; '
    '(2) High-margin PCMO (PKR 501/L) declining 4.7%; '
    '(3) DEO, the second-largest category, declining 6.7%. '
    'Reversing category mix shift toward premium products is the key lever for margin recovery.', size=10)

doc.add_picture(chart2_buf, width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

add_heading(doc, '5.1  Regional Margin Comparison', level=2)
add_body(doc,
    'North Region has the best margin profile (PKR {:.0f}/L), benefiting from a higher '
    'premium product mix. Central\'s margin (PKR {:.0f}/L) reflects its heavy Low Grade exposure. '
    'Both Central and South showed margin compression vs LY, while North\'s trajectory '
    'should be monitored closely.'.format(
        reg_df.set_index('Sales office Region').loc['North','mgn_pl_cy'],
        reg_df.set_index('Sales office Region').loc['Central','mgn_pl_cy']), size=10)

# Margin table (region x category)
mgn_cat_reg = (lubes.groupby(['Sales office Region','LubeCategory'])
               .agg(vol_cy=('SalesLtr_CY','sum'), mgn_cy=('NetMargin_CY','sum'))
               .reset_index()
               .assign(mgn_pl=lambda d: d.mgn_cy/d.vol_cy.replace(0,np.nan))
               .pivot(index='Sales office Region', columns='LubeCategory', values='mgn_pl')
               .reindex(columns=main_cats).fillna(0))

m_hdr = ['Region'] + [c.title() + ' (PKR/L)' for c in mgn_cat_reg.columns]
tbl9 = doc.add_table(rows=len(mgn_cat_reg)+1, cols=len(m_hdr))
tbl9.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl9.style = 'Table Grid'
for ci, h in enumerate(m_hdr):
    set_cell_bg(tbl9.cell(0, ci), '00479D')
    cell_text(tbl9.cell(0, ci), h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
for ri, (reg_name, row) in enumerate(mgn_cat_reg.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    cell_text(tbl9.cell(ri, 0), reg_name, size=9)
    set_cell_bg(tbl9.cell(ri, 0), bg)
    for ci, col in enumerate(mgn_cat_reg.columns, 1):
        set_cell_bg(tbl9.cell(ri, ci), bg)
        v = row[col]
        cell_text(tbl9.cell(ri, ci), f'{v:.0f}' if v else '—', size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6: SALES GROWTH IMPLICATIONS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6.  Sales Growth Implications', level=1)

add_body(doc,
    'Based on the data-driven analysis above, the following strategic implications and action '
    'priorities are recommended for the Retail Lubricants team.', size=10)

implications = [
    ('6.1  Defend & Grow Premium Category Mix',
     'DEO and PCMO together generate PKR 400+ per litre in net margin yet are both declining (DEO: -6.7%, PCMO: -4.7%). '
     'Priority actions: (a) Identify stations that historically sold DEO/PCMO but have shifted to Low Grade — '
     'conduct targeted product availability audits. (b) Sales staff at high-volume Central region stations '
     'should be incentivised on premium product revenue, not just volume. (c) Customer-facing promotions '
     'for engine oil grades at petrol stations (e.g., display boards, mechanic tie-ups) can shift mix.'),

    ('6.2  Accelerate Growth in High-Potential Cities',
     'Sargodha (+17.9%), Rahim Yar Khan (+15.7%), and Hyderabad (+14.0%) are growing rapidly and have '
     'demonstrated market receptiveness. These cities should receive: (a) priority stock availability to '
     'prevent out-of-stock situations that cap growth. (b) Additional station activations (new dealer '
     'appointments). (c) Dedicated sales officer visits to maintain momentum and drive premium up-sell.'),

    ('6.3  Turnaround Plan for Declining Cities',
     f'The {len(under_cities)} cities with >100 KL volume but declining trends represent PKR '
     f'{bn(under_cities["vol_ly"].sum()-under_cities["vol_cy"].sum())*1000:.0f} Mn in lost volume '
     'versus the prior year. A structured recovery programme should include: (a) Root-cause analysis '
     'per city — competitor activity, product mix issue, or distribution gap. (b) City-specific sales '
     'action plans with monthly tracking. KOHLU MARI BUGTI (-10.6%) needs particular attention given '
     'its large volume base and concentration of high-volume stations.'),

    ('6.4  Station-Level Performance Management',
     f'{len(bot_stns)} stations with meaningful prior-year sales have declined >20%. These accounts '
     'should be flagged for sales officer review immediately. A station that sold 5 KL LY but only '
     '2 KL CY may have shifted to a competitor brand — early intervention with product offers, '
     'credit terms, or visibility support can recover the account. High-volume stations (>20 KL) '
     'generating above-average margin per litre should be used as benchmark models for the network.'),

    ('6.5  Central Region Premium Upsell Programme',
     'Central Region has the highest volume (47.6%) but the lowest margin per litre (PKR 198/L vs '
     'National PKR 226/L). Even a PKR 10/L margin improvement across 6,329 KL would add PKR 63 Mn '
     'to net margin. The lever is premium category penetration — specifically, increasing PCMO sales '
     'at Central region stations through mechanic recommendation programmes and consumer education.'),

    ('6.6  North Region Capitalise on Margin Leadership',
     'North Region generates the highest margin per litre (PKR 261/L) but trails in volume '
     '(25.5% share). The priority here is volume growth — adding new dealer points in '
     'under-penetrated cities, leveraging North\'s stronger premium mix as a competitive '
     'differentiator. If North can grow volume by 10%, the high margin per litre means '
     f'incremental margin of approximately PKR {3396.1*10*261/1e6:.0f} Mn.'),
]

for title, body in implications:
    add_heading(doc, title, level=2)
    add_body(doc, body, size=10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# APPENDIX
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Appendix A — All Significant Cities (Top 50)', level=1)
add_body(doc, 'Cities ranked by CY volume. Includes all cities with >50 KL CY volume.', size=9)

top50 = city_df[kl(city_df['vol_cy']) > 50].copy()
app_hdr = ['#','City','Vol CY (KL)','Vol LY (KL)','YoY Chg','Stns','Mgn/L (PKR)']
tbl_app = doc.add_table(rows=len(top50)+1, cols=len(app_hdr))
tbl_app.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_app.style = 'Table Grid'
for ci, h in enumerate(app_hdr):
    set_cell_bg(tbl_app.cell(0, ci), '00479D')
    cell_text(tbl_app.cell(0, ci), h, bold=True, size=8, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
for ri, (_, row) in enumerate(top50.iterrows(), 1):
    bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    chg_clr = PSO_GREEN if row.vol_chg >= 0 else PSO_RED
    vals = [(str(ri),None),(row['CityNorm'],None),(f'{kl(row.vol_cy):,.1f}',None),
            (f'{kl(row.vol_ly):,.1f}',None),(chg_str(row.vol_chg),chg_clr),
            (str(int(row.stns)),None),
            (f'{row.mgn_pl_cy:.0f}' if not np.isnan(row.mgn_pl_cy) else 'N/A', None)]
    for ci, (v, fc) in enumerate(vals):
        set_cell_bg(tbl_app.cell(ri, ci), bg)
        cell_text(tbl_app.cell(ri, ci), v, size=8, color=fc,
                  align=WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT)

# ── save ──────────────────────────────────────────────────────────────────────
out = out_path('PSO_Lubricants_Report', 'docx', df)
doc.save(out)
print(f"\nReport saved: {out}")
print(f"Sections: Title + 6 sections + Appendix A")
