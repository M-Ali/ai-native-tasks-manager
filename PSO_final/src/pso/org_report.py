"""Org Report — lightweight single-Sales-Org summary.

For every Sales Org other than Retail Business (Aviation, LPG, Chemicals,
Agency, Strategic Accounts, Key Accounts, Head Office, Power Projects,
Marine — see config.SALES_ORG_CATEGORIES) there is no further category split
worth a full deep-dive. This module produces one small Excel overview
workbook + one short Word summary per Org instead.
"""

from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from openpyxl import Workbook

from pso.config import COL_ORG, COL_PRODUCT, COL_REGION
from pso.analyze import _agg, _vol_share, _pct_chg
from pso.excel_report import _write_header, _write_section_title, _write_df, PCT_CHG_COLS

PSO_BLUE = RGBColor(0x00, 0x47, 0x9D)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
LGREY    = RGBColor(0xF2, 0xF2, 0xF2)


def build(df: pd.DataFrame, org: str, period: str, out_dir: str | Path) -> tuple[Path, Path]:
    """Build the Excel overview + Word summary for a single Sales Org.
    Returns (excel_path, docx_path)."""
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    sub = df[df[COL_ORG] == org].copy()

    overview = _agg(sub, [COL_ORG])

    n_products = sub[COL_PRODUCT].nunique()
    by_product = _agg(sub, [COL_PRODUCT]).pipe(_vol_share) if n_products > 1 else pd.DataFrame()

    n_regions = sub[sub[COL_REGION] != ""][COL_REGION].nunique()
    by_region = _agg(sub, [COL_REGION]).pipe(_vol_share) if n_regions > 1 else pd.DataFrame()

    excel_path = _build_excel(overview, by_product, by_region, org, period, out_dir)
    docx_path  = _build_docx(overview, by_product, by_region, org, period, out_dir)
    return excel_path, docx_path


def _safe_name(org: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", org).strip("_")


# ── Excel ─────────────────────────────────────────────────────────────────────

def _build_excel(overview, by_product, by_region, org, period, out_dir) -> Path:
    wb = Workbook()
    wb.remove(wb.active)
    ws = wb.create_sheet("00_Overview")
    ws.freeze_panes = "A3"
    row = _write_header(ws, f"{org} — Overview (CY vs SPLY)", period)

    row = _write_section_title(ws, "Totals", row)
    row = _write_df(ws, overview, row, pct_cols=PCT_CHG_COLS, table_name="tbl_overview")

    if not by_product.empty:
        row = _write_section_title(ws, "By Product", row + 1)
        row = _write_df(ws, by_product, row, pct_cols=PCT_CHG_COLS, table_name="tbl_by_product")

    if not by_region.empty:
        row = _write_section_title(ws, "By Region", row + 1)
        _write_df(ws, by_region, row, pct_cols=PCT_CHG_COLS, table_name="tbl_by_region")

    path = out_dir / f"PSO_{_safe_name(org)}_Overview_{period}.xlsx"
    wb.save(path)
    return path


# ── Word ──────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{hex_color}"/>'))


def _cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color


def _chg_str(v):
    if pd.isna(v):
        return "—"
    return f"+{v:.1f}%" if v >= 0 else f"{v:.1f}%"


def _write_table(doc, df: pd.DataFrame, headers: list[str], cols: list[str], fmt: dict):
    tbl = doc.add_table(rows=len(df) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(headers):
        _cell_text(tbl.cell(0, ci), h, bold=True, size=8.5, color=WHITE)
        _set_cell_bg(tbl.cell(0, ci), "00479D")
    for ri, (_, row) in enumerate(df.iterrows(), 1):
        bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, col in enumerate(cols):
            val = row[col]
            text = fmt.get(col, str)(val) if col in fmt else (f"{val:.2f}" if isinstance(val, float) else str(val))
            _cell_text(tbl.cell(ri, ci), text, size=8.5,
                       align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_bg(tbl.cell(ri, ci), bg)
    return tbl


def _build_docx(overview, by_product, by_region, org, period, out_dir) -> Path:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"PSO — {org}")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = PSO_BLUE
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Business Overview  |  {period}")
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    o = overview.iloc[0]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(
        f"{org} generated PKR {o['GRS_CY_B']:.3f}Bn in gross revenue and "
        f"{o['Vol_CY_ML']:.2f}M litres of volume, with PKR {o['NMgn_CY_B']:.3f}Bn net margin. "
        f"Volume moved {_chg_str(o['Vol_Chg_Pct'])} and net margin moved {_chg_str(o['NMgn_Chg_Pct'])} vs SPLY."
    )
    r.font.size = Pt(10.5)

    if not by_product.empty:
        h = doc.add_paragraph()
        r = h.add_run("Product Mix")
        r.bold = True; r.font.size = Pt(13); r.font.color.rgb = PSO_BLUE
        _write_table(
            doc, by_product,
            headers=["Product", "Vol CY (ML)", "Vol SPLY (ML)", "Chg%", "GRS CY (B)", "Vol Share%"],
            cols=["ProductCategory", "Vol_CY_ML", "Vol_SPLY_ML", "Vol_Chg_Pct", "GRS_CY_B", "Vol_Share_Pct"],
            fmt={"Vol_CY_ML": lambda v: f"{v:.3f}", "Vol_SPLY_ML": lambda v: f"{v:.3f}",
                 "Vol_Chg_Pct": _chg_str, "GRS_CY_B": lambda v: f"{v:.3f}", "Vol_Share_Pct": lambda v: f"{v:.1f}%"},
        )

    if not by_region.empty:
        h = doc.add_paragraph()
        r = h.add_run("Region Split")
        r.bold = True; r.font.size = Pt(13); r.font.color.rgb = PSO_BLUE
        _write_table(
            doc, by_region,
            headers=["Region", "Vol CY (ML)", "Vol SPLY (ML)", "Chg%", "GRS CY (B)", "Vol Share%"],
            cols=["Sales office Region", "Vol_CY_ML", "Vol_SPLY_ML", "Vol_Chg_Pct", "GRS_CY_B", "Vol_Share_Pct"],
            fmt={"Vol_CY_ML": lambda v: f"{v:.3f}", "Vol_SPLY_ML": lambda v: f"{v:.3f}",
                 "Vol_Chg_Pct": _chg_str, "GRS_CY_B": lambda v: f"{v:.3f}", "Vol_Share_Pct": lambda v: f"{v:.1f}%"},
        )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(
        f"Data: PSO Working File, {org} channel. SPLY = Same Period Last Year "
        f"(like-for-like comparison, not full-year LY). {period}"
    )
    r.font.size = Pt(8); r.font.italic = True; r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    path = out_dir / f"PSO_{_safe_name(org)}_Summary_{period}.docx"
    doc.save(path)
    return path
