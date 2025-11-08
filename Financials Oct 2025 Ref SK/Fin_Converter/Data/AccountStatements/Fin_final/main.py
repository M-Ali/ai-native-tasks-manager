
import os
import re
import pandas as pd
from docx import Document

def parse_amount(amount_str):
    """
    Parses a string amount into a float.
    Handles amounts in parentheses for negative values.
    Returns None if parsing fails.
    """
    try:
        amount_str = amount_str.replace(',', '').replace('$', '')
        if amount_str.startswith('(') and amount_str.endswith(')'):
            return -float(amount_str[1:-1])
        return float(amount_str)
    except (ValueError, TypeError):
        return None

def extract_data_from_docx(file_path):
    """
    Extracts account summary and transaction data from a .docx file.
    """
    try:
        document = Document(file_path)
        
        filename = os.path.basename(file_path)
        month_match = re.search(r'([\w\.]+)\s(\d{4})', filename)
        if month_match:
            month, year = month_match.groups()
            month = month.replace('.', '')
        else:
            month, year = None, None

        account_summary = {'Year': year, 'Month': month}
        transactions = []

        # Table 0: Account Summary
        if len(document.tables) > 0:
            summary_table = document.tables[0]
            for row in summary_table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) > 2:
                    if 'Beginning Balance' in cells[1]:
                        account_summary['Beginning Balance (USD)'] = parse_amount(cells[2])
                    elif 'Checks' in cells[1]:
                        account_summary['Checks (Amount)'] = parse_amount(cells[2])
                    elif 'Withdrawals / Debits' in cells[1]:
                        account_summary['Withdrawals / Debits (USD)'] = parse_amount(cells[2])
                    elif 'Deposits / Credits' in cells[1]:
                        account_summary['Deposits / Credits (USD)'] = parse_amount(cells[2])
                    elif 'Ending Balance' in cells[1]:
                        account_summary['Ending Balance (USD)'] = parse_amount(cells[2])

        # Table 2: Checks
        if len(document.tables) > 2:
            checks_table = document.tables[2]
            for row in checks_table.rows[1:]: # Skip header
                cells = [cell.text.strip() for cell in row.cells]
                for i in range(0, len(cells), 3):
                    if len(cells) > i+2 and cells[i+1] and cells[i+2]:
                        amount = parse_amount(cells[i+2])
                        if amount is not None:
                            transactions.append({
                                'Year': year, 'Month': month, 'Type': 'Checks',
                                'Date': cells[i+1],
                                'Amount (USD)': amount,
                                'Description': f"Check {cells[i]}"
                            })

        # Table 3: Withdrawals / Debits and Deposits / Credits
        if len(document.tables) > 3:
            transaction_table = document.tables[3]
            transaction_type = 'Withdrawals / Debits'
            for row in transaction_table.rows[1:]: # Skip header
                cells = [cell.text.strip() for cell in row.cells]
                if 'Deposits / Credits' in cells[0]:
                    transaction_type = 'Deposits / Credits'
                    continue
                if len(cells) > 2 and cells[0] and cells[1]:
                    amount = parse_amount(cells[1])
                    if amount is not None:
                        transactions.append({
                            'Year': year, 'Month': month, 'Type': transaction_type,
                            'Date': cells[0],
                            'Amount (USD)': amount,
                            'Description': cells[2]
                        })

        return account_summary, transactions
    except Exception as e:
        print(f"Error processing file {file_path}: {e}")
        return None, None

def main():
    try:
        os.remove("account_summary.csv")
        os.remove("debit_credit_summary.csv")
    except FileNotFoundError:
        pass
        
    account_summaries = []
    all_transactions = []
    
    files_to_process = [f for f in os.listdir('.') if f.endswith('.docx') and not f.startswith('~$')]

    for docx_file in files_to_process:
        summary, transactions = extract_data_from_docx(docx_file)
        if summary:
            account_summaries.append(summary)
        if transactions:
            all_transactions.extend(transactions)

    account_summary_df = pd.DataFrame(account_summaries)
    debit_credit_summary_df = pd.DataFrame(all_transactions)

    account_summary_df.to_csv('account_summary.csv', index=False)
    debit_credit_summary_df.to_csv('debit_credit_summary.csv', index=False)

    print("Successfully created account_summary.csv and debit_credit_summary.csv")

if __name__ == "__main__":
    main()
